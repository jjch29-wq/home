"""
doc_template_filler.py - docxtpl 기반 절차서 템플릿 생성 도구

사용법:
  Word 문서에서 바꿀 위치에 {{ 변수명 }} 형식의 태그를 삽입
  예) {{ 개정번호 }}, {{ 검사일자 }}, {{ 담당자 }}

  반복 블록 (표 행 자동 생성):
    {% for item in 목록 %}
      {{ item.항목 }}  {{ item.값 }}
    {% endfor %}

  조건부 내용:
    {% if 고온배관 %}
      고온 배관 특수 조건 내용...
    {% endif %}
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import os
import json
import re
import datetime
import shutil
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm

PRESETS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "template_presets.json")


# ─────────────────────────────────────────────────────────────
# 태그 스캐너: 템플릿 파일에서 {{ 변수명 }} 자동 추출
# ─────────────────────────────────────────────────────────────

def scan_tags(template_path):
    """
    docx 파일에서 {{ 변수 }} 태그를 추출.
    반복/조건 블록은 제외하고 단순 변수만 반환.
    """
    from docx import Document
    doc = Document(template_path)
    tag_pattern = re.compile(r'\{\{\s*(\w+)\s*\}')

    tags = set()

    def _scan_text(text):
        for m in tag_pattern.finditer(text):
            tags.add(m.group(1))

    for para in doc.paragraphs:
        _scan_text(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _scan_text(para.text)

    for section in doc.sections:
        for container in [section.header, section.footer]:
            try:
                for para in container.paragraphs:
                    _scan_text(para.text)
                for table in container.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                _scan_text(para.text)
            except Exception:
                pass

    # 반복 블록용 변수 제외 (for/if 키워드)
    exclude = {'for', 'endfor', 'if', 'endif', 'else', 'elif', 'item'}
    return sorted(tags - exclude)


# ─────────────────────────────────────────────────────────────
# GUI
# ─────────────────────────────────────────────────────────────

class TemplateFiller:
    def __init__(self, root):
        self.root = root
        self.root.title("📄 절차서 템플릿 작성 도구 (docxtpl)")
        self.root.geometry("860x720")
        self.root.configure(bg="#f3f4f6")

        self.template_path = tk.StringVar()
        self.tag_rows = {}        # tag_name → (StringVar, frame)
        self.presets = {}

        self._load_presets()
        self._build_ui()

    # ── UI ───────────────────────────────────────────────────
    def _build_ui(self):
        style = ttk.Style()
        style.theme_use("clam")

        # ── 템플릿 파일 선택 ──────────────────────────────────
        tmpl_frame = ttk.LabelFrame(self.root, text=" 📂 템플릿 파일 (.docx) ", padding=10)
        tmpl_frame.pack(fill="x", padx=15, pady=(12, 5))

        tmpl_row = tk.Frame(tmpl_frame, bg="#f3f4f6")
        tmpl_row.pack(fill="x")
        ttk.Entry(tmpl_row, textvariable=self.template_path,
                  font=("Consolas", 9)).pack(side="left", fill="x", expand=True, padx=(0, 5))
        ttk.Button(tmpl_row, text="찾기", command=self._browse_template).pack(side="left")
        ttk.Button(tmpl_row, text="🔍 태그 자동 스캔",
                   command=self._scan).pack(side="left", padx=(8, 0))

        hint = tk.Label(tmpl_frame,
                        text="Word 문서에 {{ 변수명 }} 형태로 태그를 입력하면 아래에 자동으로 입력 칸이 생성됩니다.",
                        font=("Malgun Gothic", 9), fg="#555", bg="#f3f4f6")
        hint.pack(anchor="w", pady=(4, 0))

        # ── 태그 입력 영역 ────────────────────────────────────
        tag_outer = ttk.LabelFrame(self.root, text=" ✏️  변수 값 입력 ", padding=10)
        tag_outer.pack(fill="both", expand=True, padx=15, pady=5)

        # 스크롤 캔버스
        self.canvas = tk.Canvas(tag_outer, bg="#f3f4f6", highlightthickness=0)
        vscroll = ttk.Scrollbar(tag_outer, orient="vertical", command=self.canvas.yview)
        self.tag_container = tk.Frame(self.canvas, bg="#f3f4f6")
        self.tag_container.bind("<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self._cwin = self.canvas.create_window((0, 0), window=self.tag_container, anchor="nw")
        self.canvas.bind("<Configure>",
            lambda e: self.canvas.itemconfig(self._cwin, width=e.width))
        self.canvas.configure(yscrollcommand=vscroll.set)
        self.canvas.pack(side="left", fill="both", expand=True)
        vscroll.pack(side="right", fill="y")

        def _wheel(e):
            self.canvas.yview_scroll(int(-1*(e.delta/120)), "units")
        self.canvas.bind("<Enter>", lambda e: self.canvas.bind_all("<MouseWheel>", _wheel))
        self.canvas.bind("<Leave>", lambda e: self.canvas.unbind_all("<MouseWheel>"))

        # 수동 태그 추가
        manual_frame = tk.Frame(tag_outer, bg="#f3f4f6")
        manual_frame.pack(fill="x", pady=(6, 0))
        self._manual_tag = tk.StringVar()
        ttk.Entry(manual_frame, textvariable=self._manual_tag, width=20,
                  font=("Malgun Gothic", 10)).pack(side="left")
        ttk.Button(manual_frame, text="＋ 태그 직접 추가",
                   command=self._add_manual_tag).pack(side="left", padx=5)
        ttk.Button(manual_frame, text="전체 비우기",
                   command=self._clear_tags).pack(side="left", padx=5)

        # ── 프리셋 ────────────────────────────────────────────
        preset_frame = tk.Frame(self.root, bg="#f3f4f6")
        preset_frame.pack(fill="x", padx=15, pady=3)

        tk.Label(preset_frame, text="프리셋:", bg="#f3f4f6",
                 font=("Malgun Gothic", 9)).pack(side="left")
        self.preset_combo = ttk.Combobox(preset_frame, width=20, state="readonly")
        self.preset_combo.pack(side="left", padx=4)
        ttk.Button(preset_frame, text="불러오기", command=self._load_preset).pack(side="left", padx=2)
        ttk.Button(preset_frame, text="저장", command=self._save_preset).pack(side="left", padx=2)
        ttk.Button(preset_frame, text="삭제", command=self._delete_preset).pack(side="left", padx=2)
        self._refresh_preset_combo()

        # ── 옵션 + 실행 ───────────────────────────────────────
        run_frame = tk.Frame(self.root, bg="#f3f4f6")
        run_frame.pack(fill="x", padx=15, pady=5)

        self.suffix_var = tk.StringVar(value="_완성본")
        tk.Label(run_frame, text="저장 접미사:", bg="#f3f4f6",
                 font=("Malgun Gothic", 9)).pack(side="left")
        ttk.Entry(run_frame, textvariable=self.suffix_var, width=12).pack(side="left", padx=4)

        ttk.Button(run_frame, text="📋  미리보기",
                   command=self._preview).pack(side="right", padx=4)
        ttk.Button(run_frame, text="▶  문서 생성",
                   command=self._generate).pack(side="right", padx=4)
        ttk.Button(run_frame, text="다른 이름으로 저장",
                   command=lambda: self._generate(save_as=True)).pack(side="right", padx=4)

        # ── 로그 ─────────────────────────────────────────────
        log_frame = ttk.LabelFrame(self.root, text=" 📋 로그 ", padding=8)
        log_frame.pack(fill="x", padx=15, pady=(0, 10))
        self.log_text = tk.Text(log_frame, height=5, font=("Consolas", 9),
                                state="disabled", bg="white")
        self.log_text.pack(side="left", fill="both", expand=True)
        lsb = ttk.Scrollbar(log_frame, orient="vertical", command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=lsb.set)
        lsb.pack(side="right", fill="y")

    # ── 템플릿 스캔 ──────────────────────────────────────────
    def _browse_template(self):
        path = filedialog.askopenfilename(
            filetypes=[("Word 템플릿", "*.docx"), ("모든 파일", "*.*")])
        if path:
            self.template_path.set(path)
            self._scan()

    def _scan(self):
        path = self.template_path.get()
        if not path or not os.path.exists(path):
            messagebox.showwarning("경고", "먼저 템플릿 파일을 선택하세요.")
            return
        try:
            tags = scan_tags(path)
            if not tags:
                self._log("태그를 찾지 못했습니다. 문서에 {{ 변수명 }} 형태로 태그를 입력하세요.")
                return

            # 기존 값 보존 후 재생성
            old_vals = {name: var.get() for name, (var, _) in self.tag_rows.items()}
            self._clear_tags(confirm=False)
            for tag in tags:
                self._add_tag_row(tag, old_vals.get(tag, ""))
            self._log(f"태그 {len(tags)}개 감지: {', '.join(tags)}")
        except Exception as e:
            messagebox.showerror("오류", f"스캔 실패: {e}")

    # ── 태그 행 관리 ─────────────────────────────────────────
    def _add_tag_row(self, tag_name, default_val=""):
        if tag_name in self.tag_rows:
            return  # 중복 방지

        row = tk.Frame(self.tag_container, bg="#f3f4f6", pady=2)
        row.pack(fill="x", padx=5)

        # 태그명 라벨 (클릭 시 복사)
        lbl = tk.Label(row, text=f"{{{{ {tag_name} }}}}", width=22,
                       font=("Consolas", 10), bg="#e8f0fe", fg="#1a237e",
                       relief="groove", cursor="hand2", anchor="w", padx=6)
        lbl.pack(side="left")
        lbl.bind("<Button-1>", lambda e, t=tag_name: self._copy_tag(t))

        tk.Label(row, text="→", bg="#f3f4f6", fg="#777",
                 font=("Arial", 12)).pack(side="left", padx=6)

        var = tk.StringVar(value=default_val)

        # 값이 길 수 있으면 멀티라인 토글 버튼
        is_multi = tk.BooleanVar(value=False)

        entry_frame = tk.Frame(row, bg="#f3f4f6")
        entry_frame.pack(side="left", fill="x", expand=True)

        single_entry = ttk.Entry(entry_frame, textvariable=var,
                                 font=("Malgun Gothic", 10))
        single_entry.pack(fill="x")

        multi_text = tk.Text(entry_frame, height=4, font=("Malgun Gothic", 10),
                             wrap="word", relief="solid", borderwidth=1)
        multi_text.insert("1.0", default_val)

        def _sync_multi_to_var(*_):
            var.set(multi_text.get("1.0", tk.END).rstrip("\n"))
        multi_text.bind("<KeyRelease>", _sync_multi_to_var)

        def _toggle_multi():
            if is_multi.get():
                single_entry.pack_forget()
                multi_text.pack(fill="x")
                multi_text.delete("1.0", tk.END)
                multi_text.insert("1.0", var.get())
            else:
                multi_text.pack_forget()
                single_entry.pack(fill="x")

        btn_multi = ttk.Button(row, text="↕", width=3,
                               command=lambda: [is_multi.set(not is_multi.get()), _toggle_multi()])
        btn_multi.pack(side="left", padx=2)

        def _remove():
            row.destroy()
            self.tag_rows.pop(tag_name, None)

        ttk.Button(row, text="✕", width=3, command=_remove).pack(side="left", padx=2)
        self.tag_rows[tag_name] = (var, row)

    def _add_manual_tag(self):
        tag = self._manual_tag.get().strip()
        tag = re.sub(r'[^a-zA-Z0-9가-힣_]', '', tag)
        if not tag:
            return
        self._add_tag_row(tag)
        self._manual_tag.set("")

    def _clear_tags(self, confirm=True):
        if confirm and self.tag_rows:
            if not messagebox.askyesno("확인", "모든 태그 입력 칸을 비울까요?"):
                return
        for _, (_, frame) in self.tag_rows.items():
            frame.destroy()
        self.tag_rows.clear()

    def _copy_tag(self, tag_name):
        self.root.clipboard_clear()
        self.root.clipboard_append(f"{{{{{tag_name}}}}}")
        self._log(f"클립보드 복사: {{{{ {tag_name} }}}}")

    # ── 문서 생성 ─────────────────────────────────────────────
    def _get_context(self):
        return {name: var.get() for name, (var, _) in self.tag_rows.items()}

    def _generate(self, save_as=False):
        path = self.template_path.get()
        if not path or not os.path.exists(path):
            messagebox.showwarning("경고", "템플릿 파일을 선택하세요.")
            return

        context = self._get_context()
        if not any(v.strip() for v in context.values()):
            if not messagebox.askyesno("확인", "입력된 값이 없습니다. 계속하시겠습니까?"):
                return

        # 저장 경로 결정
        base, ext = os.path.splitext(path)
        suffix = self.suffix_var.get() or "_완성본"
        default_save = base + suffix + ext

        if save_as:
            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word 문서", "*.docx")],
                initialfile=os.path.basename(default_save))
            if not save_path:
                return
        else:
            save_path = default_save

        try:
            tpl = DocxTemplate(path)
            tpl.render(context)
            tpl.save(save_path)
            self._log(f"✅ 저장 완료: {os.path.basename(save_path)}")
            if messagebox.askyesno("완료", f"저장 완료!\n{save_path}\n\n파일을 열어볼까요?"):
                os.startfile(save_path)
        except Exception as e:
            self._log(f"❌ 오류: {e}")
            messagebox.showerror("오류", f"생성 실패:\n{e}")

    # ── 미리보기 ─────────────────────────────────────────────
    def _preview(self):
        path = self.template_path.get()
        if not path or not os.path.exists(path):
            messagebox.showwarning("경고", "템플릿 파일을 선택하세요.")
            return

        context = self._get_context()

        win = tk.Toplevel(self.root)
        win.title("미리보기 (텍스트)")
        win.geometry("700x550")

        txt = scrolledtext.ScrolledText(win, font=("Malgun Gothic", 10), wrap="word")
        txt.pack(fill="both", expand=True, padx=10, pady=10)

        try:
            from docx import Document
            doc = Document(path)

            txt.insert(tk.END, "=== 입력된 변수 값 ===\n")
            for k, v in context.items():
                txt.insert(tk.END, f"  {{{{ {k} }}}} = {v!r}\n")
            txt.insert(tk.END, "\n=== 적용 후 본문 (텍스트만) ===\n\n")

            def _replace_preview(text):
                for k, v in context.items():
                    text = text.replace(f"{{{{{k}}}}}", v)
                    text = text.replace(f"{{{{ {k} }}}}", v)
                return text

            for para in doc.paragraphs:
                line = _replace_preview(para.text)
                if line.strip():
                    txt.insert(tk.END, line + "\n")

            for table in doc.tables:
                txt.insert(tk.END, "\n[표]\n")
                for row in table.rows:
                    cells = [_replace_preview(c.text.strip()) for c in row.cells]
                    txt.insert(tk.END, " | ".join(cells) + "\n")

        except Exception as e:
            txt.insert(tk.END, f"오류: {e}")

        txt.config(state="disabled")

        ttk.Button(win, text="💾 이 내용으로 저장",
                   command=lambda: self._generate(save_as=True)).pack(pady=8)

    # ── 프리셋 ────────────────────────────────────────────────
    def _load_presets(self):
        if os.path.exists(PRESETS_FILE):
            try:
                with open(PRESETS_FILE, "r", encoding="utf-8") as f:
                    self.presets = json.load(f)
            except Exception:
                self.presets = {}

    def _save_presets_file(self):
        with open(PRESETS_FILE, "w", encoding="utf-8") as f:
            json.dump(self.presets, f, ensure_ascii=False, indent=2)

    def _refresh_preset_combo(self):
        self.preset_combo["values"] = list(self.presets.keys())

    def _save_preset(self):
        name = self._ask_string("프리셋 이름:", default="절차서_기본")
        if not name:
            return
        self.presets[name] = {
            "template": self.template_path.get(),
            "values": self._get_context()
        }
        self._save_presets_file()
        self._refresh_preset_combo()
        self.preset_combo.set(name)
        self._log(f"프리셋 '{name}' 저장 완료")

    def _load_preset(self):
        name = self.preset_combo.get()
        if not name or name not in self.presets:
            return
        data = self.presets[name]
        if "template" in data and os.path.exists(data["template"]):
            self.template_path.set(data["template"])
            self._scan()
        vals = data.get("values", {})
        for tag, (var, _) in self.tag_rows.items():
            if tag in vals:
                var.set(vals[tag])
        # 저장된 값 중 새 태그 추가
        for tag, val in vals.items():
            if tag not in self.tag_rows:
                self._add_tag_row(tag, val)
        self._log(f"프리셋 '{name}' 불러오기 완료")

    def _delete_preset(self):
        name = self.preset_combo.get()
        if name and name in self.presets:
            del self.presets[name]
            self._save_presets_file()
            self._refresh_preset_combo()
            self.preset_combo.set("")
            self._log(f"프리셋 '{name}' 삭제")

    # ── 유틸 ─────────────────────────────────────────────────
    def _log(self, msg):
        ts = datetime.datetime.now().strftime("%H:%M:%S")
        self.log_text.config(state="normal")
        self.log_text.insert(tk.END, f"[{ts}] {msg}\n")
        self.log_text.see(tk.END)
        self.log_text.config(state="disabled")

    def _ask_string(self, prompt, default=""):
        win = tk.Toplevel(self.root)
        win.title("입력")
        win.geometry("320x110")
        win.grab_set()
        tk.Label(win, text=prompt, font=("Malgun Gothic", 10)).pack(pady=(15, 5))
        var = tk.StringVar(value=default)
        entry = ttk.Entry(win, textvariable=var, width=30)
        entry.pack()
        entry.select_range(0, tk.END)
        entry.focus()
        result = [None]

        def _ok(e=None):
            result[0] = var.get().strip()
            win.destroy()

        entry.bind("<Return>", _ok)
        ttk.Button(win, text="확인", command=_ok).pack(pady=8)
        self.root.wait_window(win)
        return result[0]


# ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    root = tk.Tk()
    app = TemplateFiller(root)
    root.mainloop()
