from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk
from PIL import Image, ImageTk
import os
import shutil
import json
import copy
import subprocess
from datetime import datetime

def iter_block_items(parent, document_ref=None):
    """문단과 표를 원래 순서대로 순회"""
    parent_element = parent.element.body if hasattr(parent, 'element') and hasattr(parent.element, 'body') else parent._element
    owner = document_ref or parent

    for child in parent_element.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, owner)
        elif child.tag == qn('w:tbl'):
            yield Table(child, owner)

def extract_block_content(container, area='body'):
    """본문/머릿글/바닥글의 문단과 표를 구조적으로 추출"""
    items = []

    for block in iter_block_items(container):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                items.append({
                    'type': 'text',
                    'text': text,
                    'style': getattr(block.style, 'name', 'Normal'),
                    'area': area
                })

                if area != 'body':
                    continue

            if area == 'body':
                for run in block.runs:
                    if run._element.xpath('.//a:blip'):
                        blip = run._element.xpath('.//a:blip')[0]
                        rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if rId in container.part.rels:
                            rel = container.part.rels[rId]
                            image_part = rel.target_part
                            image_data = image_part.blob
                            content_type = image_part.content_type
                            ext = '.jpg' if 'jpeg' in content_type else '.png' if 'png' in content_type else '.gif'
                            temp_dir = os.path.join(os.path.dirname(container.part.package.part_related_by(rId).partname if False else ''), '.temp_images')
        elif isinstance(block, Table):
            table_data = [[cell.text.strip() for cell in row.cells] for row in block.rows]
            if table_data:
                items.append({
                    'type': 'table',
                    'data': table_data,
                    'area': area
                })

    return items

def add_bordered_table(target, data):
    """지정한 컨테이너에 실선 테두리 표 추가"""
    if not data:
        return None

    if hasattr(target, 'add_table') and hasattr(target, 'sections'):
        table = target.add_table(rows=len(data), cols=len(data[0]))
    else:
        try:
            table = target.add_table(rows=len(data), cols=len(data[0]), width=Inches(6.5))
        except TypeError:
            table = target.add_table(rows=len(data), cols=len(data[0]))

    table.style = 'Table Grid'
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = cell_text
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="12" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="12" w:space="0" w:color="000000"/></w:tcBorders>')
            tcPr.append(tcBorders)
    return table

def clear_story_container(container):
    """머릿글/바닥글의 기존 문단/표 제거"""
    for element in list(container._element):
        if element.tag in (qn('w:p'), qn('w:tbl')):
            container._element.remove(element)

    if not container.paragraphs:
        container.add_paragraph()

def safe_style_name(paragraph):
    """문단 스타일명 안전 조회 (일부 문서의 스타일 오류 회피)"""
    try:
        style = paragraph.style
        if style is not None and hasattr(style, 'name'):
            return style.name
    except Exception:
        pass
    return 'Normal'

def load_existing_doc(file_path):
    """Word 문서에서 모든 내용 추출 (문단 및 이미지 순서대로)"""
    doc = Document(file_path)
    temp_dir = os.path.join(os.path.dirname(file_path), ".temp_images")
    
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)
    
    all_content = []
    image_counter = 0

    try:
        header_items = []
        for block in iter_block_items(doc.sections[0].header):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if text:
                    header_items.append({
                        'type': 'text',
                        'text': text,
                        'style': safe_style_name(block),
                        'area': 'header'
                    })
            elif isinstance(block, Table):
                table_data = [[cell.text.strip() for cell in row.cells] for row in block.rows]
                if table_data:
                    header_items.append({
                        'type': 'table',
                        'data': table_data,
                        'area': 'header'
                    })
        all_content.extend(header_items)
    except:
        pass

    for element in doc.element.body:
        if element.tag == qn('w:p'):
            para = Paragraph(element, doc)
            if para.text.strip():
                all_content.append({
                    'type': 'text',
                    'text': para.text,
                    'style': safe_style_name(para),
                    'area': 'body'
                })

            for run in para.runs:
                if run._element.xpath('.//a:blip'):
                    blip = run._element.xpath('.//a:blip')[0]
                    rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if rId in doc.part.rels:
                        rel = doc.part.rels[rId]
                        image_part = rel.target_part
                        image_data = image_part.blob
                        content_type = image_part.content_type
                        ext = '.jpg' if 'jpeg' in content_type else '.png' if 'png' in content_type else '.gif'
                        image_file = os.path.join(temp_dir, f"image_{image_counter}{ext}")
                        with open(image_file, 'wb') as f:
                            f.write(image_data)
                        all_content.append({
                            'type': 'image',
                            'path': image_file,
                            'element': element,  # 원본 w:p 요소 (deepcopy 사용 시 관계 보존)
                            'area': 'body'
                        })
                        image_counter += 1
        elif element.tag == qn('w:tbl'):
            table = Table(element, doc)
            table_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if table_data:
                all_content.append({
                    'type': 'table',
                    'data': table_data,
                    'element': element,  # 병합 셀·서식 보존을 위해 원본 XML 요소 저장
                    'area': 'body'
                })

    try:
        footer_items = []
        for block in iter_block_items(doc.sections[0].footer):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if text:
                    footer_items.append({
                        'type': 'text',
                        'text': text,
                        'style': safe_style_name(block),
                        'area': 'footer'
                    })
            elif isinstance(block, Table):
                table_data = [[cell.text.strip() for cell in row.cells] for row in block.rows]
                if table_data:
                    footer_items.append({
                        'type': 'table',
                        'data': table_data,
                        'area': 'footer'
                    })
        all_content.extend(footer_items)
    except:
        pass
    
    return all_content

def extract_images_from_doc(file_path):
    """Word 문서에서 이미지 추출"""
    doc = Document(file_path)
    image_paths = []
    temp_dir = os.path.join(os.path.dirname(file_path), ".temp_images")
    
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)
    
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_part = rel.target_part
                image_data = image_part.blob
                content_type = image_part.content_type
                ext = '.jpg' if 'jpeg' in content_type else '.png' if 'png' in content_type else '.gif'
                image_file = os.path.join(temp_dir, f"image_{len(image_paths)}{ext}")
                with open(image_file, 'wb') as f:
                    f.write(image_data)
                image_paths.append(image_file)
    except Exception as e:
        print(f"이미지 추출 오류: {e}")
    
    return image_paths

def create_ndt_procedure_doc(paragraphs, image_paths=None):
    """문단 정보와 이미지로 새로운 문서 생성"""
    doc = Document()
    
    for para_info in paragraphs:
        text = para_info.get('text', '').strip()
        if not text:
            continue
            
        style = para_info.get('style', 'Normal')
        
        # Heading 스타일 처리
        if 'Heading' in style:
            try:
                level = int(style.split()[-1])
            except:
                level = 1
            doc.add_heading(text, level=level)
        else:
            doc.add_paragraph(text, style=style)
    
    # 이미지 추가
    if image_paths:
        doc.add_heading('첨부 사진', level=1)
        for img_path in image_paths:
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(5))
                    doc.add_paragraph()
                except:
                    pass
    
    return doc

class NDTProcedureApp:
    CONFIG_FILE = "app_config.json"
    
    def __init__(self, root):
        self.root = root
        self.root.title("비파괴 검사 절차서 생성 및 관리 시스템")
        self.root.resizable(True, True)
        
        # 저장된 창 크기 복원 또는 기본값 사용
        geometry = self.load_window_geometry()
        self.root.geometry(geometry)
        
        # 창 크기 변경 이벤트 바인드
        self.root.bind('<Configure>', self.on_window_configure)
        self.root.bind('<Control-z>', lambda e: self.undo())
        self.root.bind('<Control-y>', lambda e: self.redo())
        self.root.bind('<Control-Z>', lambda e: self.undo())
        self.root.bind('<Control-Y>', lambda e: self.redo())
        
        self.paragraphs = []
        self.image_paths = []
        self.content = []
        self.source_file = None  # 원본 Word 파일 경로 (바닥글/헤더 이미지 보존용)
        self._undo_stack = []   # (content deepcopy, source_file) 최대 30단계
        self._redo_stack = []
        self._tree_filter_var = tk.StringVar()
        self.standards = {
            "ASME Section V, Article 4 (PAUT 기본 절차)":
                "ASME Section V, Article 4 - Ultrasonic Examination Methods (PAUT 기본 절차)\n\n적용 범위:\n위상배열 초음파검사(PAUT)의 핵심 절차 요구사항을 규정하는 기본 코드.\n\n주요 요구사항:\n- 위상배열 프로브 사양 및 선정 기준\n- 스캔 계획(Scan Plan) 수립 및 시뮬레이션\n- 보정 블록(Calibration Block) 규격 및 보정 절차\n- 감도 설정 및 DAC/TCG 적용\n- 결함 탐지, 위치, 크기 측정 기준\n- 기록 요구사항: A-scan, S-scan 데이터 보존\n\n검사원 자격: ASNT SNT-TC-1A 또는 CP-189 기준 Level II 이상\n보정 주기: 검사 전·후 및 8시간마다 보정 확인 필요",
            "ASME Section V, Article 4, Appendix III (PAUT 전용 요구사항)":
                "ASME Section V, Article 4, Appendix III - Phased Array Ultrasonic Examination\n\n적용 범위:\n위상배열(Phased Array) 전용 부록으로, PA 시스템 고유 요구사항 규정.\n\n주요 요구사항:\n- 초점 법칙(Focal Law) 설계 및 검증\n- 섹터 스캔(S-scan) 각도 범위 및 해상도 설정\n- 선형 스캔(Linear Scan) 인덱스 포인트 설정\n- 유효 빔 프로파일 검증 (Beam Profile Verification)\n- 개구수(Aperture) 및 소자 수 설정 기준\n- 데이터 획득 설정: 피치, 펄스 반복 주파수\n\n비고: Article 4 본문과 함께 적용 필수",
            "ASME Section VIII Div.1, Appendix 12 (PAUT - 압력용기)":
                "ASME Section VIII Division 1, Mandatory Appendix 12\n"
                "Ultrasonic Examination of Welds — 압력용기 용접부 RT 대체 UT/PAUT 적용 기준\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 적용 범위\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- 적용 대상: ASME Sec. VIII Div.1 압력용기 맞대기 용접부\n"
                "- 목적: 방사선투과검사(RT)를 초음파검사(UT / PAUT)로 대체\n"
                "- 적용 두께: t ≥ 1/2\" (13 mm) 이상 용접부\n"
                "- 근거 조항: UW-11(a)(4), UW-53, Appendix 12\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ RT 대체 적용 조건 (UW-11(a)(4))\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "RT 대체가 허용되는 경우:\n"
                "  ① 재료가 RT 불가 형상·재질인 경우\n"
                "  ② 방사선 접근 불가 환경인 경우\n"
                "  ③ PAUT 동등 검사 능력 기술 입증 및 절차 승인 완료 시\n\n"
                "RT 대체 시 필수 문서:\n"
                "  - 서면 UT/PAUT 절차서 (Written Procedure)\n"
                "  - 절차 검증 기록 (Procedure Qualification Record)\n"
                "  - 발주처 / AI (Authorized Inspector) 서면 승인\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 검사 범위\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- 용접부 전체 체적 (Full Volumetric Coverage)\n"
                "- 열영향부 (HAZ) 포함\n"
                "- 최소 2방향 각도 스캔으로 용접 단면 100% 커버\n"
                "- 루트부 / 캡부 / 측벽 용합 불량 탐지 각도 포함 필수\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 보정 블록 및 감도 (App. 12, Para. 12-4)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "기준 반사체:\n"
                "  - SDH Ø 3/64\" (≈ 1.2 mm)\n"
                "  - 배치 깊이: t/4, t/2, 3t/4\n\n"
                "감도 기준:\n"
                "  - 기록 수준: DAC 20% (−14 dB)\n"
                "  - 평가 수준: DAC 50% (−6 dB)\n"
                "  - 거부 수준: DAC 100% 초과\n\n"
                "보정 블록 재질:\n"
                "  - 검사 대상과 동일 P-Number 재질\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 수용 기준 (Table UW-53)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "불합격 지시 (Rejectable Indications):\n"
                "  ① 균열 (Crack): 크기·길이 무관 전부 불합격\n"
                "  ② 용합 불량 / 용입 불량:\n"
                "     - t < 3/4\" (19 mm)  : 누적 길이 > 1\" (25 mm) 불합격\n"
                "     - 3/4\" ≤ t < 2-1/4\" : 누적 길이 > t/3 불합격\n"
                "     - t ≥ 2-1/4\" (57 mm): 누적 길이 > 3/4\" (19 mm) 불합격\n"
                "     (임의 12\" (300 mm) 구간 내 누적 기준)\n"
                "  ③ 내부 결함 (기공·슬래그 등):\n"
                "     - 개별 지시 길이 > 1/4\" (6 mm) 불합격\n"
                "     - 임의 6\" (150 mm) 내 합계 > 1/2\" (12 mm) 불합격\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 탐촉자 선정 기준\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- 주파수: 2 ~ 5 MHz (재질·두께에 따라 선택)\n"
                "- 굴절각: 용접 형상·결함 방향에 따라 45°~70° 복수 각도\n"
                "- PAUT: S-scan으로 각도 범위 커버, Focal Law 사전 검증\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 검사원 자격\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- ASNT SNT-TC-1A 또는 CP-189 기준 UT Level II 이상\n"
                "- 합부판정: UT Level II/III 자격자만 수행\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 기록 요건\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- 절차서 번호, 장비 시리얼, 검사원 자격\n"
                "- 결함 위치 (스캔 축, 깊이), 최대 진폭 (%DAC)\n"
                "- 크기 측정 방법 및 결과\n"
                "- 합부 판정 결과 및 근거 조항 (UW-53)\n"
                "- A-scan / S-scan 디지털 데이터 보존",
            "ASME Section VIII Div.2, Para. 7.5.5 (PAUT - 고압용기)":
                "ASME Section VIII Division 2, Paragraph 7.5.5 - Ultrasonic Examination\n"
                "고압용기 용접부 PAUT 체적 검사 요구사항\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 적용 범위\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- 적용 대상: ASME Sec. VIII Div.2 고압용기 완전 용입 맞대기 용접부\n"
                "- Div.2는 Div.1 대비 설계 허용 응력을 높게 허용하는 대신\n"
                "  검사·제조 요건이 더 엄격하게 적용됨\n"
                "- 근거 조항: Para. 7.5.5, Table 7.5.5-1\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 검사 범위\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- 전체 용접 체적 (Full Volumetric) + HAZ 100% 검사 필수\n"
                "- 루트부 / 캡부 / 측벽 용합 불량 탐지 각도 모두 포함\n"
                "- 스캔 커버리지: 용접 단면 100% 입증 필수 (Scan Plan 문서화)\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 절차 요건\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- 서면 절차서: ASME Sec. V Article 4 + Appendix III 준수\n"
                "- 스캔 계획 (Scan Plan): 프로브 배치, 각도, 인덱스 포인트 명시\n"
                "- 절차 검증 (Procedure Qualification): 보정 블록으로 사전 검증\n"
                "- Focal Law 검증: 설계된 각도·깊이 범위에서 SDH 탐지 확인\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 보정 블록 및 감도 (Table 7.5.5-1 기준)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "기준 반사체:\n"
                "  - SDH Ø 2.0 mm\n"
                "  - 배치 깊이: t/4, t/2, 3t/4\n"
                "  - 재질: 검사 대상과 동일 P-Number\n\n"
                "감도 기준:\n"
                "  - 기준 에코: SDH Ø 2.0 mm = 80% FSH (Full Screen Height)\n"
                "  - 기록 수준: DAC 20% (−14 dB)\n"
                "  - 평가 수준: DAC 50% (−6 dB)\n"
                "  - 검사 감도: 평가 수준 +6 dB 증폭 후 스캔\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 수용 기준 (Table 7.5.5-1)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "【평면형 결함 (Planar Defects)】\n"
                "  - 균열 / 용합 불량 / 용입 불량: 크기 무관 전부 불합격\n\n"
                "【체적형 결함 (Volumetric Defects)】\n"
                "  결함 높이(a) 및 길이(ℓ) 기준:\n"
                "  ┌────────────────┬───────────────────┐\n"
                "  │ 결함 높이 (a)   │ 허용 길이 (ℓ)      │\n"
                "  ├────────────────┼───────────────────┤\n"
                "  │ a ≤ 3 mm       │ ℓ ≤ 6 mm          │\n"
                "  │ 3 mm < a ≤ 6 mm│ ℓ ≤ 2a            │\n"
                "  │ a > 6 mm       │ 불합격              │\n"
                "  └────────────────┴───────────────────┘\n\n"
                "【표면 연결 결함 (Surface-Breaking)】\n"
                "  ┌────────────────┬───────────────────┐\n"
                "  │ 결함 높이 (a)   │ 허용 길이 (ℓ)      │\n"
                "  ├────────────────┼───────────────────┤\n"
                "  │ a ≤ 1.5 mm     │ ℓ ≤ 6 mm          │\n"
                "  │ a > 1.5 mm     │ 불합격              │\n"
                "  └────────────────┴───────────────────┘\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 결함 크기 측정 요건\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- 결함 높이(a): −6 dB drop법 또는 TOFD 병행 (정확도 ±1 mm)\n"
                "- 결함 길이(ℓ): −6 dB drop법 또는 −20 dB 끝점법\n"
                "- 위치 정확도: 스캔 축·인덱스 축 ±1 mm 이내\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 탐촉자 선정 기준\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- 주파수: 2 ~ 5 MHz (두께·재질에 따라 선택)\n"
                "- 굴절각: 용접 형상·결함 방향에 따라 45°~70° 복수 적용\n"
                "- PAUT S-scan: 각도 범위 내 Focal Law 사전 검증 필수\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ Div.1 vs Div.2 주요 차이점\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "┌──────────────────┬──────────────────┬──────────────────┐\n"
                "│ 항목              │ Div.1 App.12     │ Div.2 Para.7.5.5 │\n"
                "├──────────────────┼──────────────────┼──────────────────┤\n"
                "│ 기준 반사체 SDH   │ Ø 3/64\" (1.2mm) │ Ø 2.0 mm         │\n"
                "│ 수용 기준 조항    │ Table UW-53      │ Table 7.5.5-1    │\n"
                "│ 설계 허용 응력    │ 낮음 (보수적)     │ 높음 (엄격 검사) │\n"
                "│ 체적형 결함 기준  │ 길이 기반         │ 높이(a)+길이(ℓ)  │\n"
                "└──────────────────┴──────────────────┴──────────────────┘\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 검사원 자격\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- ASNT SNT-TC-1A 또는 CP-189 기준 UT Level II 이상\n"
                "- 합부판정: UT Level II/III 자격자만 수행\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 기록 요건\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- A-scan + S-scan / B-scan 디지털 데이터 전체 보존\n"
                "- 결함 위치 (스캔 축, 인덱스 축, 깊이 상단·하단)\n"
                "- 결함 크기 (높이 a, 길이 ℓ), 최대 진폭 (%DAC)\n"
                "- 크기 측정 방법 (−6 dB / TOFD 등)\n"
                "- 합부 판정 결과 및 근거 조항 (Table 7.5.5-1)",
            "ASME B31.1 (PAUT - 발전 배관)":
                "ASME B31.1 Power Piping - PAUT 적용 요구사항\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 적용 범위\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- 적용 대상: 발전소 배관 용접부 (증기, 급수, 블로우다운 등)\n"
                "- 근거 조항: B31.1 Para. 136.4, Table 136.4.1\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 보정 블록 (Calibration Block)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "B31.1은 보정 블록으로 두 가지를 허용:\n\n"
                "【1】 IIW Block (International Institute of Welding Block)\n"
                "  - 규격: ISO 2400\n"
                "  - 구조물:\n"
                "    · Ø 1.5 mm FBH (Flat Bottom Hole) — 거리 진폭 교정용\n"
                "    · R25 / R100 mm 곡면 — 음속 측정 및 영점 설정용\n"
                "    · 50 mm 두께 범위 홈 — 굴절각 확인용 (45°/60°/70°)\n"
                "  - 주요 용도:\n"
                "    · 탐촉자 굴절각 확인\n"
                "    · 음속(Velocity) 측정\n"
                "    · 영점(Zero Offset) 설정\n"
                "    · 표면 파형 및 빔 프로파일 확인\n"
                "  - 재질: 탄소강 (ASME SA-36 동등)\n"
                "  - 제한: 현장 두께별 SDH 기반 DAC 설정 불가 → 감도 설정 시 ASME Basic Block 병행 필요\n\n"
                "【2】 ASME Basic Calibration Block\n"
                "  - 규격: ASME Section V, Article 4, Table T-434.2.1\n"
                "  - 구조물: SDH (Side Drilled Hole) — 두께별 직경 기준 적용\n"
                "    ┌─────────────────────┬────────────────────┐\n"
                "    │ 검사 두께 (t)        │ SDH 직경            │\n"
                "    ├─────────────────────┼────────────────────┤\n"
                "    │ t ≤ 1\"  (25 mm)     │ Ø 3/32\" (2.4 mm)  │\n"
                "    │ 1\" < t ≤ 2\" (50 mm)│ Ø 1/8\"  (3.2 mm)  │\n"
                "    │ 2\" < t ≤ 4\" (100mm)│ Ø 3/16\" (4.8 mm)  │\n"
                "    │ t > 4\"  (100 mm↑)  │ Ø 1/4\"  (6.4 mm)  │\n"
                "    └─────────────────────┴────────────────────┘\n"
                "  - SDH 배치 깊이: t/4, t/2, 3t/4 (최소 3개 지점)\n"
                "  - 주요 용도:\n"
                "    · DAC (Distance Amplitude Correction) 곡선 작성\n"
                "    · 검사 감도 (Search Sensitivity) 설정\n"
                "    · 거리별 진폭 보정 기준 확립\n"
                "  - 재질: 검사 대상과 동일 P-Number 재질\n\n"
                "【블록 선택 기준 요약】\n"
                "  ┌──────────────────┬────────────────┬──────────────────┐\n"
                "  │ 용도              │ IIW Block       │ ASME Basic Block │\n"
                "  ├──────────────────┼────────────────┼──────────────────┤\n"
                "  │ 굴절각 확인       │ ✔ 적합          │ 불가             │\n"
                "  │ 음속 / 영점       │ ✔ 적합          │ 불가             │\n"
                "  │ DAC 감도 설정     │ 제한적          │ ✔ 적합           │\n"
                "  │ 두께별 SDH 탐지   │ 불가            │ ✔ 적합           │\n"
                "  └──────────────────┴────────────────┴──────────────────┘\n"
                "  → 실무: IIW로 탐촉자 검증 + ASME Basic Block으로 DAC 설정\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 감도 기준\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- DAC 작성 후 검사 감도: +6 dB 추가 (Search Level)\n"
                "- 기록 수준: DAC 20% (−14 dB)\n"
                "- 평가 수준: DAC 50% (−6 dB)\n"
                "- 거부 수준: DAC 100% 초과\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 보정 주기\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- 검사 시작 전 / 매 8시간마다 / 검사 종료 후\n"
                "- 장비 교체·온도 차 ±14°C 초과 시 재보정\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 절차 / 수용 기준 / 자격\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- 절차: ASME Sec. V Article 4 + Appendix III 기반 서면 절차\n"
                "- 수용 기준: Table 136.4.1 (결함 유형별 기준)\n"
                "- RT 대체: 동등 검사 능력 기술 입증 시 PAUT로 RT 대체 가능\n"
                "- 자격: ASNT SNT-TC-1A Level II/III\n"
                "- 기록: 결함 위치, 크기, 평가 결과, 검사 조건 문서화\n"
                "- 보고서: 절차, 장비, 인원, 결과 및 판정 포함",
            "ASME B31.3 (PAUT - 공정 배관)":
                "ASME B31.3 Process Piping - PAUT 적용 요구사항\n\n적용 범위:\n석유화학·정유·가스 플랜트 공정 배관 용접부 PAUT 검사.\n\n주요 요구사항:\n- 배관 등급별 검사 범위:\n  · Normal Fluid Service: 용접부의 5% 이상\n  · Category M / High Pressure: 100% 검사\n- 절차: ASME Section V Article 4 준수\n- 수용 기준: Table 341.3.2 적용\n- RT 대체: PAUT로 RT 대체 시 동등 이상 감도 입증 필요\n- 자격: ASNT SNT-TC-1A Level II 이상\n- 결과 기록: 검사 부위, 결함 지시, 판정 결과 문서화",
            "ASME Section XI (PAUT - 원자력 가동 중 검사)":
                "ASME Section XI - Rules for Inservice Inspection of Nuclear Power Plant Components\n\n적용 범위:\n원자력 발전소 핵심 기기·배관·용기의 가동 중 검사(ISI).\n\n주요 요구사항:\n- 검사 주기: IWB/IWC/IWD 조항별 10년 주기 검사 계획\n- PAUT 적용: Appendix VIII (성능 실증, PDI) 필수 통과\n- 절차 검증: Performance Demonstration Initiative (PDI) 인증\n- 적용 부위: 반응기 압력용기, 1차 배관, 노심 내부 구조물\n- 결함 크기 기준: IWB-3500 시리즈 수용 기준\n- 자격: ASNT Level III + PDI 자격 보유자 감독 하 수행\n- 기록: 전체 스캔 데이터 10년 이상 보존",
            "ASME Section I (PAUT - 보일러)":
                "ASME Section I - Power Boilers (PAUT 적용)\n\n적용 범위:\n발전용 보일러 동체, 헤더, 고온 배관 용접부 검사.\n\n주요 요구사항:\n- 적용 부위: 보일러 동체 용접부, 드럼, 노즐, 헤더 용접부\n- 절차: ASME Section V Article 4 기반\n- 수용 기준: PW-51 조항 적용\n- 두께 범위: 주로 25mm 이상 후판 용접부에 PAUT 적용\n- 보정: 검사 두께에 맞는 ASME 보정 블록 사용\n- 자격: ASNT SNT-TC-1A Level II 이상\n- 기록: 스캔 데이터 및 판정 결과 보존",
            "ASME Sec. VIII Div.2 Para. 7.5.5 (Ultrasonic Examination)": 
                "ASME Section VIII Division 2, Paragraph 7.5.5: Ultrasonic Examination\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 원문 (Original Text)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "\"Ultrasonic examination shall be performed in accordance with the\n"
                "requirements of Article 7.5. The examination shall be conducted by\n"
                "qualified personnel using calibrated equipment. Acceptance criteria\n"
                "shall meet the requirements of Table 7.5.5-1.\"\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 조항 해설\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "① \"in accordance with the requirements of Article 7.5\"\n"
                "   → Article 7.5 전체가 상위 요건\n"
                "   - Para. 7.5.1: 일반 요건 (검사 범위, 절차 승인)\n"
                "   - Para. 7.5.2: 검사원 자격 (ASNT Level II/III)\n"
                "   - Para. 7.5.3: 검사 장비 및 보정 요건\n"
                "   - Para. 7.5.4: RT 요건\n"
                "   - Para. 7.5.5: UT / PAUT 요건 (본 조항)\n\n"
                "② \"qualified personnel\"\n"
                "   → ASNT SNT-TC-1A 또는 CP-189 기준 UT Level II 이상\n"
                "   → 합부판정은 반드시 Level II / III 자격자만 수행\n\n"
                "③ \"calibrated equipment\"\n"
                "   → 검사 전·후 및 매 8시간마다 보정 확인 필수\n"
                "   → 보정 기준: SDH Ø 2.0 mm (t/4, t/2, 3t/4)\n"
                "   → 기준 에코: 80% FSH (Full Screen Height)\n\n"
                "④ \"Acceptance criteria shall meet the requirements of Table 7.5.5-1\"\n"
                "   → 수용 기준은 오직 Table 7.5.5-1 기준만 적용\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ Table 7.5.5-1 수용 기준 (Acceptance Criteria)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "【평면형 결함 (Planar)】\n"
                "  - 균열 / 용합 불량 / 용입 불량: 크기 무관 전부 불합격\n\n"
                "【체적형 결함 (Volumetric) — 내부】\n"
                "  ┌────────────────┬───────────────┐\n"
                "  │ 결함 높이 (a)   │ 허용 길이 (ℓ) │\n"
                "  ├────────────────┼───────────────┤\n"
                "  │ a ≤ 3 mm       │ ℓ ≤ 6 mm      │\n"
                "  │ 3 mm < a ≤ 6 mm│ ℓ ≤ 2a        │\n"
                "  │ a > 6 mm       │ 불합격         │\n"
                "  └────────────────┴───────────────┘\n\n"
                "【표면 연결 결함 (Surface-Breaking)】\n"
                "  ┌────────────────┬───────────────┐\n"
                "  │ 결함 높이 (a)   │ 허용 길이 (ℓ) │\n"
                "  ├────────────────┼───────────────┤\n"
                "  │ a ≤ 1.5 mm     │ ℓ ≤ 6 mm      │\n"
                "  │ a > 1.5 mm     │ 불합격         │\n"
                "  └────────────────┴───────────────┘\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ Div.1 vs Div.2 수용 기준 비교\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "┌──────────────────┬──────────────────┬──────────────────┐\n"
                "│ 항목              │ Div.1 (UW-53)    │ Div.2 (7.5.5-1)  │\n"
                "├──────────────────┼──────────────────┼──────────────────┤\n"
                "│ 평가 기준         │ 결함 길이 위주    │ 높이(a)+길이(ℓ)  │\n"
                "│ 평면형 결함       │ 전부 불합격       │ 전부 불합격      │\n"
                "│ 체적형 최대 허용  │ 길이 ≤ t/3 등    │ a≤6mm, ℓ≤2a     │\n"
                "│ 표면 결함         │ 별도 언더컷 기준  │ a≤1.5mm 엄격     │\n"
                "│ 설계 허용 응력    │ 낮음 (보수적)     │ 높음 (엄격 검사) │\n"
                "└──────────────────┴──────────────────┴──────────────────┘",
            "ASME Sec. VIII Div.2 Para. 7.5.4 (Radiographic Examination)": 
                "ASME Section VIII Division 2, Paragraph 7.5.4: Radiographic Examination\n\nRadiographic examination shall be performed in accordance with the requirements of Article 7.5. The examination shall be conducted by qualified personnel using approved techniques. Acceptance criteria shall meet the requirements of Table 7.5.4-1.",
            "ASME B31.1 PAUT 관련 코드":
                "ASME B31.1 PAUT 관련 코드\n\n- 적용 기준: ASME B31.1 Table 136.4.1 수용 기준.\n- 검사 범위: B31.1 136.5 용접부 검사 요구 사항 준수.\n- 절차 기준: ASME V Article 4 및 Appendix III에 따른 서면 시험 절차.\n- 자격: SNT-TC-1A 또는 ASNT Level II / III 자격.\n- 장비 보정: IIW 블록 또는 ASME 보정 블록을 사용한 보정 및 감도 확인.\n- 스캔 범위: 용접부 및 열영향부 전체 체적 커버리지.\n- 기록: 결함 위치, 크기, 평가 결과, 테스트 조건을 포함한 문서화.\n- 보고서: 절차 식별, 장비, 인원, 검사 결과 및 판정 포함.",
            "ASME B31.1 PAUT (Phased Array Ultrasonic Testing)": 
                "ASME B31.1 Power Piping - PAUT Requirements (8대 필수 요건 상세)\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "1. Procedure (서면 절차서)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "Written procedure in accordance with ASME V Article 4 and Appendix III.\n"
                "필수 기재 항목:\n"
                "  - 검사 범위 및 적용 코드 (B31.1 Para. 136.4)\n"
                "  - 탐촉자 사양 (주파수, 소자 수, 피치, 굴절각 범위)\n"
                "  - Focal Law 설계 및 검증 방법\n"
                "  - S-scan / Linear scan 스캔 계획 (Scan Plan)\n"
                "  - 보정 블록 사양 및 보정 절차\n"
                "  - 감도 설정 방법 (DAC/TCG)\n"
                "  - 합부판정 기준 (Table 136.4.1 참조)\n"
                "  - 기록 및 보고 요건\n"
                "  → 절차서는 AI (Authorized Inspector) 또는 발주처 사전 승인 필요\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "2. Personnel (검사원 자격)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "Qualified Level II or III per SNT-TC-1A, ASNT, or equivalent.\n"
                "  - Level II: 검사 수행 + 합부판정 가능\n"
                "  - Level III: 절차 개발·승인 + Level II 감독\n"
                "  - 자격 유효기간: SNT-TC-1A 기준 5년 (사내 갱신 프로그램)\n"
                "  - PAUT 전용 실무 훈련 시간 별도 요구 (UT 경력 포함)\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "3. Equipment (장비)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "Calibrated PAUT system with appropriate probes and wedge angles.\n"
                "  - PA 인스트루먼트: 멀티채널 위상배열 장비 (최소 16~32 채널)\n"
                "  - 프로브: 검사 두께·재질에 맞는 주파수 (2~5 MHz)\n"
                "  - 웨지: 굴절각 45°~70° (검사 형상에 따라 선택)\n"
                "  - 스캐너: 인코더 부착 (위치 정확도 ±1 mm)\n"
                "  - 장비 교정 주기: 제조사 권고 주기 준수 (보통 1년)\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "4. Calibration (보정)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "Use IIW or ASME calibration blocks and perform sensitivity checks.\n"
                "  IIW Block 용도:\n"
                "    - 굴절각 확인 / 음속 측정 / 영점 설정\n"
                "  ASME Basic Block (Art.4 Table T-434.2.1) 용도:\n"
                "    - SDH 기반 DAC 곡선 작성 / 검사 감도 설정\n"
                "  보정 주기:\n"
                "    - 검사 시작 전 / 매 8시간마다 / 검사 종료 후\n"
                "    - 장비 이동 또는 온도 차 ±14°C 초과 시 재보정\n"
                "  유효성 기준:\n"
                "    ±2 dB 이내 → 유효 / ±4 dB 초과 → 전체 재검사\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "5. Scanning (스캔)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "Ensure full volumetric coverage of the weld and HAZ.\n"
                "  - 용접부 전체 단면 체적 (루트 / 충전부 / 캡) 100% 커버\n"
                "  - 열영향부 (HAZ) 양쪽 포함\n"
                "  - S-scan: 각도 범위 내 용접 단면 연속 스캔\n"
                "  - 최소 2방향 이상 스캔 (종방향·횡방향 결함 모두 탐지)\n"
                "  - 스캔 커버리지 도해 (Coverage Map) 문서화 필수\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "6. Evaluation (평가)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "Apply acceptance criteria in Table 136.4.1 or project-specific criteria.\n"
                "  DAC 진폭 기준:\n"
                "    - DAC 20% 미만     → 기록 불요\n"
                "    - DAC 20% ~ 100%  → 기록 + 크기 측정 후 Table 136.4.1 적용\n"
                "    - DAC 100% 초과    → 불합격 추정, 크기 측정 필수\n"
                "  결함 크기 측정법:\n"
                "    - −6 dB 강하법 (길이·높이)\n"
                "    - −20 dB 끝점법 (길이)\n"
                "    - TOFD 병행 가능 (높이 정밀 측정)\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "7. Records (기록)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "Maintain complete inspection records including defect sizing,\n"
                "location, probe data, and disposition.\n"
                "  필수 기록 항목:\n"
                "    - 절차서 번호 및 개정 이력\n"
                "    - 장비 시리얼 번호 및 교정 유효기간\n"
                "    - 검사원 이름 및 자격 번호\n"
                "    - 보정 블록 사양 및 보정 결과\n"
                "    - 결함 위치 (스캔 축, 인덱스 축, 깊이)\n"
                "    - 결함 크기 (길이 ℓ, 높이 a) 및 측정 방법\n"
                "    - 최대 에코 진폭 (%DAC)\n"
                "    - 합부 판정 결과 (합격/불합격) 및 근거 조항\n"
                "    - A-scan / S-scan 디지털 데이터 파일 보존\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "8. Reporting (보고서)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "Include procedure identification, equipment, personnel,\n"
                "and results in the test report.\n"
                "  보고서 필수 포함 항목:\n"
                "    - 검사 절차서 번호 및 개정 번호\n"
                "    - 적용 코드 및 수용 기준 조항\n"
                "    - 장비 명칭, 모델, 시리얼, 교정 유효기간\n"
                "    - 검사원 성명, 자격 등급, 자격 번호\n"
                "    - 검사 일시 및 장소\n"
                "    - 검사 대상 (용접부 번호, 이음 유형, 두께)\n"
                "    - 검사 결과 요약 (지시 목록 및 합부 판정)\n"
                "    - 검사원 서명 및 Level III 확인 서명",
            "ASME B31.1 PAUT 필수 항목":
                "ASME B31.1 PAUT 필수 항목\n\n- 검사 절차: ASME V Article 4, Appendix III에 따른 서면 절차.\n- 자격: SNT-TC-1A 또는 ASNT 기준의 Level II/III 검사자.\n- 장비: 적절한 위상배열 초음파 시스템 및 프로브.\n- 보정: IIW 블록 또는 ASME 보정 블록을 이용한 보정 및 감도 확인.\n- 스캔: 용접부와 열영향부의 전체 체적 커버리지.\n- 평가: B31.1 Table 136.4.1 기준 또는 지정된 수용 기준.\n- 기록: 결함 위치, 크기, 평가 및 처분을 포함한 완전한 기록.\n- 보고: 절차, 장비, 인원, 결과를 포함한 보고서 작성.",
            "ASME B31.1 PAUT 합부판정 기준 (Table 136.4.1)":
                "ASME B31.1 Power Piping - Para. 136.4 및 Table 136.4.1 상세\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ Para. 136.4 구조 (검사 요건 계층)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "Para. 136.4.1  검사 방법 선택 기준\n"
                "  → RT, UT(PAUT 포함), 또는 대체 NDE 선택 근거 규정\n"
                "  → RT 대체 시 동등 검사 능력 입증 필수\n\n"
                "Para. 136.4.2  검사 범위 (Required Examination)\n"
                "  → 검사 대상 용접부 유형 및 비율 규정\n"
                "  → 진보적 검사 (Progressive Examination) 조항:\n"
                "     최초 용접부 검사 불합격 시 추가 검사 범위 확대\n\n"
                "Para. 136.4.3  수용 기준 (Acceptance Criteria)\n"
                "  → \"Acceptance criteria shall be in accordance with\n"
                "      Table 136.4.1\"\n"
                "  → Table 136.4.1이 B31.1의 유일한 공식 합부판정 기준\n\n"
                "Para. 136.4.4  재검사 (Re-examination)\n"
                "  → 불합격 지시 수정 후 재검사 절차\n"
                "  → 동일 검사 방법·절차로 재실시\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ Table 136.4.1 — 불합격 지시 기준\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "【1】 균열 (Cracks)\n"
                "  → 크기·길이·위치 무관 전부 불합격\n\n"
                "【2】 완전 용입 불량 (Incomplete Penetration) / 용합 불량 (Incomplete Fusion)\n"
                "  두께(t)별 임의 300 mm 구간 내 누적 길이 기준:\n"
                "  ┌──────────────────────┬─────────────────────────────┐\n"
                "  │ 검사 두께 (t)         │ 최대 허용 누적 길이          │\n"
                "  ├──────────────────────┼─────────────────────────────┤\n"
                "  │ t < 19 mm (3/4\")     │ 25 mm (1\") 이하             │\n"
                "  │ 19 mm ≤ t < 57 mm    │ t/3 이하                    │\n"
                "  │ t ≥ 57 mm (2-1/4\")   │ 19 mm (3/4\") 이하           │\n"
                "  └──────────────────────┴─────────────────────────────┘\n"
                "  → IP / IF 동일 기준 적용\n\n"
                "【3】 내부 결함 (Porosity / Slag Inclusion 등)\n"
                "  - 개별 지시 길이 > 6 mm (1/4\") → 불합격\n"
                "  - 임의 150 mm (6\") 구간 내 지시 길이 합계 > 12 mm → 불합격\n"
                "  예외 허용 (모두 충족 시):\n"
                "    · 최대 개별 지시 < 3 mm (1/8\")\n"
                "    · 군집 면적 < 6 cm² (1 in²) 이하\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ PAUT 진폭 기반 평가 (DAC 기준)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "┌──────────────────┬────────────────────────────────────┐\n"
                "│ 진폭 수준         │ 조치                                │\n"
                "├──────────────────┼────────────────────────────────────┤\n"
                "│ DAC 20% 미만      │ 기록 불요                           │\n"
                "│ DAC 20% ~ 100%   │ 기록 필수 + 크기 측정 후 Table 적용 │\n"
                "│ DAC 100% 초과     │ 불합격 추정 → 크기 측정·평가 필수  │\n"
                "└──────────────────┴────────────────────────────────────┘\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ B31.1 vs B31.3 수용 기준 비교\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "┌──────────────────┬──────────────────┬──────────────────────┐\n"
                "│ 결함 유형         │ B31.1 T.136.4.1  │ B31.3 T.341.3.2      │\n"
                "├──────────────────┼──────────────────┼──────────────────────┤\n"
                "│ 균열              │ 전부 불합격       │ 전부 불합격          │\n"
                "│ IF/IP 구간       │ 300 mm 기준       │ 100 mm 기준 (엄격)   │\n"
                "│ IF/IP t<19mm     │ 누적 ≤ 25 mm      │ 누적 ≤ t/3 ~ 6 mm   │\n"
                "│ 기공 개별         │ ≤ 6 mm           │ ≤ 3 mm (엄격)        │\n"
                "│ 기공 구간 합계    │ 150 mm 내 12 mm  │ 100 mm 내 6 mm (엄격)│\n"
                "└──────────────────┴──────────────────┴──────────────────────┘\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 진보적 검사 (Progressive Examination, Para. 136.4.2)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "초기 샘플 검사 결과에 따른 추가 검사 범위:\n"
                "  1차 불합격 → 동일 용접사·절차의 용접부 2개 추가 검사\n"
                "  2차 불합격 → 해당 용접사·절차의 나머지 전체 용접부 검사\n"
                "  ※ 완전 합격 시 이후 샘플링 비율 유지",
            "ASME B31.3 PAUT 합부판정 기준 (Table 341.3.2)":
                "ASME B31.3 Process Piping - 유체 서비스 분류 및 Table 341.3.2 수용 기준\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ B31.3 유체 서비스 분류 (Fluid Service Categories)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "B31.3는 유체의 위험도에 따라 배관을 5가지 서비스로 분류하고\n"
                "각각 다른 검사 범위·수용 기준을 적용함.\n\n"
                "【1】 Normal Fluid Service (일반 유체 서비스) — 기본값\n"
                "  - 정의: 위험 유체 / 고압 / 극저온이 아닌 일반적인 공정 배관\n"
                "  - 해당 유체 예시: 냉각수, 저압 스팀, 공기, 일반 화학약품\n"
                "  - 검사 요건: 용접부의 5% 이상 (Random Examination)\n"
                "  - 수용 기준: Table 341.3.2 Normal 기준 적용\n"
                "  - ⚠ B31.3의 가장 기본 등급 — 별도 분류 없으면 여기 해당\n\n"
                "【2】 Category D Fluid Service (위험도 낮은 유체)\n"
                "  - 정의: 비가연성·무독성, 설계 압력 ≤ 1035 kPa (150 psi),\n"
                "    설계 온도 −29°C ~ +186°C 범위\n"
                "  - 해당 유체 예시: 물, 저압 스팀, 압축공기, 냉매\n"
                "  - 검사 요건: 육안 검사(VT)만으로 RT/UT 생략 가능\n"
                "  - 수용 기준: Para. 341.4 (완화 기준)\n\n"
                "【3】 Category M Fluid Service (극히 위험한 유체)\n"
                "  - 정의: 미량 누출도 인체에 치명적인 독성 유체\n"
                "    (TLV ≤ 1 ppm, 또는 이에 준하는 독성)\n"
                "  - 해당 유체 예시: 염소(Cl₂), 포스겐(COCl₂), HF, 일산화탄소(CO)\n"
                "  - 검사 요건: 용접부 100% RT 또는 UT\n"
                "  - 수용 기준: Table 341.3.2 Category M 기준 (Normal보다 엄격)\n\n"
                "【4】 High Pressure Fluid Service (고압 유체)\n"
                "  - 정의: 설계 압력이 ASME B16.5 Class 2500 플랜지 허용 압력 초과\n"
                "  - 적용: B31.3 Chapter IX 별도 적용\n"
                "  - 검사 요건: 용접부 100% RT 또는 UT (Code Case 2235 권장)\n"
                "  - 수용 기준: Chapter IX 별도 기준\n\n"
                "【5】 Severe Cyclic Conditions (심한 사이클 하중)\n"
                "  - 정의: 응력 범위가 허용 응력의 80% 초과 또는 사이클 횟수 과다\n"
                "  - 검사 요건: 용접부 100% RT 또는 UT\n"
                "  - 수용 기준: Table 341.3.2 Severe Cyclic 기준 (가장 엄격)\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ Table 341.3.2 — 수용 기준 비교\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "【Normal Fluid Service 기준】\n"
                "  [균열] → 전부 불합격\n"
                "  [IF / IP] 임의 100 mm 구간 내 누적 길이 기준:\n"
                "    ┌──────────────┬─────────────────┐\n"
                "    │ 두께 (t)      │ 최대 허용 누적 ℓ │\n"
                "    ├──────────────┼─────────────────┤\n"
                "    │ t ≤ 6 mm     │ 2 mm            │\n"
                "    │ 6 < t ≤ 19 mm│ t/3             │\n"
                "    │ t > 19 mm    │ 6 mm            │\n"
                "    └──────────────┴─────────────────┘\n"
                "  [기공 / 슬래그]\n"
                "    - 개별 지시 > 3 mm → 불합격\n"
                "    - 임의 100 mm 내 합계 > 6 mm → 불합격\n"
                "    - 군집 기공: 25 cm² 내 > 1 cm² → 불합격\n\n"
                "【Category M / Severe Cyclic 기준 (엄격)】\n"
                "  [IF / IP] → 길이 > 0 (어떠한 지시도 불합격)\n"
                "  [기공]    → 개별 > 1.5 mm → 불합격\n"
                "  [균열]    → 전부 불합격\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 검사 범위 요약\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "┌────────────────────┬───────────────────────┐\n"
                "│ 서비스 분류         │ 최소 검사 비율         │\n"
                "├────────────────────┼───────────────────────┤\n"
                "│ Category D          │ VT만 (RT/UT 불필요)   │\n"
                "│ Normal              │ 5% Random             │\n"
                "│ Severe Cyclic       │ 100%                  │\n"
                "│ Category M          │ 100%                  │\n"
                "│ High Pressure       │ 100% (Chapter IX)     │\n"
                "└────────────────────┴───────────────────────┘\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ PAUT 진폭 평가 기준 (DAC)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- 기록 수준: DAC 20% (−14 dB)\n"
                "- 평가 수준: DAC 50% (−6 dB)\n"
                "- DAC 100% 초과: 불합격 추정, 크기 측정·평가 필수",
            "ASME Sec. VIII Div.1 PAUT 합부판정 기준 (UW-51/App.12)":
                "ASME Section VIII Div.1 PAUT 합부판정 기준\n\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ UW-51 (전수 RT 대체 PAUT) 수용 기준\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n[1] 무조건 불합격 결함\n- 균열 (Cracks): 크기·위치 무관 전부 불합격\n- 용합 불량 (Incomplete Fusion)\n- 용입 불량 (Incomplete Penetration)\n\n[2] 내부 결함 (기공, 슬래그)\n두께(t)별 최대 허용 개별 지시 길이:\n  · t ≤ 19mm  → 최대 6mm\n  · 19mm < t ≤ 57mm → 최대 t/3\n  · t > 57mm  → 최대 19mm\n\n임의 12t 구간(단, 최대 152mm) 내 지시 길이 합계:\n  → 위 개별 기준치 이내\n\n[3] 언더컷 (Undercut)\n- 표면 언더컷 깊이 > 1mm: 불합격\n- 0.4mm 미만: 허용\n\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ Appendix 12 (UT 대체) 수용 기준\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n- 평면형 결함 (균열, IF, IP): 전부 불합격\n- 체적형 결함 (기공, 슬래그):\n  · 개별 지시 높이 > 25% t 또는 6mm 중 작은 값: 불합격\n  · 지시 길이: UW-51 기준과 동일 적용\n- DAC 기준:\n  · 평가 수준: DAC 50%\n  · 기록 수준: DAC 20%",
            "ASME PAUT 파괴역학 합부판정 (Code Case 2235 / ECA)":
                "ASME PAUT 파괴역학 합부판정 기준\n"
                "Code Case 2235 / Engineering Critical Assessment (ECA)\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 개요 (Overview)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "파괴역학(Fracture Mechanics) 기반 합부판정은 결함의 실제\n"
                "구조적 영향을 평가하여 전통적 RT/UT 기준보다 완화된 허용\n"
                "기준을 적용하는 방법. ASME Code Case 2235가 대표적.\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ ASME Code Case 2235 (UT/PAUT로 RT 대체)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "적용 조건:\n"
                "- 적용 코드: ASME Section VIII Div.1, Div.2 및 Section I\n"
                "- 두께 범위: 20mm ≤ t ≤ 250mm\n"
                "- 대상 용접부: 완전 용입 맞대기 용접 (Full Penetration Butt Weld)\n"
                "- 적용 재질 (P-Number):\n"
                "  P-No.1(탄소강), 3(저합금), 4(Cr-Mo), 5A/5B(고Cr-Mo),\n"
                "  6/7(마르텐사이트/페라이트 스테인리스), 8(오스테나이트 STS),\n"
                "  9A/9B(Ni합금강), 10A/10F(고강도 저합금강)\n"
                "- 용접 조인트 범주: Category A (종방향·환형 맞대기) 우선 적용\n"
                "- 표면 상태: 검사 전 기계 가공 또는 연마 (Ra ≤ 6.3 μm 권장)\n"
                "- 검사 온도: 15°C ~ 50°C (재질·절차에 따라 조정)\n\n"
                "절차 자격 인정 (Procedure Qualification):\n"
                "- ASME Section V, Article 4 Appendix III 준수 서면 절차 필수\n"
                "- 실제 검사 두께와 동일한 데모 블록(Demonstration Block)으로\n"
                "  절차 유효성 검증 (Blind Test / Mock-Up)\n"
                "- POD (Probability of Detection): 목표 결함 크기에서 ≥ 90%\n"
                "- 절차 변경(장비·프로브·각도·설정) 시 재자격 인정 필요\n\n"
                "보정 블록 기준 (Calibration Block):\n"
                "- 재질: 검사 대상과 동종 또는 음향 임피던스 동등 재질\n"
                "- 기준 반사체 (SDH, Side Drilled Hole):\n"
                "  · t ≤ 50mm  → Ø 1.5mm SDH\n"
                "  · t > 50mm  → Ø 2.0mm SDH\n"
                "- SDH 배치 깊이: t/4, t/2, 3t/4 (최소 3개 지점)\n"
                "- 보정 주기: 검사 시작 전, 종료 후, 매 4시간마다, 장비 이동 시\n"
                "- 온도 보정: 검사체와 보정 블록 온도 차 ±14°C 초과 시 재보정\n\n"
                "스캔 커버리지 (Scan Coverage):\n"
                "- 전체 용접 체적의 100% 커버리지 확보 필수\n"
                "- HAZ(열영향부) 포함 모재 측 최소 t/4 이상 구간 검사\n"
                "- 인덱스 방향 스캔 간격(Scan Increment): ≤ 1.0mm\n"
                "- S-scan 빔 각도 범위: 40°~70° (스텝 ≤ 2° 권장)\n"
                "- 모든 각도에서 -6dB 중첩(Beam Overlap) 유지\n\n"
                "신호 수준 기준 (Amplitude Criteria):\n"
                "- 기록 수준 (Recording Level)  : DAC 20% (−14 dB) 이상\n"
                "- 평가 수준 (Evaluation Level) : DAC 50% (−6 dB) 이상 → 크기 측정\n"
                "- 거부 수준 (Rejection Level)  : DAC 100% 초과 → 즉시 크기 평가\n"
                "- S/N 비: ≥ 3:1 (9.5 dB) 유지 필수\n\n"
                "결함 허용 기준 (Allowable Flaw Size):\n"
                "아래 조건을 모두 만족 시 합격\n\n"
                "  (1) 결함 높이 (a, Through-Thickness):\n"
                "      a ≤ 0.1t  (단, a ≤ 6mm)\n\n"
                "  (2) 결함 길이 (ℓ, Along Weld):\n"
                "      ℓ ≤ 6a  (최대 50mm)\n\n"
                "  (3) 표면 연결 결함 (Surface-Breaking):\n"
                "      허용 높이 기준 50% 감소 → a_allow × 0.5 적용\n\n"
                "  (4) 결함 간격 규칙 (Flaw Spacing Rule):\n"
                "      인접 결함 간격 S < max(a₁, a₂) 이면\n"
                "      두 결함을 단일 결함으로 합산하여 평가\n\n"
                "평면형 결함 (Planar Flaws - 균열, 용합불량, 용입불량):\n"
                "  위 (1)~(4) 기준 동시 적용\n"
                "  균열성 결함: 보수적 평가 필수, 재질의 K_IC, ΔK_th 확인\n\n"
                "검사원 자격 (Personnel Qualification):\n"
                "- ASNT SNT-TC-1A 또는 CP-189 기준 PAUT Level II 이상\n"
                "- 해당 장비·소프트웨어 교육 이수 및 실기 평가 기록 보유\n"
                "- 합부판정은 Level II 이상만 수행 가능\n\n"
                "보고서 요건 (Reporting Requirements):\n"
                "- 장비 식별(SN), 프로브 사양, 보정 데이터 첨부\n"
                "- A-scan, S-scan 원시 데이터(Raw Data) 파일 보존\n"
                "- 검사된 용접부 번호, 길이, 커버리지 맵 포함\n"
                "- 지시 목록: 위치(X, Y, 깊이), 높이, 길이, 판정 결과\n"
                "- 데이터 저장 형식: DICONDE 또는 제조사 전용 포맷\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 파괴역학 ECA (Engineering Critical Assessment)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "ECA 적용 기준 (BS 7910 / API 579 방법론):\n\n"
                "[1] 임계 결함 크기 계산 (Critical Flaw Size)\n"
                "    K_I = Y × σ × √(π·a)  ≤  K_IC / SF\n"
                "    - K_I : 응력 확대 계수\n"
                "    - Y   : 형상 계수 (결함 형상·위치 의존)\n"
                "    - σ   : 작용 응력 (막응력 + 굽힘응력 + 잔류응력)\n"
                "    - K_IC: 파괴인성 (재질별 시험값 또는 Charpy 변환식)\n"
                "    - SF  : 안전계수 (일반적으로 2.0~2.5)\n\n"
                "[2] 피로 균열 성장 평가\n"
                "    da/dN = C × (ΔK)^m  (Paris Law)\n"
                "    - 설계 수명 내 결함 성장 후 최종 크기 ≤ a_critical\n"
                "    - 검사 주기 결정에도 활용\n\n"
                "[3] 파손 평가 다이어그램 (FAD - Failure Assessment Diagram)\n"
                "    Kr = K_I / K_IC  (파괴비)\n"
                "    Lr = P / P_L     (소성 붕괴비)\n"
                "    → FAD 곡선 내부에 위치 시 합격\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ PAUT 결함 측정 요건 (ECA 입력 데이터)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- 결함 높이(a): TOFD 또는 -6dB drop 법, 정확도 ±1mm 이내\n"
                "- 결함 길이(ℓ): -6dB 또는 -20dB 끝점법\n"
                "- 결함 위치: 표면 연결 여부, 깊이(d), 두께 위치 확인\n"
                "- 결함 형상: 타원형(Elliptical) 가정 → a/c 비율 결정\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 적용 참고 코드\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- ASME Code Case 2235 (최신 개정판 확인 필수)\n"
                "- API 579-1/ASME FFS-1 : Part 9 (균열형 결함 FFS)\n"
                "- BS 7910 : Guide to methods for assessing the acceptability\n"
                "  of flaws in metallic structures\n"
                "- ASME Section XI : IWB-3600 (원자력 배관 ECA)\n"
                "- 적용 시 재질의 파괴인성(K_IC) 시험값 또는\n"
                "  Charpy → K_IC 변환식 사용 근거 문서화 필수",

            "Appendix 12 vs Code Case 2235 - 구조·차이·적용 흐름 상세":
                "ASME Sec. VIII Div.1 — Appendix 12 vs Code Case 2235\n"
                "RT 대체 UT/PAUT 적용 체계 완전 해설\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 1. 두 기준의 위치와 역할\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "┌─────────────────┬──────────────────────┬──────────────────────┐\n"
                "│ 구분            │ Appendix 12          │ Code Case 2235       │\n"
                "├─────────────────┼──────────────────────┼──────────────────────┤\n"
                "│ 문서 유형       │ Mandatory Appendix   │ Code Case (임시 허용)│\n"
                "│                 │ (코드 본문 일부)     │ (별도 승인 문서)     │\n"
                "│ 발행 기관       │ ASME Sec. VIII 위원회│ ASME Standards Comm. │\n"
                "│ 적용 범위       │ Div.1 전용           │ Div.1 + Div.2 +      │\n"
                "│                 │                      │ Section I            │\n"
                "│ 두께 하한       │ t ≥ 1/2\" (13mm)     │ t ≥ 20mm (25mm 실무) │\n"
                "│ SDH 직경 기준   │ Ø 3/64\" (≈1.2mm)    │ t≤50mm: Ø1.5mm       │\n"
                "│                 │                      │ t>50mm : Ø2.0mm      │\n"
                "│ 절차 검증       │ 규정 없음            │ 데모 블록 필수       │\n"
                "│ 결함 허용 기준  │ 길이 기반 (UW-53)    │ 높이+길이 복합 기준  │\n"
                "│ 표면 결함 기준  │ 언더컷 기준 준용     │ 50% 감소 적용        │\n"
                "│ 발주처 승인     │ 필요                 │ 필요 + AI 별도 확인  │\n"
                "│ 데이터 보존     │ 일반 기록            │ Raw Data 전체 보존   │\n"
                "└─────────────────┴──────────────────────┴──────────────────────┘\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 2. Appendix 12 상세 구조\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "Appendix 12는 Sec. VIII Div.1 코드 본문에 포함된\n"
                "\"Mandatory\"(강제) 부록으로, UW-53과 연계하여\n"
                "RT 대체 UT 시 최소 요건을 규정한다.\n\n"
                "[12-1] Scope (적용 범위)\n"
                "  - 맞대기 용접부 완전 체적 검사 목적\n"
                "  - UW-11(a)(4): RT 불가 시 UT 대체 허용 근거 제공\n"
                "  - UW-53: UT 수용 기준 조항 (Appendix 12와 연동)\n\n"
                "[12-2] Procedure Requirements (절차서 요건)\n"
                "  - ASME Section V Article 4 기준 서면 절차서 필수\n"
                "  - 절차서에 포함해야 할 최소 항목:\n"
                "    · 장비 유형, 주파수, 프로브 각도\n"
                "    · 보정 블록 사양 (재질, SDH 직경·깊이)\n"
                "    · 스캔 방향·범위·인덱스 간격\n"
                "    · 감도 설정 방법 (DAC 곡선 작성)\n"
                "    · 기록·평가·거부 수준 dB 값\n"
                "    · 결함 크기 측정 방법\n"
                "  - 절차 변경 시 재승인 필요\n\n"
                "[12-3] Equipment (장비)\n"
                "  - A-scan 표시 장치 필수 (실시간 파형 확인)\n"
                "  - PAUT 사용 시: S-scan 표시 및 Focal Law 설계 요건 추가\n"
                "  - 장비 교정 유효성: 제조사 권고 주기 (연 1회 이상)\n\n"
                "[12-4] Calibration (보정)\n"
                "  보정 블록 (Basic Calibration Block):\n"
                "    - 재질: 검사 대상과 동일 P-Number\n"
                "    - 기준 반사체: SDH Ø 3/64\" (≈1.19mm)\n"
                "    - SDH 배치: t/4, t/2, 3t/4 (최소 3점)\n"
                "  DAC 곡선 작성:\n"
                "    - 각 깊이 SDH 최대 에코를 연결한 거리-진폭 곡선\n"
                "    - Search Level = DAC + 6 dB\n"
                "    - 기록 수준: DAC 20% (−14 dB)\n"
                "    - 평가 수준: DAC 50% (−6 dB)\n"
                "    - 거부 수준: DAC 100%\n"
                "  보정 주기:\n"
                "    - 검사 시작 전·종료 후 필수\n"
                "    - 매 2시간마다 (또는 절차서 명시 주기)\n"
                "    - 장비 이동, 충격, 온도 변화 ±14°C 초과 시\n"
                "  보정 유효성:\n"
                "    - ±2 dB 이내 → 계속 유효\n"
                "    - 2~4 dB 이탈 → 직전 보정 이후 지시 재평가\n"
                "    - 4 dB 초과 이탈 → 직전 보정 이후 전 용접부 재검사\n\n"
                "[12-5] Coverage (검사 커버리지)\n"
                "  - 용접부 전체 단면 체적 100% 필수\n"
                "  - HAZ 포함 (최소 모재 측 13mm)\n"
                "  - 최소 2방향 각도 스캔:\n"
                "    · 루트부 결함 탐지: 45° 또는 60° 빔\n"
                "    · 측벽 용합 불량 탐지: 45°~60°\n"
                "    · 종방향 결함 (횡단 균열): 추가 스캔 필요\n"
                "  - 스캔 커버리지 도해 (Coverage Plan) 문서화\n\n"
                "[12-6] Acceptance Criteria (수용 기준 - Table UW-53)\n"
                "  평면형 결함:\n"
                "    - 균열: 전부 불합격\n"
                "    - 용합불량 (IF), 용입불량 (IP):\n"
                "      t < 19mm  → 누적 길이 > 25mm: 불합격\n"
                "      19≤t<57mm → 누적 길이 > t/3 : 불합격\n"
                "      t ≥ 57mm  → 누적 길이 > 19mm: 불합격\n"
                "      (임의 300mm 구간 내 기준)\n"
                "  체적형 결함:\n"
                "    - 개별 지시 길이 > 6mm: 불합격\n"
                "    - 임의 150mm 내 합계 > 12mm: 불합격\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 3. Code Case 2235 상세 구조\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "Code Case 2235는 Appendix 12를 기반으로 하되,\n"
                "결함 허용 기준을 파괴역학(ECA) 방법론으로 강화하고\n"
                "절차 자격 인정(POD 시험)을 추가한 상위 기준이다.\n\n"
                "[CC 2235] 적용 조건\n"
                "  - 발행 코드: Section VIII Div.1, Div.2, Section I\n"
                "  - 두께 범위: 20mm ≤ t ≤ 250mm\n"
                "  - 대상 용접부: Full Penetration Butt Weld\n"
                "  - 재질 P-Number: P1/3/4/5A/5B/6/7/8/9A/9B/10A/10F\n"
                "  - 표면 조건: 검사 전 연마 Ra ≤ 6.3μm\n"
                "  - 온도 조건: 15°C ~ 50°C\n\n"
                "[CC 2235] Procedure Qualification (절차 자격 인정)\n"
                "  Appendix 12 대비 추가 요건:\n"
                "  ① 데모 블록 (Demonstration Block) 필수:\n"
                "     - 실제 검사 두께와 동일 재질·두께로 제작\n"
                "     - 인공 결함 매입 (SDH, FBH, 노치 등)\n"
                "     - 블라인드 테스트(Blind Test) 수행\n"
                "       → 목표 결함 크기 탐지율 POD ≥ 90% 달성 필수\n"
                "  ② 절차 변경 재자격:\n"
                "     - 장비 교체, 프로브 변경, 각도 변경, 설정 변경 시\n"
                "       데모 블록 재시험 필요\n\n"
                "[CC 2235] SDH 보정 블록 강화 기준\n"
                "  - Appendix 12의 Ø1.2mm보다 큰 SDH 적용:\n"
                "    · t ≤ 50mm : SDH Ø 1.5mm (약 25% 증가)\n"
                "    · t > 50mm : SDH Ø 2.0mm (약 67% 증가)\n"
                "  → 더 작은 결함까지 탐지·기록 의무화\n\n"
                "[CC 2235] 보정 주기 강화\n"
                "  - Appendix 12(2시간)보다 강화:\n"
                "    · 검사 전·후 필수\n"
                "    · 매 4시간마다 (일부 절차서: 2시간)\n"
                "    · 장비 이동·충격 후\n"
                "    · 온도 차 ±14°C 초과 시\n\n"
                "[CC 2235] 결함 허용 기준 (Allowable Flaw Size)\n"
                "  Appendix 12(길이 기반)에서 높이+길이 복합 기준으로 전환:\n\n"
                "  평면형 결함:\n"
                "    (1) 결함 높이: a ≤ 0.1t  (단, a ≤ 6mm)\n"
                "    (2) 결함 길이: ℓ ≤ 6a  (단, ℓ ≤ 50mm)\n"
                "    (3) 표면 연결 결함: 허용치 × 0.5 (50% 감소)\n"
                "    (4) 결함 간격 규칙:\n"
                "        인접 결함 S < max(a₁, a₂) → 단일 결함으로 합산\n\n"
                "  체적형 결함:\n"
                "    - Appendix 12 기준 동일 적용\n"
                "    - 단, 표면 연결 체적 결함: 허용 높이 30% 감소\n\n"
                "[CC 2235] 스캔 파라미터 강화\n"
                "  - 인덱스 스캔 간격: ≤ 1.0mm (Appendix 12는 명시 없음)\n"
                "  - S-scan 각도 스텝: ≤ 2°\n"
                "  - 빔 중첩(-6dB): 모든 각도에서 확인\n"
                "  - S/N 비: ≥ 3:1 (9.5dB) 유지\n\n"
                "[CC 2235] 기록·보고 강화\n"
                "  - A-scan + S-scan 원시 데이터(Raw Data) 전량 보존\n"
                "  - 저장 형식: DICONDE 또는 제조사 전용 포맷\n"
                "  - 결함 목록: 위치(X, Y, 깊이), 높이, 길이, 판정\n"
                "  - 커버리지 맵(Coverage Map) 첨부\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 4. 적용 흐름도 (Decision Flow)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "RT 적용 가능? ─Yes─→ RT 실시 (UW-51/52)\n"
                "     │\n"
                "    No\n"
                "     ↓\n"
                "두께 ≥ 13mm? ─No─→ UT 일반 절차 (Art.4)\n"
                "     │\n"
                "    Yes\n"
                "     ↓\n"
                "Appendix 12 UT/PAUT 절차 작성 + 발주처 승인\n"
                "     │\n"
                "     ├─ [기본 적용] → Appendix 12 기준\n"
                "     │                SDH Ø1.2mm / UW-53 길이 기준\n"
                "     │\n"
                "     └─ [강화 적용] → Code Case 2235 추가 채택\n"
                "                      SDH Ø1.5~2.0mm / 높이+길이 기준\n"
                "                      데모 블록 POD ≥90% 필수\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 5. 실무 선택 기준\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "Appendix 12만 적용 시:\n"
                "  - 두께 13~20mm 비교적 얇은 용기\n"
                "  - 결함 허용 기준을 RT 수준(길이 기반)으로 유지\n"
                "  - 데모 블록 준비 시간·비용 절약 필요 시\n\n"
                "Code Case 2235 추가 적용 시:\n"
                "  - 두께 20mm 이상 고두께·고압 용기\n"
                "  - 발주처/고객이 CC 2235 명시 요구 시\n"
                "  - RT 완전 대체를 위한 높은 신뢰성 입증 필요 시\n"
                "  - TOFD 또는 고해상도 PAUT 장비 보유 시\n\n"
                "⚠ 주의: Code Case는 ASME 위원회 갱신 여부 확인 필수\n"
                "        (최신 Annex 또는 개정 CC 2235 Rev. 번호 확인)",

            "ASME Sec. VIII Div.2 PAUT 합부판정 기준 (Para. 7.5.5)":
                "ASME Section VIII Div.2 PAUT 합부판정 기준 - Paragraph 7.5.5\n\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ Table 7.5.5-1 수용 기준\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n[1] 평면형 결함 (Planar Defects)\n- 균열, 용합 불량, 용입 불량: 크기 무관 전부 불합격\n\n[2] 체적형 결함 (Volumetric Defects)\n결함 높이(a) 및 길이(ℓ) 기준:\n  · a ≤ 3mm : ℓ ≤ 6mm 허용\n  · 3mm < a ≤ 6mm : ℓ ≤ 2a 허용\n  · a > 6mm : 불합격\n\n표면 결함 (Surface-Breaking):\n  · a ≤ 1.5mm : ℓ ≤ 6mm 허용\n  · a > 1.5mm : 불합격\n\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ PAUT 결함 크기 측정 요건\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n- 결함 높이(a): -6dB drop법 또는 TOFD(Tip Diffraction) 병행\n- 결함 길이(ℓ): -6dB drop법 또는 -20dB 끝점법\n- 위치 정확도: ±1mm 이내\n\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ 보정 블록 기준 (Reference Sensitivity)\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n- 2mm SDH(Side Drilled Hole)을 기준 반사체로 DAC 설정\n- 평가 수준: DAC 50% (-6dB)\n- 기록 수준: DAC 20% (-14dB)\n- 검사 감도: 평가 수준보다 +6dB 추가 증폭하여 검사",
            "RT 표준 항목":
                "RT (Radiographic Testing) 표준 항목\n\n- 적용 코드: ASME Section V Article 2 (또는 프로젝트 지정 기준).\n- 필름/디지털 기법: 노출 조건, IQI, 감도 및 판독 조건 준수.\n- 평가 기준: 해당 제작 코드(예: ASME B31.1, ASME VIII)의 수용 기준 적용.\n- 기록 및 보고: 노출 조건, 판독 결과, 결함 위치/길이 및 판정 포함.",
            "ASME Section V, Article 2 (RT 기본 절차)":
                "ASME Section V, Article 2 - Radiographic Examination (RT 기본 절차)\n\n적용 범위:\n방사선투과검사(RT)의 핵심 절차 요구사항을 규정하는 기본 코드.\n\n주요 요구사항:\n- 선원(Source): X선관 또는 감마선원(Ir-192, Co-60, Se-75) 선택 기준\n- 기하학적 조건: 선원-피사체 거리(SFD), 필름-피사체 거리 설정\n- IQI (Image Quality Indicator): 선형(Wire) 또는 구멍형(Hole) IQI 선택 및 배치\n- 최소 감도: 2-2T (Wire IQI) 또는 2% 구멍 IQI 기준 충족\n- 필름: ASTM E1815 Class 1 또는 2 이상\n- 현상 처리: 온도·시간 기준 준수 또는 디지털 이미징 시스템 적용\n- 판독 조건: 밝기 조도, 차폐 조건 규정\n\n검사원 자격: ASNT SNT-TC-1A Level II 이상",
            "ASME Section V, Article 2, Appendix D (디지털 RT)":
                "ASME Section V, Article 2, Appendix D - Digital Radiography (DR)\n\n적용 범위:\n디지털 방사선 시스템(DR/CR)을 이용한 RT 절차.\n\n주요 요구사항:\n- 디지털 검출기: DR(평판), CR(이미징 플레이트) 사용 가능\n- IQI: 전통 RT와 동일 요건, 디지털 콘트라스트 감도 별도 검증\n- 공간 해상도: Basic Spatial Resolution(BSR) 측정 및 기준 충족\n- 시스템 검증: 동일 두께 필름 RT와 동등성 입증\n- 데이터 저장: 무결성 보장 포맷(DICONDE 권장)으로 보존\n- 판독: 모니터 최소 해상도 및 밝기 기준 준수",
            "ASME B31.1, Para. 136.4 (RT - 발전 배관)":
                "ASME B31.1, Paragraph 136.4 - Radiographic Examination (발전 배관 RT)\n\n적용 범위:\n발전소 배관 용접부에 대한 RT 요구사항.\n\n주요 요구사항:\n- 검사 범위: Table 136.4 기준 (범주별 RT 비율 규정)\n  · P1 ~ P15 재질 및 두께별 의무 검사 비율\n- 절차: ASME Section V Article 2 준수\n- 수용 기준: Table 136.4.1 적용\n  · 균열, 융합불량, 미용융, 기공 등 유형별 기준\n- 100% RT 대상: 고온·고압 배관, 카테고리 D 이상\n- 보고: 노출 조건, IQI 확인, 판독 결과 기록\n\n검사원 자격: ASNT Level II 이상",
            "ASME B31.3, Para. 344.5 (RT - 공정 배관)":
                "ASME B31.3, Paragraph 344.5 - Radiographic Examination (공정 배관 RT)\n\n적용 범위:\n석유화학·정유·가스 플랜트 공정 배관 RT.\n\n주요 요구사항:\n- 검사 비율 (Table 341.3.2):\n  · Normal: 5% 이상 무작위 검사\n  · Severe Cyclic: 100%\n  · Category M: 100%\n  · High Pressure: 100%\n- 절차: ASME Section V Article 2 준수\n- 수용 기준: Table 341.3.2 및 Appendix A 규정\n- 선원: X-Ray 또는 Gamma-Ray 프로젝트 승인 조건 사용\n- 추가 검사: 거부 지시 발견 시 동일 용접사 전수 검사",
            "ASME Sec. VIII Div.1, UW-51/52 (RT - 압력용기)":
                "ASME Section VIII Division 1, UW-51/52 - Radiographic Examination (압력용기 RT)\n\n적용 범위:\n압력용기 동체·헤드·노즐 용접부 RT.\n\n주요 요구사항:\n- UW-11: 필수 RT 대상 용접부 결정 기준\n  · P-No.1 ~ 15 재질 및 두께 기준\n- UW-51 (전수 RT): 전체 용접 길이 100% 검사\n  · 1.0 이음효율(E=1.0) 적용 가능\n- UW-52 (부분 RT): 용접 길이의 일부 검사\n  · 0.85 이음효율(E=0.85) 적용\n- 수용 기준: UW-51(b) - 균열, 미융합, 불완전 용입 불허\n  · 기공: Table UW-51 면적 기준\n- 기록: 투과사진(필름 또는 디지털) 3년 이상 보존",
            "MT 표준 항목":
                "MT (Magnetic Particle Testing) 표준 항목\n\n- 적용 코드: ASME Section V Article 7 (또는 프로젝트 지정 기준).\n- 자분 방식: 건식/습식, 형광/비형광, 자화 방법 및 방향성 확인.\n- 평가 기준: 해당 제작 코드의 표면 결함 수용 기준 적용.\n- 기록 및 보고: 자화 조건, 사용 매질, 지시 길이/위치 및 판정 포함.",
            "ASME Section V, Article 7 (MT 기본 절차)":
                "ASME Section V, Article 7 - Magnetic Particle Examination (MT 기본 절차)\n\n적용 범위:\n자분탐상검사(MT)의 핵심 절차 요구사항.\n\n주요 요구사항:\n- 자화 방법:\n  · 연속법(Continuous Method): 자화 유지 중 자분 적용\n  · 잔류법(Residual Method): 자화 후 자분 적용 (고보자력 재질)\n- 자화 방식: 요크(Yoke), 코일, 헤드샷, 프로드(Prod) 선택\n- 자분: 건식(Dry Powder) 또는 습식(Wet Suspension), 형광·비형광\n- 자장 강도: 요크 - 4.5kgf(AC) 또는 18kgf(DC) 리프팅 파워 확인\n- 검사 방향: 최소 2방향(90° 교차) 자화 필요\n- 조명: 가시광선 MT - 최소 100fc(1000 lux), 형광 MT - 최대 2fc 암실\n- 온도: 자분 적용 표면 10~52°C 범위\n\n검사원 자격: ASNT SNT-TC-1A Level II 이상",
            "ASME B31.1, Para. 136.4 (MT - 발전 배관)":
                "ASME B31.1, Paragraph 136.4 - Magnetic Particle Examination (발전 배관 MT)\n\n적용 범위:\n발전소 배관 용접부 표면 결함 MT 검사.\n\n주요 요구사항:\n- 적용 부위: 표면 및 표면 직하 결함 검출 (최대 3mm 깊이)\n- 절차: ASME Section V Article 7 준수\n- 수용 기준: Table 136.4.1 (선형·원형 지시 기준)\n  · 선형 지시: 1.6mm 이상 불허\n  · 원형 지시: 직경 4.8mm 이상 불허\n- 자화 방법: 요크, 코일, 프로드 (프로드는 균열 민감 재질 주의)\n- 후처리: 검사 후 잔류 자분 제거 및 탈자 필요 시 실시",
            "ASME B31.3, Para. 344.3 (MT - 공정 배관)":
                "ASME B31.3, Paragraph 344.3 - Magnetic Particle Examination (공정 배관 MT)\n\n적용 범위:\n공정 배관 용접부 및 모재 표면 결함 MT.\n\n주요 요구사항:\n- 절차: ASME Section V Article 7 기반\n- 검사 범위: Table 341.3.2 기준 (배관 등급별 MT 비율)\n- 수용 기준: Table 341.3.2 또는 Appendix A 적용\n  · 선형 지시 2mm 이상, 원형 지시 5mm 이상 불허\n- 강자성체 재질(탄소강, 저합금강) 전용\n- 오스테나이트계 스테인리스강: PT 적용 (MT 불가)\n- 후열처리 후 잔류 자분 및 자장 제거 확인",
            "ASME Sec. VIII Div.1, App. 6 (MT - 압력용기)":
                "ASME Section VIII Division 1, Mandatory Appendix 6 - MT (압력용기)\n\n적용 범위:\n압력용기 용접부 및 표면 결함 MT.\n\n주요 요구사항:\n- 적용 시점: 용접 완료 후, 최종 열처리 후 검사\n- 검사 면적: 용접부 + 양쪽 열영향부 각 13mm 포함\n- 자화 방법: 요크(AC 권장), 코일, 헤드샷\n- 자분: 형광 습식 자분 권장 (감도 우수)\n- 수용 기준: Appendix 6, Para. 6-5 적용\n  · 선형 지시 1.6mm, 원형 지시 4.8mm 이상 불허\n  · 4개 이상 지시 열 배열 불허\n- 기록: 지시 위치, 크기, 자화 조건, 판정 포함",
            "PT 표준 항목":
                "PT (Penetrant Testing) 표준 항목\n\n- 적용 코드: ASME Section V Article 6 (또는 프로젝트 지정 기준).\n- 절차: 전처리, 침투 시간, 제거, 현상, 관찰 시간 조건 준수.\n- 평가 기준: 해당 제작 코드의 표면 결함 수용 기준 적용.\n- 기록 및 보고: 재료, 표면 상태, 지시 위치/크기 및 판정 포함.",
            "ASME Section V, Article 6 (PT 기본 절차)":
                "ASME Section V, Article 6 - Liquid Penetrant Examination (PT 기본 절차)\n\n적용 범위:\n침투탐상검사(PT) 핵심 절차 요구사항. 비자성체 포함 모든 금속 재질 적용 가능.\n\n주요 요구사항:\n- PT 시스템 종류:\n  · Type 1 (형광): 자외선(UV-A, 320~400nm) 조사 하 관찰\n  · Type 2 (가시광선): 가시광선 하 관찰\n- 제거 방법:\n  · Method A: 수세성(Water Washable)\n  · Method B: 후유화성(Post-emulsifiable, 지용성)\n  · Method C: 용제 제거성(Solvent Removable)\n  · Method D: 후유화성 (수용성)\n- 절차 단계: 전처리 → 침투 → 제거 → 현상 → 관찰 → 후처리\n- 침투 시간: 최소 5분 (재질·온도에 따라 조정)\n- 현상 시간: 10분 이상 (비수성 습식 현상제)\n- 온도: 표면 10~52°C 범위\n- 조명: 가시광선 PT - 최소 100fc(1000 lux)\n\n검사원 자격: ASNT SNT-TC-1A Level II 이상",
            "ASME B31.1, Para. 136.4 (PT - 발전 배관)":
                "ASME B31.1, Paragraph 136.4 - Liquid Penetrant Examination (발전 배관 PT)\n\n적용 범위:\n발전소 배관 용접부 표면 개구 결함 PT 검사. 비자성 재질(STS, 니켈합금 등) 주로 적용.\n\n주요 요구사항:\n- 절차: ASME Section V Article 6 준수\n- 검사 대상: 용접부 표면 + 열영향부 각 13mm 이상\n- 수용 기준: Table 136.4.1 (선형·원형 지시 기준)\n  · 선형 지시(장축/단축 ≥ 3): 1.6mm 이상 불허\n  · 원형 지시: 직경 4.8mm 이상 불허\n  · 4개 이상 지시 열 배열 불허\n- 적용 시점: 용접 완료 후 최소 24시간 경과 권장 (지연 균열 고려)\n- 후처리: 검사 후 침투제·현상제 완전 제거",
            "ASME B31.3, Para. 344.4 (PT - 공정 배관)":
                "ASME B31.3, Paragraph 344.4 - Liquid Penetrant Examination (공정 배관 PT)\n\n적용 범위:\n공정 배관 용접부·모재 표면 결함 PT. 오스테나이트계·비철금속 배관에 주로 사용.\n\n주요 요구사항:\n- 절차: ASME Section V Article 6 기반\n- 검사 범위: Table 341.3.2 기준 (배관 등급별 PT 비율)\n- 수용 기준: Table 341.3.2 또는 Appendix A\n  · 선형 지시 2mm 이상, 원형 지시 5mm 이상 불허\n- 적용 대상: 비자성체(STS 304/316, Inconel, 알루미늄 등)\n- 저온 PT: 10°C 미만 시 특수 저온 침투제 사용 및 별도 검증\n- 기록: 침투제 종류·로트번호, 침투 시간, 지시 위치·크기, 판정",
            "ASME Sec. VIII Div.1, App. 8 (PT - 압력용기)":
                "ASME Section VIII Division 1, Mandatory Appendix 8 - PT (압력용기)\n\n적용 범위:\n압력용기 용접부 및 표면 결함 PT.\n\n주요 요구사항:\n- 적용 시점: 용접 완료 후, 열처리 완료 후 최종 검사\n- 검사 면적: 용접부 + 양쪽 각 13mm 열영향부 포함\n- PT 방법: 형광 PT 권장 (감도 우수), 가시광선 PT 허용\n- 제거 방법: 용제 제거성(Method C) 또는 수세성(Method A) 적용\n- 수용 기준: Appendix 8, Para. 8-4\n  · 선형 지시 1.6mm, 원형 지시 4.8mm 이상 불허\n- 검사원: ASNT Level II 이상\n- 기록: 지시 위치·크기, 침투제 정보, 검사 조건 및 판정",
            "PMI 표준 항목":
                "PMI (Positive Material Identification) 표준 항목\n\n- 적용 기준: 프로젝트 재질 관리 절차 및 관련 코드 요구사항.\n- 장비: XRF/OES 등 교정된 장비 사용, 표준 시편으로 정확도 확인.\n- 판정: 규정 재질 성분 범위와 비교하여 적합/부적합 판정.\n- 기록 및 보고: 부재 식별번호, 측정값, 장비 정보, 검사자 및 판정 포함.",
            "API RP 578 (PMI - 공정 배관·압력용기)":
                "API Recommended Practice 578 - Material Verification Program for New and Existing Alloy Piping Systems\n\n적용 범위:\n공정 배관 및 압력용기 합금 재질 검증 프로그램의 핵심 기준.\n\n주요 요구사항:\n- 적용 대상: 합금강(P-No. 3 이상), 스테인리스강, 니켈합금, 이종금속 용접부\n- PMI 장비:\n  · XRF(X선 형광분석): 비파괴, 현장 적용 용이, 경원소 검출 한계\n  · OES(광학 방출 분광): 파괴적, 탄소 검출 가능\n- 검사 범위:\n  · 신규 배관: 합금 성분 재질 100% PMI 권장\n  · 기존 배관: 위험 기반 PMI(Risk-Based PMI) 프로그램 수립\n- 판정: 재질 규격(ASTM, ASME) 성분 범위와 비교\n- 기록: 부재 번호, 측정값, 장비 S/N, 검사자, 날짜, 판정\n- 불합격: 즉시 격리, 재질 확인 후 교체 또는 재시험",
            "ASME B31.3, Para. 342.2 (PMI - 공정 배관 재질 관리)":
                "ASME B31.3, Paragraph 342.2 - Material Identification (공정 배관 PMI)\n\n적용 범위:\n공정 배관 재질 식별 및 검증 요구사항.\n\n주요 요구사항:\n- 재질 확인 의무: 설계 사양과 다른 재질 혼용 방지\n- PMI 적용 대상:\n  · 합금강 배관 구성품(파이프, 피팅, 플랜지, 밸브)\n  · 이종금속 용접부\n- 장비: XRF 또는 OES 교정된 장비 사용\n- 검사 시점: 제작 중, 설치 전, 최종 검사 단계\n- 적용 코드와 병행: API RP 578 권고사항 반영\n- 기록 유지: 전체 PMI 결과를 배관 사양서(Line List)와 연계 관리\n- 재질 표식: PMI 확인 후 색상 코딩 또는 금속 태그 부착",
            "ASTM E1476 (PMI - XRF 기법)":
                "ASTM E1476 - Standard Guide for Metals Identification, Grade Verification, and Sorting (PMI XRF 기법)\n\n적용 범위:\nXRF(X선 형광분석)를 이용한 금속 재질 식별·등급 검증 가이드.\n\n주요 요구사항:\n- 장비 종류:\n  · 휴대형 XRF(pXRF): 현장 비파괴 분석 (Ni, Cr, Mo, V, Nb 등 검출)\n  · 벤치탑 XRF: 실험실용, 정밀도 높음\n- 교정: 매 측정 전 NIST 추적 가능 인증 표준 시편으로 교정\n- 측정 조건: 측정 시간, 면적, 표면 상태(스케일 제거 필요) 설정\n- 한계:\n  · 탄소(C), 황(S), 인(P): XRF로 검출 불가 → OES 병행 필요\n  · 도막·산화층: 표면 연마 후 측정\n- 재질 판정: ASTM, ASME, EN 규격 성분 데이터베이스 비교\n- 기록: 측정값, 장비 모델·S/N, 교정 결과, 검사자, 날짜",
            "ASME PCC-2 (PMI - 수리·교체 재질 검증)":
                "ASME PCC-2 - Repair of Pressure Equipment and Piping (수리·교체 PMI)\n\n적용 범위:\n압력 기기·배관 수리 및 교체 시 재질 검증 요구사항.\n\n주요 요구사항:\n- 적용 시점: 수리 전 기존 모재 확인, 수리 후 교체 재질 최종 검증\n- PMI 필수 대상:\n  · 합금강·고합금 재질 수리 부위\n  · 이종금속 용접 수리\n  · 재질 불명 부품 교체\n- 장비: XRF(주) + OES(탄소 확인 필요 시 보완)\n- 절차:\n  1) 기존 모재 재질 확인\n  2) 수리 재료 입고 검증\n  3) 용접 완료 후 용착금속 및 HAZ 인접 모재 재확인\n- 기록: 수리 전·후 PMI 결과, 장비 정보, 검사자, 판정 포함\n- 불합격 조치: 즉시 작업 중단, 재질 재확인 후 적합 재료로 교체",
            "ASME Section VIII Div.1 vs Div.2 비교 (설계·재료·적용)":
                "ASME Section VIII Div.1 vs Div.2 - 설계·재료·적용 종합 비교\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 설계 철학\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "┌──────────────────┬──────────────────┬──────────────────┐\n"
                "│ 항목             │ Div.1            │ Div.2            │\n"
                "├──────────────────┼──────────────────┼──────────────────┤\n"
                "│ 설계 방식        │ Design by Rule   │ Design by Analy- │\n"
                "│                  │ (규칙 기반)      │ sis (해석 기반)  │\n"
                "│ 안전계수(UTS)    │ 4.0              │ 3.0              │\n"
                "│ 허용응력(Sm)     │ UTS / 4          │ UTS / 3          │\n"
                "│ 피로 해석        │ 불필요           │ 필수 (Screening  │\n"
                "│                  │                  │  → 상세 해석)    │\n"
                "│ 응력 분류        │ 없음             │ 1차/2차/피크 필수│\n"
                "│ 이음 효율 (E)    │ 0.65 ~ 1.0       │ 항상 1.0         │\n"
                "│ 설계 문서        │ 간략 계산서      │ 상세 설계 보고서 │\n"
                "│ Code Stamp       │ U Stamp          │ U2 Stamp         │\n"
                "│ 제작 비용        │ 낮음             │ 높음             │\n"
                "│ 재료 두께        │ 두꺼운 쉘        │ 더 얇은 쉘 가능  │\n"
                "└──────────────────┴──────────────────┴──────────────────┘\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 적용 대상\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "Div.1 적합:\n"
                "  - 소형·중압 일반 압력용기\n"
                "  - 반복 하중 없는 정상운전\n"
                "  - 설계 기간 단축 필요 시\n"
                "Div.2 적합:\n"
                "  - 고압·대형 용기 (두께 절감 효과 > 해석 비용)\n"
                "  - 반복 하중·사이클 피로 적용 서비스\n"
                "  - 핵·화학 플랜트 등 고신뢰성 요구\n\n"
                "실무 판단 기준:\n"
                "  설계 압력이 높고 쉘 두께가 두꺼울수록\n"
                "  Div.2 재료 절감 효과가 해석 비용 초과 → Div.2 경제적 우위",
            "ASME Section VIII Div.1 vs Div.2 비교 (비파괴검사)":
                "ASME Section VIII Div.1 vs Div.2 - 비파괴검사(NDT) 요건 상세 비교\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 1. NDT 전체 요건 개요\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "┌──────────────────────┬─────────────────────┬─────────────────────┐\n"
                "│ 항목                 │ Div.1               │ Div.2               │\n"
                "├──────────────────────┼─────────────────────┼─────────────────────┤\n"
                "│ 근거 조항            │ Part UW (용접 요건) │ Part 7 (검사 요건)  │\n"
                "│ RT/UT 적용 범위      │ 조건부 (UW-11 기준) │ 전수 100% 필수      │\n"
                "│ 이음 효율(E) 연동    │ E=0.65/0.85/1.0 선택│ 항상 E=1.0          │\n"
                "│ RT 완전 생략 가능    │ E=0.65 적용 시 가능 │ 불가                │\n"
                "│ PAUT 허용 근거       │ App.12 + CC 2235    │ Para.7.5.5 (명시)   │\n"
                "│ TOFD 요구            │ 없음                │ 고두께 권고         │\n"
                "│ MT 근거              │ Mandatory App.6     │ Para.7.5.6          │\n"
                "│ PT 근거              │ Mandatory App.8     │ Para.7.5.7          │\n"
                "│ VT 근거              │ UW-38, UG-97        │ Para.7.5.1          │\n"
                "│ PMI                  │ 계약 조건 의존      │ 고합금 계열 필수    │\n"
                "│ 검사 계획서 (EP)     │ 불필요              │ 필수 제출           │\n"
                "│ Level III 승인       │ 권고                │ 절차서 개발·승인 필수│\n"
                "└──────────────────────┴─────────────────────┴─────────────────────┘\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 2. RT (방사선검사) 상세 비교\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "[Div.1 - UW-51 / UW-52]\n"
                "  적용 결정 기준 (UW-11):\n"
                "    - Category A, B 맞대기 용접부 → UW-11(a) 기준으로\n"
                "      전수(UW-51) 또는 부분(UW-52) 선택\n"
                "    - 재질 P-No., 두께, 설계 온도·압력 조건으로 결정\n"
                "  전수 RT (UW-51):\n"
                "    - 용접부 전체 길이 100% 검사\n"
                "    - 이음 효율 E=1.0 적용 → 두께 최소화\n"
                "    - 수용 기준: UW-51(b)\n"
                "      · 균열, 미융합, 용입불량: 전부 불합격\n"
                "      · 기공: Table UW-51 면적 기준\n"
                "      · 슬래그: 길이 ≤ t/3 (최대 6mm ~ 19mm)\n"
                "  부분 RT (UW-52):\n"
                "    - 용접부 길이의 일부 (최소 1 spot / 50ft)\n"
                "    - 이음 효율 E=0.85 적용 → 두께 약 18% 증가\n"
                "    - 수용 기준: UW-51(b) 동일 (해당 spot 기준)\n"
                "    - 불합격 spot 발생 시 → 전수 RT로 확대\n"
                "  RT 생략 (E=0.65):\n"
                "    - 검사 전혀 없음 → 두께 약 54% 증가로 보완\n"
                "    - 저압 비중요 용기에 한정 적용\n\n"
                "[Div.2 - Para.7.5.3]\n"
                "  - 모든 맞대기 용접부 (Category A, B): 전수 RT 또는 UT\n"
                "  - E=1.0으로 설계 → 검사 생략 시 설계 자체 무효화\n"
                "  - 100% UT(PAUT 포함)로 RT 대체 가능 (Para.7.5.5)\n"
                "  수용 기준 (Table 7.5.3-1):\n"
                "    · 균열, 미융합, 용입불량: 크기 무관 전부 불합격\n"
                "    · 기공: 최대 직경 3mm, 면적 기준 강화\n"
                "    · 슬래그/개재물: 길이 ≤ min(t/3, 6mm)\n"
                "    · 표면 언더컷: 깊이 ≤ 0.8mm\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 3. UT / PAUT 상세 비교\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "[Div.1 - Appendix 12 + Code Case 2235]\n"
                "  허용 조건:\n"
                "    - App.12 요건 충족 + 발주처 승인 시 RT 대체\n"
                "    - Code Case 2235: PAUT로 RT 완전 대체 허용\n"
                "      (두께 25mm 이상 맞대기 용접부 대상)\n"
                "  절차서 기준:\n"
                "    - ASME Section V Article 4\n"
                "    - App.XII (PAUT 전용 요건)\n"
                "  보정 블록:\n"
                "    - ASME Basic Calibration Block\n"
                "    - SDH 직경: Art.4 Table T-434.2.1 기준\n"
                "    - IIW Block: 굴절각·음속 확인용\n"
                "  감도 설정:\n"
                "    - DAC 곡선 작성 (기준 SDH → 각 깊이별 DAC점)\n"
                "    - Search Level: DAC +6dB\n"
                "    - 기록: DAC 20% / 평가: DAC 50% / 거부: DAC 100%\n"
                "  수용 기준:\n"
                "    - UW-51(b) 준용 또는 CC 2235 별도 기준 적용\n"
                "    - 결함 높이 측정: 선택적 (−6dB 또는 −20dB법)\n\n"
                "[Div.2 - Para.7.5.5 (UT / PAUT)]\n"
                "  허용 조건:\n"
                "    - Para.7.5.5에서 PAUT 명시적 허용\n"
                "    - 모든 맞대기 용접부 RT 대체 가능\n"
                "  절차서 기준:\n"
                "    - ASME Section V Article 4 + App.XII\n"
                "    - 상세 Scan Plan 첨부 필수\n"
                "  보정 블록:\n"
                "    - 검사 모재와 동일 재질·두께\n"
                "    - SDH 직경: Table T-434.2.1 + Para.7.5.5 요건 병행\n"
                "  감도 설정:\n"
                "    - TCG (Time Corrected Gain) 또는 DAC 적용 가능\n"
                "    - 전 깊이 범위에서 균일한 SDH 감도 유지\n"
                "  결함 크기 측정 (필수):\n"
                "    - 길이: −6dB 강하법 또는 −20dB 끝점법\n"
                "    - 높이: −6dB 강하법 (TOFD 병행 권고)\n"
                "    - 결함 치수 측정 결과를 Table 7.5.5-1에 대입\n"
                "  수용 기준 (Table 7.5.5-1):\n"
                "    평면형 결함 (균열, 미융합, 용입불량):\n"
                "      → 크기 무관 전부 불합격\n"
                "    체적형 결함 (기공, 슬래그):\n"
                "      · 결함 높이 a ≤ 3mm : 길이 ℓ ≤ 6mm 허용\n"
                "      · 3mm < a ≤ 6mm    : 길이 ℓ ≤ 2a 허용\n"
                "      · a > 6mm          : 불합격\n"
                "    표면 연결 결함:\n"
                "      · a ≤ 1.5mm : 길이 ℓ ≤ 6mm 허용\n"
                "      · a > 1.5mm : 불합격\n"
                "  디지털 데이터 보존:\n"
                "    - A-scan, S-scan, B-scan 전체 파일 보존 필수\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 4. MT (자분탐상검사) 상세 비교\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "[Div.1 - Mandatory Appendix 6]\n"
                "  적용 대상:\n"
                "    - 용접 완료 후 표면·표면 직하 결함 탐지\n"
                "    - PWHT(용접 후 열처리) 완료 후 재검사\n"
                "  검사 면적:\n"
                "    - 용접부 + 양쪽 HAZ 각 13mm (1/2인치) 포함\n"
                "  자화 방법:\n"
                "    - 요크(AC 권장): 표면 결함 감도 우수\n"
                "    - 코일법, 헤드샷: 원통형 부재 적용\n"
                "    - 최소 2방향 자화 필수 (직교 방향 결함 탐지)\n"
                "  자분 종류:\n"
                "    - 형광 습식 자분 (Fluorescent Wet): 감도 최우수\n"
                "    - 건식 자분 (Dry Powder): 고온 부위 허용\n"
                "  조도:\n"
                "    - 형광 자분: 자외선 UV-A 최소 1000 μW/cm²\n"
                "    - 가시광선 자분: 최소 100 fc (1076 lux)\n"
                "  수용 기준 (App.6, Para. 6-5):\n"
                "    - 선형 지시: 1.6mm 초과 불합격\n"
                "    - 원형 지시: 4.8mm 초과 불합격\n"
                "    - 지시 군집: 150mm² 이내 4개 이상 군집 불합격\n\n"
                "[Div.2 - Para.7.5.6]\n"
                "  적용 시점 추가:\n"
                "    - 피로 해석 대상 용접부: 최종 형상 가공 완료 후\n"
                "    - 고사이클 피로 부위: 모든 표면 처리 후 추가 MT\n"
                "  수용 기준 (Table 7.5.6-1):\n"
                "    - 선형 지시: 1.6mm 초과 불합격 (동일)\n"
                "    - 원형 지시: 4.8mm 초과 불합격 (동일)\n"
                "    - 표면 결함 높이 기준 병행 적용:\n"
                "      · 피로 해석 부위 → 높이 a ≤ 1.5mm 추가 요건\n"
                "  추가 요건:\n"
                "    - 압력 시험 후 접근 가능한 용접부 MT 재검사\n"
                "    - 검사 결과 디지털 기록 및 보존 의무\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 5. PT (침투탐상검사) 상세 비교\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "[Div.1 - Mandatory Appendix 8]\n"
                "  적용 대상:\n"
                "    - 비자성 재질 (오스테나이트 스테인리스, 알루미늄 등)\n"
                "    - MT 적용 불가 부위의 표면·개구 결함 탐지\n"
                "  PT 방법:\n"
                "    - Method A: 수세성 침투제\n"
                "    - Method C: 용제 제거성 (현장 적용 용이)\n"
                "    - Method D: 후유화성 (고감도 요구 시)\n"
                "  Type 선택:\n"
                "    - Type 1: 형광 PT (감도 우수, 암실 필요)\n"
                "    - Type 2: 가시광선 PT (현장 편의성)\n"
                "  침투 시간: 최소 5분 (재질·온도 따라 최대 60분)\n"
                "  현상 시간: 최소 10분 이상\n"
                "  수용 기준 (App.8, Para. 8-4):\n"
                "    - 선형 지시: 1.6mm 초과 불합격\n"
                "    - 원형 지시: 4.8mm 초과 불합격\n"
                "    - 지시 열배열: 150mm 내 4개 이상 불합격\n\n"
                "[Div.2 - Para.7.5.7]\n"
                "  추가 요건:\n"
                "    - 형광 PT (Type 1) 권장 (가시광선 PT는 동등 감도 입증 필요)\n"
                "    - 침투 시간: 재질별 최소 시간 표 준수\n"
                "    - 피로 해석 부위: 표면 거칠기 Ra ≤ 6.3μm 확보 후 검사\n"
                "  수용 기준 (Table 7.5.7-1):\n"
                "    - 선형·원형 지시 기준: Div.1과 동일\n"
                "    - 피로 해석 대상 부위 표면 결함:\n"
                "      · 균열성 선형 지시: 모두 불합격\n"
                "      · 깊이 a > 1.5mm 추정 지시: 불합격\n"
                "  추가:\n"
                "    - 압력 시험 후 PT 재검사 (접근 가능 용접부)\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 6. VT (육안검사) 비교\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "[Div.1 - UW-38, UG-97]\n"
                "  - 용접 중간 패스 육안 검사 (UW-38)\n"
                "  - 최종 수압 시험 중 육안 검사 (UG-99/UG-100)\n"
                "  - 검사거리: 600mm 이내, 조도 최소 50 fc\n"
                "[Div.2 - Para.7.5.1]\n"
                "  - 동일 + 용접 비드 형상, 표면 결함 상세 기준 추가\n"
                "  - 용접 완료 후 100% VT 의무 (중간 패스 포함)\n"
                "  - 최종 압력 시험 중 전체 외면 VT\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 7. 검사원 자격 상세 비교\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "┌──────────────────┬──────────────────────┬──────────────────────┐\n"
                "│ 항목             │ Div.1                │ Div.2                │\n"
                "├──────────────────┼──────────────────────┼──────────────────────┤\n"
                "│ 최소 자격        │ SNT-TC-1A Level II   │ SNT-TC-1A Level II   │\n"
                "│ 합부판정 권한    │ Level II / III       │ Level II / III       │\n"
                "│ 절차서 개발·승인 │ 권고                 │ Level III 필수       │\n"
                "│ 검사 계획서(EP)  │ 불필요               │ 필수 제출 및 승인    │\n"
                "│ PAUT 전용 훈련   │ 기록 권고            │ 별도 기록 의무       │\n"
                "│ AI 최종 확인     │ 필수                 │ 필수                 │\n"
                "│ 자격 갱신        │ 사내 프로그램        │ 사내 프로그램        │\n"
                "└──────────────────┴──────────────────────┴──────────────────────┘\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 8. 검사 기록 및 문서 요건 비교\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "[Div.1]\n"
                "  - RT 필름: 최소 3년 보존\n"
                "  - 검사 보고서: 장비, 검사원, 판정 결과 포함\n"
                "  - 디지털 RT: 동등 품질 입증 시 허용\n"
                "[Div.2]\n"
                "  - UT/PAUT 디지털 데이터: 전체 A-scan/S-scan 보존\n"
                "  - 검사 보고서: 결함 위치·크기·판정 포함 상세 작성\n"
                "  - Examination Program(EP): 검사 전 제출·승인\n"
                "  - 검사 완료 보고서: AI 서명 + Level III 확인\n"
                "  - 데이터 보존 기간: 기기 수명 기간 (최소 10년 이상)",
            "기타": "사용자 정의 텍스트를 입력하세요.",

            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            # ISO PAUT 관련 코드
            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            "ISO 13588 (PAUT - 용접부 검사)":
                "ISO 13588:2019 - Non-destructive testing of welds - Ultrasonic testing - Use of automated phased array technology\n\n적용 범위:\n용접부 자동 위상배열 초음파검사(PAUT) 국제표준. ASME와 병행 사용.\n\n주요 요구사항:\n- 적용 두께: 6mm 이상 금속 용접부\n- 검사 레벨:\n  · Level A: 기본 감도, 일반 산업 적용\n  · Level B: 향상된 감도, 중요 구조물\n  · Level C: 최고 감도, 원자력·안전 중요 기기\n- 스캔 방식: 선형(Linear), 섹터(Sectorial), 복합 스캔\n- 보정 블록: 동일 재질, SDH 직경 = 검사 두께의 1/40 (최소 1mm)\n- 커버리지: 전체 용접 단면 체적 100% 스캔\n- 기록: 교정 기록, 스캔 데이터(A-scan/S-scan), 결함 지시 기록\n\n합부판정:\n- ISO 11666 (수용 기준 코드)과 연계 적용\n- 검사 레벨에 따른 기록·평가 수준 차등 적용",
            "ISO 11666 (PAUT 합부판정 수용 기준)":
                "ISO 11666:2018 - Non-destructive testing of welds - Ultrasonic testing - Acceptance levels\n\n적용 범위:\n초음파검사(UT/PAUT) 용접부 합부판정 수용 기준 국제표준.\n\n수용 레벨 (Acceptance Levels):\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ Level 1 (엄격)\n- 평면형 결함 (균열, 용합불량, 용입불량):\n  · 모든 지시 불합격\n- 체적형 결함:\n  · 반사 진폭 ≥ 기준 반사체(SDH)의 100%: 불합격\n  · 길이: 10mm 이상 불합격\n\n■ Level 2 (표준)\n- 평면형 결함: 모든 지시 불합격\n- 체적형 결함:\n  · 반사 진폭 ≥ 기준의 100%: 불합격\n  · 길이: 두께별 기준 (t < 15mm → 15mm, t ≥ 15mm → t)\n\n■ Level 3 (완화)\n- 평면형 결함: 모든 지시 불합격\n- 체적형 결함:\n  · 반사 진폭 ≥ 기준의 100%: 불합격\n  · 길이: 25mm 또는 t 중 큰 값 이상 불합격\n\n기록 수준: 기준 반사체 대비 −6dB 이상 모든 지시 기록\n평가 수준: 기준 반사체 대비 0dB 초과 지시 평가",
            "ISO 19285 (PAUT - 합부판정 대체 기준)":
                "ISO 19285:2017 - Non-destructive testing of welds - Phased array ultrasonic testing (PAUT) - Acceptance levels\n\n적용 범위:\nPAUT 전용 합부판정 국제표준. ISO 11666의 PAUT 특화 보완 기준.\n\n주요 내용:\n- PAUT 결함 크기 측정 방법 규정:\n  · -6dB 강하법 (Half Maximum Amplitude Method)\n  · -20dB 강하법 (끝점 탐지법)\n  · TOFD (Time of Flight Diffraction) 병행 가능\n- 결함 높이(a) 기반 합부판정:\n  · a < 0.5t 이고 a < 6mm: 길이 기준 추가 적용\n  · 표면 연결 결함: 길이 기준 강화\n- 스캔 인덱스 포인트 정확도: ±1mm 이내\n- 검사 레벨 (ISO 13588 Level A/B/C)과 연계\n\n비고: ASME Code Case 2235와 비교하여 적용 가능",
            "ISO PAUT 파괴역학 합부판정 (BS 7910 / ECA)":
                "ISO/BS 기반 PAUT 파괴역학 합부판정\n"
                "BS 7910 & ECA (Engineering Critical Assessment)\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ BS 7910 개요\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "BS 7910:2019 - Guide to methods for assessing the\n"
                "acceptability of flaws in metallic structures\n\n"
                "파괴역학 이론을 적용하여 결함의 실제 구조 안전성을\n"
                "평가하는 국제 표준 방법론. ISO 국제화 진행 중.\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ ECA 적용 레벨 (3단계)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "Level 1 (간편법 - 보수적):\n"
                "  - 별도 재료 물성 없이 표준 FAD 곡선 적용\n"
                "  - 간단한 응력 상태, 보수적 안전계수\n"
                "  - 결과: 합격/불합격 여부만 판단\n\n"
                "Level 2 (표준법 - 일반 적용):\n"
                "  - 재료의 항복강도, 인장강도, 파괴인성(K_IC 또는 J_IC) 필요\n"
                "  - FAD 식: f(Lr) = (1 - 0.14·Lr²)[0.3 + 0.7·exp(-0.65·Lr⁶)]\n"
                "  - Kr = K_I / K_mat ≤ f(Lr)  → FAD 내부 시 합격\n"
                "  - Lr = σ_ref / σ_Y  (소성 붕괴 비율)\n\n"
                "Level 3 (정밀법 - J 적분 / R-곡선):\n"
                "  - J 적분 또는 CTOD 파괴인성 데이터 필요\n"
                "  - 연성 찢김(Ductile Tearing) 포함 평가\n"
                "  - 피로·크리프 복합 손상 평가 가능\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 결함 특성화 (PAUT 측정값 입력)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "PAUT → ECA 입력 파라미터:\n"
                "  - 결함 높이 (a): TOFD / -6dB drop 법\n"
                "  - 결함 길이 (2c): -6dB 또는 -20dB 끝점법\n"
                "  - 결함 깊이 (d): 표면 기준 결함 상단 위치\n"
                "  - 결함 형상: 타원형 (a/c ≤ 1 가정)\n"
                "  - 표면 연결 여부: 표면 결함 vs 매립 결함 구분\n\n"
                "결함 보수화 (Flaw Characterisation):\n"
                "  - 측정 불확도 고려: a_char = a_measured + δa (측정 불확도)\n"
                "  - 인접 결함 상호작용 규칙 (Interaction Criterion):\n"
                "    두 결함 간격 < min(a1, a2) 시 합산 처리\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 피로 균열 성장 평가 (BS 7910 Annex A)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "  da/dN = A × ΔK^m  (Paris-Erdogan 법칙)\n"
                "  - ΔK_th 이하: 성장 없음\n"
                "  - 초기 결함 크기(a_0) → 설계 수명 후 최종 크기(a_f)\n"
                "  - a_f ≤ a_critical 조건 만족 시 합격\n"
                "  → PAUT 재검사 주기 결정에 활용\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 관련 코드 / 표준\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "  - BS 7910:2019 (파괴역학 기반 ECA 주요 기준)\n"
                "  - API 579-1/ASME FFS-1 : Part 9 (균열형 결함)\n"
                "  - SINTAP / FITNET : 유럽 ECA 방법론\n"
                "  - ISO 15653 : 파괴인성 시험 (용접부 J/CTOD)\n"
                "  - ISO 12135 : K_IC 시험법\n"
                "  - DNV-RP-C210 : 해양 구조물 피로 ECA",

            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            # ISO RT 관련 코드
            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            "ISO 17636-1 (RT - 필름 방사선 투과)":
                "ISO 17636-1:2013 - Non-destructive testing of welds - Radiographic testing - Part 1: X- and gamma-ray techniques with film\n\n적용 범위:\n필름 방사선투과검사(RT) 국제표준.\n\n주요 요구사항:\n- 검사 기법 등급:\n  · Class A (기본): 일반 산업 적용\n  · Class B (향상): 항공, 압력 기기, 안전 중요 구조물\n- 선원 종류: X선관, Ir-192, Se-75, Co-60, Yb-169\n- IQI 종류:\n  · Wire IQI (ISO 19232-1 기준)\n  · Step/Hole IQI (ISO 19232-2 기준)\n- 최소 감도: Class A - W13 wire, Class B - W14 wire 이상\n- 필름 종류: ISO 11699-1 기준 C3~C5 (Class A), C4~C6 (Class B)\n- 노출 기하: 선원-피사체 거리(f), 필름-피사체 거리(b) 규정\n- 암실 처리: 온도·시간 기준 준수\n\n검사원 자격: ISO 9712 Level 2 이상",
            "ISO 17636-2 (디지털 RT - DR/CR)":
                "ISO 17636-2:2013 - Non-destructive testing of welds - Radiographic testing - Part 2: X- and gamma-ray techniques with digital detectors\n\n적용 범위:\n디지털 방사선투과검사(DR: Digital Radiography, CR: Computed Radiography) 국제표준.\n\n주요 요구사항:\n- 검사 기법 등급: Class A / Class B (필름 RT와 동일 구분)\n- 디지털 검출기 종류:\n  · DR (Flat Panel Detector): 직접·간접 변환 방식\n  · CR (Storage Phosphor Imaging Plate): 이미징 플레이트\n- 기본 공간 해상도(BSR): IQI 감도와 함께 검증 필수\n- 콘트라스트 감도: 동일 두께 필름 RT와 동등 이상\n- 데이터 형식: DICONDE 형식 권장, 무손실 압축만 허용\n- 시스템 검증: 처음 사용 시 및 주요 설정 변경 시 검증\n- 영상 판독: 보정된 모니터(최소 2MP, ISO 14396 기준) 사용",
            "ISO 10675-1 (RT 합부판정 - 강재)":
                "ISO 10675-1:2016 - Non-destructive testing of welds - Acceptance levels for radiographic testing - Part 1: Steel, nickel, titanium and their alloys\n\n적용 범위:\n강재·니켈합금·티타늄 용접부 RT 합부판정 국제표준.\n\n수용 레벨:\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ Level 1 (엄격)\n- 균열, 용합불량, 용입불량: 전부 불합격\n- 기공: 단독 기공 직경 ≤ 0.25t (최대 3mm), 군집 면적 ≤ 1%\n- 슬래그: 길이 ≤ t/3 (최대 10mm)\n\n■ Level 2 (표준 - 일반 적용)\n- 균열, 용합불량, 용입불량: 전부 불합격\n- 기공: 단독 기공 직경 ≤ 0.3t (최대 4mm), 군집 면적 ≤ 2%\n- 슬래그: 길이 ≤ 0.6t (최대 25mm)\n\n■ Level 3 (완화)\n- 균열: 불합격\n- 기공: 단독 기공 직경 ≤ 0.4t (최대 6mm), 군집 면적 ≤ 4%\n- 슬래그: 길이 ≤ t (최대 50mm)\n- 용입불량: 길이 ≤ t/4",

            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            # ISO MT 관련 코드
            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            "ISO 17638 (MT - 용접부 자분탐상)":
                "ISO 17638:2016 - Non-destructive testing of welds - Magnetic particle testing\n\n적용 범위:\n용접부 자분탐상검사(MT) 국제표준.\n\n주요 요구사항:\n- 검사 기법:\n  · 형광 MT (UV-A 365nm 조사)\n  · 비형광(가시광선) MT\n- 자화 방법: 요크(Yoke), 헤드샷, 코일, 프로드, 전류 직접통전\n- 연속법 vs 잔류법 선택 기준 명시\n- 자장 강도 확인: 자장 지시계(Field Indicator) 또는 홀(Hall) 효과 가우스미터\n  · 연속법: 표면 접선 자장 2~6 kA/m\n- 자분 입자: 건식 또는 습식 (ISO 9934-2 기준)\n- 탐지 매질(Contrast Aid): 필요 시 흰색 반사 도막 사용\n- 조명:\n  · 형광 MT: UV-A 조도 ≥ 10 W/m², 주변 가시광 ≤ 20 lux\n  · 비형광 MT: 표면 조도 ≥ 500 lux\n\n검사원 자격: ISO 9712 Level 2 이상",
            "ISO 23278 (MT 합부판정)":
                "ISO 23278:2015 - Non-destructive testing of welds - Magnetic particle testing - Acceptance levels\n\n적용 범위:\n용접부 MT 합부판정 국제표준.\n\n수용 레벨:\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ Level 1 (엄격)\n- 선형 지시 (Linear Indications, 길이/폭 ≥ 3:1):\n  · 길이 > 1.5mm: 불합격\n- 비선형 지시 (Non-linear):\n  · 장축 > 3mm: 불합격\n- 열 배열 지시: 3개 이상 / 구간 내 합계 > 3mm 불합격\n\n■ Level 2 (표준)\n- 선형 지시: 길이 > 3mm 불합격\n- 비선형 지시: 장축 > 5mm 불합격\n- 열 배열 지시: 4개 이상 / 구간 내 합계 > 6mm 불합격\n\n■ Level 3 (완화)\n- 선형 지시: 길이 > 6mm 불합격\n- 비선형 지시: 장축 > 8mm 불합격\n\n비고: 균열 지시는 레벨 무관 전부 불합격",

            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            # ISO PT 관련 코드
            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            "ISO 3452-1 (PT - 침투탐상 기본 절차)":
                "ISO 3452-1:2021 - Non-destructive testing - Penetrant testing - Part 1: General principles\n\n적용 범위:\n침투탐상검사(PT) 일반 원칙 국제표준.\n\n주요 요구사항:\n- PT 계열 (System):\n  · Type I: 형광 침투제 (UV-A 하 관찰)\n  · Type II: 가시광선 침투제\n- 제거 방법:\n  · Method A: 수세성 (Water Washable)\n  · Method B: 후유화성 지용성 (Lipophilic Post-emulsifiable)\n  · Method C: 용제 제거성 (Solvent Removable)\n  · Method D: 후유화성 수용성 (Hydrophilic Post-emulsifiable)\n- 현상제 종류: 건식 분말, 수용성, 비수성 습식, 특수 형광\n- 침투 시간: 최소 5분 (재질·온도 따라 조정, ISO 3452-4 참조)\n- 현상 시간: 10분 이상 (비수성 현상제 기준)\n- 온도: 10~50°C 범위\n- 시약 계열 적합성: ISO 3452-2 기준 동일 제조사 계열 사용 원칙\n\n검사원 자격: ISO 9712 Level 2 이상",
            "ISO 23277 (PT 합부판정)":
                "ISO 23277:2015 - Non-destructive testing of welds - Penetrant testing - Acceptance levels\n\n적용 범위:\n용접부 PT 합부판정 국제표준.\n\n수용 레벨:\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ Level 1 (엄격)\n- 선형 지시 (Linear, 길이/폭 ≥ 3:1):\n  · 길이 > 1.5mm: 불합격\n- 비선형 지시:\n  · 장축 > 3mm: 불합격\n- 열 배열 지시: 3개 이상 / 합계 > 3mm 불합격\n\n■ Level 2 (표준)\n- 선형 지시: 길이 > 3mm 불합격\n- 비선형 지시: 장축 > 4mm 불합격\n- 열 배열 지시: 4개 이상 / 합계 > 6mm 불합격\n\n■ Level 3 (완화)\n- 선형 지시: 길이 > 6mm 불합격\n- 비선형 지시: 장축 > 8mm 불합격\n\n비고:\n- 균열 지시: 레벨 무관 전부 불합격\n- 표면 개구 여부 확인 후 선형/비선형 분류\n- 검사 후 침투제·현상제 완전 제거 확인",

            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            # ISO PMI 관련 코드
            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            "ISO 15011-4 (PMI - 합금 재질 식별 가이드)":
                "ISO 15011-4 - Material identification and verification methods (합금 재질 식별)\n\n적용 범위:\n합금 금속 재질 식별 및 검증을 위한 국제 가이드라인.\n\n주요 내용:\n- 재질 식별 방법 분류:\n  · XRF (X-ray Fluorescence): 비파괴, 현장 신속 분석\n  · OES (Optical Emission Spectrometry): 파괴적, 탄소 검출 가능\n  · LIBS (Laser-Induced Breakdown Spectroscopy): 비파괴, 소형화 가능\n- 분석 원소 범위:\n  · XRF: Mg ~ U (Z=12~92), 탄소·질소·산소 검출 불가\n  · OES: C, Si, Mn, P, S, Cr, Ni, Mo 등 전원소 분석 가능\n- 교정 요건: NIST/PTB 추적 인증 표준 시편 사용\n- 측정 불확도: 주요 합금 원소 ±0.05wt% 이내 권장\n- 재질 판정: ASTM, EN, JIS, KS 규격 성분 범위와 비교\n- 기록: 측정값, 장비 정보, 교정 기록, 검사자, 날짜 포함",
            "ISO 9712 (NDT 검사원 자격 인증)":
                "ISO 9712:2021 - Non-destructive testing - Qualification and certification of NDT personnel\n\n적용 범위:\n비파괴검사(NDT) 전 종목 검사원 자격 부여 및 인증 국제표준.\n적용 종목: PT, MT, RT, UT(PAUT 포함), ET(와전류), VT, ST, LT 등\n\n자격 레벨:\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ Level 1\n- 지정된 NDT 지시서에 따라 검사 수행\n- 검사 결과 기록 (합부판정 권한 없음)\n- 자격 요건: 산업 교육 + 현장 경험 + 필기·실기 시험\n\n■ Level 2 (현장 주력 자격)\n- 검사 절차 설정 및 수행\n- 합부판정 (해당 코드 기준 적용)\n- Level 1 지도·감독\n- 자격 요건: Level 1 경험 + 교육 시간 + 시험 합격\n\n■ Level 3\n- NDT 절차·기술 개발 및 승인\n- 합부판정 기준 해석\n- Level 1/2 자격 시험 감독 및 인증\n- 자격 요건: 학력 + 광범위한 실무 경험 + 종합 시험\n\n인증 유효기간: 5년 (중간 재확인 + 갱신 시험)\n인증 기관: 각국 ISO 9712 인정 인증 기관 (한국: KSNT 등)",

            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            # PAUT 보정 블록 (Calibration Block)
            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            "PAUT 보정 블록 종류 및 재질 요건":
                "PAUT Calibration Block - 종류 및 재질 요건\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 보정 블록 종류\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "┌─────────────┬──────────────┬────────────────────────┐\n"
                "│ 블록         │ 규격          │ 특징                    │\n"
                "├─────────────┼──────────────┼────────────────────────┤\n"
                "│ IIW Block   │ ISO 2400     │ 범용, UT/PAUT 공통      │\n"
                "│ ASME Basic  │ Sec.V Art.4  │ 미국 코드 기본          │\n"
                "│ SDH Block   │ 프로젝트 제작 │ PAUT 주력 기준 반사체   │\n"
                "│ 사용자 제작  │ 동일재질 제작 │ 현장 맞춤형            │\n"
                "└─────────────┴──────────────┴────────────────────────┘\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 재질 요건\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- 검사 대상과 동일 재질 (또는 음향 임피던스 동등 재질)\n"
                "- 표면 상태: 검사면과 동등 이상의 표면 거칠기\n"
                "- 열처리 이력: 검사 대상과 동일하게 적용 권장\n"
                "- 온도 보정: 보정 블록과 검사체 온도 차 ±14°C 초과 시 재보정",

            "PAUT 보정 블록 - SDH 직경 기준 (코드별 비교)":
                "PAUT Calibration Block - SDH 직경 기준 코드별 비교\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ ASME Section V, Article 4 — Table T-434.2.1 (기본 기준)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "┌─────────────────────┬────────────────────┐\n"
                "│ 검사 두께 (t)        │ SDH 직경            │\n"
                "├─────────────────────┼────────────────────┤\n"
                "│ t ≤ 1\"  (25mm)      │ Ø 3/32\" (≈ 2.4mm) │\n"
                "│ 1\" < t ≤ 2\" (50mm) │ Ø 1/8\"  (≈ 3.2mm) │\n"
                "│ 2\" < t ≤ 4\" (100mm)│ Ø 3/16\" (≈ 4.8mm) │\n"
                "│ t > 4\"  (100mm↑)   │ Ø 1/4\"  (≈ 6.4mm) │\n"
                "└─────────────────────┴────────────────────┘\n"
                "※ B31.1, B31.3 일반 적용 시 이 기준 사용\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ ASME Section VIII Div.1, Appendix 12\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- 기준 반사체: Ø 3/64\" (≈ 1.2mm) SDH\n"
                "- 배치 깊이: t/4, t/2, 3t/4\n"
                "- 근거 조항: App. 12, Para. 12-4\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ ASME Section VIII Div.2, Para. 7.5.5\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- 기준 반사체: Ø 2.0mm SDH\n"
                "- 배치 깊이: t/4, t/2, 3t/4\n"
                "- 감도 기준: SDH 에코 = 80% FSH (Full Screen Height)\n"
                "- 근거 조항: Table 7.5.5-1\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ ASME Code Case 2235 (RT 대체 강화 기준)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- t ≤ 50mm → Ø 1.5mm SDH\n"
                "- t > 50mm → Ø 2.0mm SDH\n"
                "- Art.4 기본값보다 더 작은 결함 탐지를 위한 강화 기준\n"
                "- 적용 여부: 프로젝트 계약 사양서(PO/Spec) 확인 필수\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ ISO 13588 (국제 표준)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- SDH 직경: t/40 (최소 1mm)\n"
                "- 검사 레벨 A/B/C에 따라 감도 차등 적용\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 적용 코드 선택 기준\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "일반 산업 배관·압력용기\n"
                "  └─ Art.4 Table T-434.2.1 SDH 기준 적용\n\n"
                "RT 완전 대체 (고신뢰도 요구)\n"
                "  └─ Code Case 2235 강화 기준 (1.5/2.0mm)\n\n"
                "원자력 (ASME Sec. XI)\n"
                "  └─ Appendix VIII PDI 실증 별도 요구",

            "PAUT 보정 절차 - DAC 감도 설정 및 보정 주기":
                "PAUT Calibration Procedure - DAC 감도 설정 및 보정 주기\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ DAC (Distance Amplitude Correction) 설정 절차\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "① 각 SDH에서 최대 에코 확보\n"
                "② 각 깊이별 최대 진폭 포인트 연결 → DAC 곡선 작성\n"
                "③ TCG (Time Corrected Gain) 적용 시:\n"
                "   모든 깊이에서 SDH 에코를 동일 높이(%)로 보정\n"
                "④ 감도 기준 설정:\n"
                "   ├─ 기록 수준: DAC 20% (−14 dB)\n"
                "   ├─ 평가 수준: DAC 50% (−6 dB)\n"
                "   └─ 거부 수준: DAC 100% 초과\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 보정 주기\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "┌───────────────────┬──────────────────────────┐\n"
                "│ 시점               │ 내용                      │\n"
                "├───────────────────┼──────────────────────────┤\n"
                "│ 검사 시작 전       │ 전체 보정 수행            │\n"
                "│ 매 4시간마다       │ 감도 확인 (CC 2235 기준)  │\n"
                "│ 매 8시간마다       │ ASME Art.4 기본 기준      │\n"
                "│ 검사 종료 후       │ 최종 보정 확인            │\n"
                "│ 장비 이동 시       │ 재보정                    │\n"
                "│ 온도 차 ±14°C 초과│ 온도 보정 실시            │\n"
                "└───────────────────┴──────────────────────────┘\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 보정 유효성 확인 (Calibration Check)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- 감도 변화 ±2 dB 이내         → 유효, 계속 검사\n"
                "- 감도 변화 ±2 dB 초과 ~ ±4 dB → 직전 보정 후 검사 부위 재검토\n"
                "- 감도 변화 ±4 dB 초과          → 직전 유효 보정 후 전체 재검사\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ S/N 비 요건\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- 신호 대 잡음비: ≥ 3:1 (9.5 dB) 유지 필수\n"
                "- 이 이하면 프로브·각도·감도 재조정 후 재보정",

            "PAUT 결함 탐지 · 위치 · 크기 측정 기준":
                "PAUT - 결함 탐지, 위치 확인, 크기 측정 기준\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 1. 결함 탐지 (Detection)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "탐지 기준 진폭:\n"
                "┌──────────────────┬────────────────────────────────┐\n"
                "│ 진폭 수준         │ 조치                            │\n"
                "├──────────────────┼────────────────────────────────┤\n"
                "│ DAC 20% 미만      │ 기록 불요 (무시)                │\n"
                "│ DAC 20% 이상      │ 기록 필수 (기록 수준)           │\n"
                "│ DAC 50% 이상      │ 크기 측정 및 평가 수준          │\n"
                "│ DAC 100% 초과     │ 불합격 추정, 크기 측정 필수     │\n"
                "└──────────────────┴────────────────────────────────┘\n"
                "※ 검사 감도: 평가 수준(DAC 50%)보다 +6 dB 추가 증폭 후 스캔\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 2. 결함 위치 확인 (Location)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "① 위치 좌표 기록 항목:\n"
                "   - 스캔 축(Scan Axis): 용접선 기준 거리 (mm)\n"
                "   - 인덱스 축(Index Axis): 탐촉자 이동 방향 거리 (mm)\n"
                "   - 깊이(Depth): 검사면 기준 결함 상단·하단 깊이 (mm)\n\n"
                "② 위치 정확도 요건:\n"
                "   - 인덱스 포인트 정확도: ±1 mm 이내 (ISO 13588 / ASME Art.4)\n"
                "   - 깊이 정확도: ±1 mm 이내\n\n"
                "③ S-scan / Linear scan 활용:\n"
                "   - S-scan: 결함 경사 및 방향 확인에 유리\n"
                "   - Linear scan: 정확한 스캔 축 위치 확인에 유리\n"
                "   - B-scan / D-scan 복합 활용으로 3차원 위치 파악\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 3. 결함 크기 측정 (Sizing)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "【방법 1】 −6 dB 강하법 (Half Maximum Amplitude Method)\n"
                "- 최대 에코 진폭의 50% (−6 dB) 지점을 결함 끝단으로 정의\n"
                "- 결함 길이(ℓ) 및 높이(a) 측정에 모두 사용\n"
                "- 가장 널리 사용되는 표준 방법\n"
                "- 적용 기준: ASME Art.4, B31.1, B31.3, ISO 19285\n\n"
                "【방법 2】 −20 dB 강하법 (Endpoint Method)\n"
                "- 최대 에코 진폭의 10% (−20 dB) 지점을 결함 끝단으로 정의\n"
                "- 결함 길이 측정에 주로 사용\n"
                "- −6 dB법보다 더 보수적인(큰) 크기 산출\n\n"
                "【방법 3】 TOFD (Time of Flight Diffraction)\n"
                "- 결함 선단 회절파 도달 시간차로 높이(a) 측정\n"
                "- 결함 높이 정확도: ±0.5~1 mm (최고 정밀도)\n"
                "- 표면 근접 결함 탐지 취약 (Lateral wave 간섭)\n"
                "- Code Case 2235, Sec. VIII Div.2 적용 시 TOFD 병행 권장\n\n"
                "【방법 비교표】\n"
                "┌──────────────┬──────────────┬────────────┬────────────────┐\n"
                "│ 방법          │ 측정 항목     │ 정확도     │ 주 사용 코드   │\n"
                "├──────────────┼──────────────┼────────────┼────────────────┤\n"
                "│ −6 dB 강하법 │ 길이, 높이    │ ±1~2 mm    │ ASME, ISO 범용 │\n"
                "│ −20 dB 강하법│ 길이          │ ±2~3 mm    │ 보수적 평가    │\n"
                "│ TOFD          │ 높이 (우선)   │ ±0.5~1 mm  │ CC2235, Div.2  │\n"
                "└──────────────┴──────────────┴────────────┴────────────────┘\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 4. 결함 분류 (Classification)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "평면형 결함 (Planar) — 가장 위험\n"
                "  · 균열 (Crack), 용합 불량 (Lack of Fusion), 용입 불량 (Lack of Penetration)\n"
                "  · 모든 코드에서 크기 무관 불합격 처리\n\n"
                "체적형 결함 (Volumetric) — 크기 기준 평가\n"
                "  · 기공 (Porosity), 슬래그 개재물 (Slag), 텅스텐 개재물 (Tungsten)\n"
                "  · 개별 크기 및 누적 길이 기준 적용\n\n"
                "표면 연결 결함 (Surface-Breaking)\n"
                "  · 표면 균열, 루트 불용합 등\n"
                "  · 내부 결함보다 강화된 기준 적용\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 5. 기록 요건\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "기록 필수 항목:\n"
                "  - A-scan 파형 + S-scan/B-scan 데이터\n"
                "  - 결함 위치 (스캔 축, 인덱스 축, 깊이)\n"
                "  - 결함 크기 (길이 ℓ, 높이 a)\n"
                "  - 최대 에코 진폭 (%DAC 또는 dB)\n"
                "  - 사용된 크기 측정 방법\n"
                "  - 합부 판정 결과 및 근거 조항",

            "기타": "사용자 정의 텍스트를 입력하세요."
        }
        
        # 상단 저장 영역 (최우선 표시)
        top_frame = tk.Frame(root, bg="#e8f4fd", relief=tk.RIDGE, bd=1)
        top_frame.pack(fill=tk.X, padx=10, pady=(10, 2))
        
        tk.Label(top_frame, text="문서 제목:", font=("Arial", 10, "bold"), bg="#e8f4fd").pack(side=tk.LEFT, padx=(10, 5), pady=6)
        self.new_title_entry = tk.Entry(top_frame, width=50, font=("Arial", 10))
        self.new_title_entry.pack(side=tk.LEFT, padx=5, pady=6)
        self.new_title_entry.insert(0, "비파괴 검사 절차서")
        
        tk.Button(top_frame, text="💾  Word 문서 생성 / 저장", command=self.generate_document,
                  bg="#2196F3", fg="white", padx=15, pady=5,
                  font=("Arial", 10, "bold"), relief=tk.RAISED).pack(side=tk.LEFT, padx=15, pady=6)
        tk.Button(top_frame, text="📄 초안 저장", command=self.save_draft,
                  bg="#607d8b", fg="white", padx=10, pady=5,
                  font=("Arial", 9)).pack(side=tk.LEFT, padx=3, pady=6)
        tk.Button(top_frame, text="📂 초안 불러오기", command=self.load_draft,
                  bg="#607d8b", fg="white", padx=10, pady=5,
                  font=("Arial", 9)).pack(side=tk.LEFT, padx=3, pady=6)
        self._redo_btn = tk.Button(top_frame, text="↪ 다시실행", command=self.redo,
                  bg="#f5f5f5", padx=10, pady=5, font=("Arial", 9), state=tk.DISABLED)
        self._redo_btn.pack(side=tk.RIGHT, padx=3, pady=6)
        self._undo_btn = tk.Button(top_frame, text="↩ 실행취소", command=self.undo,
                  bg="#f5f5f5", padx=10, pady=5, font=("Arial", 9), state=tk.DISABLED)
        self._undo_btn.pack(side=tk.RIGHT, padx=3, pady=6)
        
        # 하단 버튼 영역
        button_frame = tk.Frame(root)
        button_frame.pack(fill=tk.X, padx=10, pady=(2, 10))
        
        tk.Button(button_frame, text="Word 파일 로드", command=self.load_document, bg="lightblue", padx=10, pady=5).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="사진 추가", command=self.add_images, bg="lightgreen", padx=10, pady=5).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="표준 추가", command=self.add_standard, bg="orange", padx=10, pady=5).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="✏ 텍스트 추가", command=self.add_text_item, bg="#e8f5e9", padx=10, pady=5).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="텍스트 편집", command=self.edit_selected_text, bg="#d6eaff", padx=10, pady=5).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="↑ 위로", command=self.move_item_up, bg="#f5f5f5", padx=8, pady=5).pack(side=tk.LEFT, padx=2)
        tk.Button(button_frame, text="↓ 아래로", command=self.move_item_down, bg="#f5f5f5", padx=8, pady=5).pack(side=tk.LEFT, padx=2)
        tk.Button(button_frame, text="📋 복사", command=self.duplicate_selected_item, bg="#fff9c4", padx=8, pady=5).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="선택 삭제", command=self.delete_selected_item, bg="#ffb3b3", padx=10, pady=5).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="전체 초기화", command=self.clear_all, bg="lightcoral", padx=10, pady=5).pack(side=tk.LEFT, padx=5)
        
        self.status_label = tk.Label(root, text="로드된 문서: 없음 | 포함된 사진: 0개", bg="lightyellow", pady=5)
        self.status_label.pack(fill=tk.X, padx=10)
        
        # 가운데 내용 영역
        main_frame = tk.Frame(root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self._paned = tk.PanedWindow(main_frame, orient=tk.HORIZONTAL, sashrelief=tk.RAISED,
                               sashwidth=6, bg="#cccccc")
        self._paned.pack(fill=tk.BOTH, expand=True)

        left_frame = tk.Frame(self._paned)
        self._paned.add(left_frame, stretch='always', minsize=400)

        right_frame = tk.Frame(self._paned, bg="white", relief=tk.FLAT)
        self._paned.add(right_frame, stretch='always', minsize=280)

        # sash 드래그 후 위치 저장
        self._paned.bind('<ButtonRelease-1>', lambda e: self._save_sash_position())

        # 창 표시 후 초기 비율 설정 (저장된 위치 복원 또는 좌 70% : 우 30%)
        def _set_sash(event=None):
            try:
                total = self._paned.winfo_width()
                if total > 100:
                    saved = self._load_sash_position()
                    pos = saved if saved and saved < total - 280 else int(total * 0.70)
                    self._paned.sash_place(0, pos, 0)
                    root.unbind('<Map>', _map_id[0])
            except Exception:
                pass
        _map_id = [None]
        _map_id[0] = root.bind('<Map>', _set_sash)
        
        # 트리뷰 검색창
        search_row = tk.Frame(left_frame)
        search_row.pack(fill=tk.X, pady=(0, 3))
        tk.Label(search_row, text="🔍", font=("Arial", 10)).pack(side=tk.LEFT, padx=(0, 3))
        self._tree_filter_var.trace_add('write', lambda *_: self._apply_tree_filter())
        ttk.Entry(search_row, textvariable=self._tree_filter_var,
                  font=("Arial", 9)).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(search_row, text="✕", width=3,
                   command=lambda: self._tree_filter_var.set('')).pack(side=tk.LEFT, padx=2)

        # 트리뷰
        tree_frame = tk.Frame(left_frame)
        tree_frame.pack(fill=tk.BOTH, expand=True)
        
        tree_scrollbar = tk.Scrollbar(tree_frame)
        tree_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.tree = ttk.Treeview(tree_frame, yscrollcommand=tree_scrollbar.set, height=35)
        self.tree.pack(fill=tk.BOTH, expand=True)
        tree_scrollbar.config(command=self.tree.yview)
        
        self.tree['columns'] = ('타입', '내용')
        self.tree.column('#0', width=50, anchor='w')
        self.tree.column('타입', width=80, anchor='w')
        self.tree.column('내용', width=950, anchor='w')
        self.tree.heading('#0', text='번호')
        self.tree.heading('타입', text='타입')
        self.tree.heading('내용', text='내용')
        
        self.tree.bind('<Double-1>', self.on_tree_double_click)
        self.tree.bind('<Button-1>', self.on_tree_click)
        self.tree.bind('<Button-3>', self.on_tree_right_click)  # 우클릭 바인드
        self.tree.bind('<Alt-Up>', lambda e: self.move_item_up())
        self.tree.bind('<Alt-Down>', lambda e: self.move_item_down())
        self.tree.bind('<Delete>', lambda e: self.delete_selected_item())
        self.tree.bind('<F2>', lambda e: self.edit_selected_text())
        self.tree.bind('<Return>', lambda e: self.edit_selected_text())
        
        # 우측 내용 및 이미지 영역
        info_label = tk.Label(right_frame, text="로드된 문서 내용", font=("Arial", 11, "bold"))
        info_label.pack(anchor='w')
        
        content_frame = tk.Frame(right_frame)
        content_frame.pack(fill=tk.BOTH, expand=True)
        
        content_scrollbar = tk.Scrollbar(content_frame)
        content_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.content_text = tk.Text(content_frame, wrap=tk.WORD,
                                     font=("Malgun Gothic", 9),
                                     yscrollcommand=content_scrollbar.set)
        self.content_text.pack(fill=tk.BOTH, expand=True)
        content_scrollbar.config(command=self.content_text.yview)
        self.content_text.config(state=tk.DISABLED)
        self.content_text.bind('<Button-1>', self._on_content_text_click)
        self.content_text.bind('<Motion>', self._on_content_text_motion)
        
        # 이미지 미리보기 영역
        image_preview_label = tk.Label(right_frame, text="이미지 미리보기", font=("Arial", 11, "bold"))
        image_preview_label.pack(anchor='w', pady=(10, 0))
        
        self.image_canvas = tk.Canvas(right_frame, height=160)
        self.image_canvas.pack(fill=tk.X, pady=5)
        self.image_frame = tk.Frame(self.image_canvas)
        self.image_canvas.create_window((0, 0), window=self.image_frame, anchor='nw')
        
        self.image_scrollbar = tk.Scrollbar(right_frame, orient=tk.HORIZONTAL, command=self.image_canvas.xview)
        self.image_scrollbar.pack(fill=tk.X)
        self.image_canvas.configure(xscrollcommand=self.image_scrollbar.set)
        self.image_frame.bind('<Configure>', lambda e: self.image_canvas.configure(scrollregion=self.image_canvas.bbox('all')))
        self.image_thumbnails = []
        self.content_images = []
        
        self.update_info_text()
    
    def on_tree_double_click(self, event):
        item = self.tree.identify_row(event.y)
        if not item:
            selected = self.tree.selection()
            if not selected:
                return
            item = selected[0]
        else:
            self.tree.selection_set(item)
            self.tree.focus(item)

        values = self.tree.item(item, 'values')
        if not values:
            return
        content_index = int(item)
        if values[0] == 'image':
            if 0 <= content_index < len(self.content):
                self.replace_image_dialog(content_index)
        elif values[0] == 'table':
            if 0 <= content_index < len(self.content):
                self.edit_table_dialog(content_index)
        else:
            self.edit_selected_text(item)

    def on_tree_click(self, event):
        item = self.tree.identify_row(event.y)
        if not item:
            selected = self.tree.selection()
            if not selected:
                return
            item = selected[0]
        else:
            self.tree.selection_set(item)
            self.tree.focus(item)

        try:
            self._highlight_content_text_item(int(item))
        except Exception:
            pass

        # 단일 클릭은 선택·하이라이트만 수행 (편집/교체는 더블클릭)

    def _highlight_content_text_item(self, idx):
        """트리 선택 항목을 content_text에서 하이라이트 및 스크롤"""
        tag = f"item_{idx}"
        try:
            ranges = self.content_text.tag_ranges(tag)
            if ranges:
                self.content_text.tag_remove("_hl_sel", "1.0", tk.END)
                self.content_text.tag_add("_hl_sel", ranges[0], ranges[1])
                # 선택 색상: 타입별 기본색보다 진하게 + 텍스트 화이트
                itype = self.content[idx].get('type', 'text') if idx < len(self.content) else 'text'
                hl_colors = {'text': '#1565c0', 'image': '#2e7d32', 'table': '#e65100'}
                hl_bg =     {'text': '#bbdefb', 'image': '#c8e6c9', 'table': '#ffe0b2'}
                self.content_text.tag_config("_hl_sel",
                    background=hl_bg.get(itype, '#fff3cd'),
                    foreground=hl_colors.get(itype, '#333'),
                    font=("Malgun Gothic", 9, "bold"))
                self.content_text.see(ranges[0])
        except Exception:
            pass

    def _on_content_text_click(self, event):
        """요약 텍스트 클릭 → 트리뷰 해당 항목 선택 및 스크롤"""
        try:
            pos = self.content_text.index(f"@{event.x},{event.y}")
            tags = self.content_text.tag_names(pos)
            for tag in tags:
                if tag.startswith("item_"):
                    item_idx = int(tag[5:])
                    iid = str(item_idx)
                    if self.tree.exists(iid):
                        self._highlight_content_text_item(item_idx)
                        self.tree.selection_set(iid)
                        self.tree.focus(iid)
                        self.tree.see(iid)
                    break
        except Exception:
            pass

    def _on_content_text_motion(self, event):
        """마우스 오버 시 커서 hand2 전환"""
        try:
            pos = self.content_text.index(f"@{event.x},{event.y}")
            tags = self.content_text.tag_names(pos)
            has_item = any(t.startswith("item_") for t in tags)
            self.content_text.config(cursor="hand2" if has_item else "")
        except Exception:
            pass

    def get_insert_index(self, mode='after'):
        selected = self.tree.selection()
        if selected:
            base_idx = int(selected[0])
            idx = base_idx if mode == 'before' else base_idx + 1
            return min(max(idx, 0), len(self.content))
        return len(self.content)

    def edit_selected_text(self, item_id=None):
        item = item_id
        if item is None:
            selected = self.tree.selection()
            if not selected:
                messagebox.showwarning("선택 필요", "편집할 텍스트 항목을 선택하세요.")
                return
            item = selected[0]

        values = self.tree.item(item, 'values')
        if not values or values[0] != 'text':
            messagebox.showwarning("선택 오류", "텍스트 항목만 편집할 수 있습니다.")
            return

        content_index = int(item)
        if 0 <= content_index < len(self.content):
            current_text = self.content[content_index].get('text', '')
            self.edit_text_dialog(content_index, current_text)
    
    def edit_text_dialog(self, index, current_text):
        STYLES = ['Normal', 'Heading 1', 'Heading 2', 'Heading 3',
                  'Heading 4', 'List Bullet', 'List Number', 'Body Text']
        dialog = tk.Toplevel(self.root)
        dialog.title("텍스트 편집")
        dialog.geometry("680x480")
        dialog.bind('<Control-Return>', lambda e: save_text())
        dialog.bind('<Escape>', lambda e: dialog.destroy())

        def save_text():
            self._push_undo()
            new_text = text_area.get("1.0", tk.END).rstrip("\n")
            new_style = style_var.get()
            if 0 <= index < len(self.content) and self.content[index].get('type') == 'text':
                self.content[index]['text'] = new_text
                self.content[index]['style'] = new_style
                self._renumber_sections()  # 삽입/편집 후 섹션 번호 재정렬
                self.refresh_content()
                iid = str(index)
                if self.tree.exists(iid):
                    self.tree.selection_set(iid)
                    self.tree.focus(iid)
                    self.tree.see(iid)
                self._highlight_content_text_item(index)
                dialog.destroy()

        # 상단 툴바 (저장 버튼)
        toolbar = tk.Frame(dialog, bg="#e8f4fd", relief=tk.RIDGE, bd=1)
        toolbar.pack(fill=tk.X, padx=0, pady=(0, 3))

        tk.Button(toolbar, text="💾  저장 (Ctrl+Enter)", command=save_text,
                  bg="#2196F3", fg="white", padx=15, pady=4,
                  font=("Arial", 10, "bold"), relief=tk.RAISED).pack(side=tk.LEFT, padx=10, pady=5)
        tk.Button(toolbar, text="✕  닫기 (Esc)", command=dialog.destroy,
                  bg="#f44336", fg="white", padx=12, pady=4,
                  font=("Arial", 10), relief=tk.RAISED).pack(side=tk.LEFT, padx=5, pady=5)
        tk.Label(toolbar, text=f"항목 #{index + 1} 편집 중",
                 font=("Arial", 9), bg="#e8f4fd", fg="#555").pack(side=tk.LEFT, padx=10)

        # 스타일 선택 행
        style_row = tk.Frame(dialog)
        style_row.pack(fill=tk.X, padx=10, pady=(0, 4))
        tk.Label(style_row, text="단락 스타일:", font=("Arial", 9)).pack(side=tk.LEFT, padx=(0, 5))
        current_style = self.content[index].get('style', 'Normal') if 0 <= index < len(self.content) else 'Normal'
        if current_style not in STYLES:
            STYLES.append(current_style)
        style_var = tk.StringVar(value=current_style)
        ttk.Combobox(style_row, textvariable=style_var, values=STYLES,
                     state='readonly', width=18, font=("Arial", 9)).pack(side=tk.LEFT)
        tk.Label(style_row, text="  ※ Word 저장 시 해당 스타일로 적용됩니다",
                 font=("Arial", 8), fg="#888").pack(side=tk.LEFT, padx=8)

        text_area = tk.Text(dialog, wrap=tk.WORD, font=("Malgun Gothic", 10))
        text_area.pack(fill=tk.BOTH, expand=True, padx=10, pady=(0, 10))
        text_area.insert(tk.END, current_text)
        text_area.focus_set()
    
    def replace_image_dialog(self, index):
        file_path = filedialog.askopenfilename(filetypes=[("Image files", "*.jpg *.jpeg *.png *.gif *.bmp")])
        if file_path:
            if not (0 <= index < len(self.content)) or self.content[index].get('type') != 'image':
                return
            self._push_undo()
            old_path = self.content[index]['path']
            # 새 이미지 복사
            import shutil
            raw_dir = os.path.dirname(old_path) if old_path else ''
            temp_dir = raw_dir if raw_dir and os.path.isdir(raw_dir) else os.path.dirname(file_path)
            new_path = os.path.join(temp_dir, f"replaced_{os.path.basename(file_path)}")
            shutil.copy(file_path, new_path)
            self.content[index]['path'] = new_path
            self.update_tree_view()
            self.update_image_preview()
            self.update_info_text()
    
    def load_document(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if not file_path:
            return
        
        try:
            self._push_undo()
            self.content = load_existing_doc(file_path)
            self.source_file = file_path  # 바닥글/헤더 이미지 보존을 위해 원본 경로 저장
            self._redo_stack.clear()
            self._update_undo_buttons()
            self.refresh_content()
            self.status_label.config(text=f"로드된 문서: {os.path.basename(file_path)} | 포함된 사진: {len(self.image_paths)}개")
            messagebox.showinfo("완료", f"문서를 로드했습니다.\n내용: {len(self.content)}개\n사진: {len(self.image_paths)}개")
        except Exception as e:
            messagebox.showerror("오류", f"문서 로드 실패: {str(e)}")
    
    def update_tree_view(self):
        # detach된 항목은 get_children()에 포함되지 않으므로
        # 삽입한 iid 목록을 직접 추적해서 삭제
        if hasattr(self, '_all_tree_iids'):
            for iid in self._all_tree_iids:
                try:
                    self.tree.delete(iid)
                except Exception:
                    pass
        self._all_tree_iids = set()
        self._detached_iids = set()
        visible_num = 1  # 트리에 실제로 표시되는 순번 (삭제 후 연속 번호 유지)
        for idx, item in enumerate(self.content):
            item_type = item.get('type')
            area = item.get('area', 'body')
            area_label = '' if area == 'body' else f"[{ '머릿글' if area == 'header' else '바닥글' }] "
            if item_type == 'text':
                text = item.get('text', '').strip().replace('\n', ' ')[:120]
                style = item.get('style', 'Normal')
                if text:
                    self.tree.insert('', tk.END, iid=str(idx), text=str(visible_num), values=('text', f"{area_label}{style}: {text}"))
                    self._all_tree_iids.add(str(idx))
                    visible_num += 1
            elif item_type == 'image':
                path = item.get('path', '')
                w = item.get('width_inches', 5.0)
                self.tree.insert('', tk.END, iid=str(idx), text=str(visible_num), values=('image', f"{area_label}{os.path.basename(path)}  [{w}인치]"))
                self._all_tree_iids.add(str(idx))
                visible_num += 1
            elif item_type == 'table':
                self.tree.insert('', tk.END, iid=str(idx), text=str(visible_num), values=('table', f"{area_label}표"))
                self._all_tree_iids.add(str(idx))
                visible_num += 1
        self._apply_tree_filter()
    
    def update_info_text(self):
        self.content_text.config(state=tk.NORMAL)
        self.content_text.delete("1.0", tk.END)
        self.content_images.clear()
        info_content = f"""
로드된 내용 요약:
- 총 요소 수: {len(self.content)}개
- 텍스트 문단: {len(self.paragraphs)}개
- 포함된 사진: {len(self.image_paths)}개
"""
        self.content_text.insert(tk.END, info_content)
        self.content_text.insert(tk.END, "\n---\n\n")
        
        visible_num = 1  # 실제 표시 순번 (연속 번호)
        item_positions = {}  # {idx: (start_pos, end_pos)} - 태그는 루프 후 일괄 적용

        for idx, item in enumerate(self.content):
            area = item.get('area', 'body')
            area_label = '' if area == 'body' else f"[{ '머릿글' if area == 'header' else '바닥글' }]\n"
            start_pos = self.content_text.index(tk.END)
            if item['type'] == 'text':
                raw_text = item.get('text', '')
                style = item.get('style', 'Normal')
                # strip()으로 공백만 있는 항목 제외 (tree_view와 동일한 기준)
                if raw_text.strip():
                    self.content_text.insert(tk.END, f"■ {visible_num}번 — [{style}]  {area_label}\n{raw_text}\n\n")
                    visible_num += 1
            elif item['type'] == 'image':
                path = item['path']
                self.content_text.insert(tk.END, f"🖼 {visible_num}번 — [이미지]  {area_label}{os.path.basename(path)}\n")
                visible_num += 1
                if os.path.exists(path):
                    try:
                        img = Image.open(path)
                        img.thumbnail((280, 280))
                        photo = ImageTk.PhotoImage(img)
                        self.content_images.append(photo)
                        self.content_text.image_create(tk.END, image=photo)
                        self.content_text.insert(tk.END, "\n\n")
                    except Exception as e:
                        self.content_text.insert(tk.END, f"[이미지 로드 실패: {e}]\n\n")
                else:
                    self.content_text.insert(tk.END, "[파일 없음]\n\n")
            elif item['type'] == 'table':
                table_data = item['data']
                self.content_text.insert(tk.END, f"🗒 {visible_num}번 — [표]  {area_label}{len(table_data)}행 × {len(table_data[0]) if table_data else 0}열\n")
                table_text = "\n".join("  " + "  |  ".join(str(c) for c in row) for row in table_data) + "\n\n"
                self.content_text.insert(tk.END, table_text)
                visible_num += 1
            # tk.END는 항상 "마지막+1줄" 위치라 다음 항목과 경계가 겹침
            # → 실제 삽입된 마지막 문자 바로 다음 위치를 end_pos로 정확히 기록
            raw_end = self.content_text.index(tk.END)
            # END는 "\n"이 자동 추가된 위치이므로 -1c 로 실제 내용 끝을 구함
            end_pos = self.content_text.index(f"{raw_end} -1c")
            if self.content_text.compare(start_pos, '<', end_pos):
                item_positions[idx] = (start_pos, end_pos)

        # ─── 모든 삽입 완료 후 일괄 태그 적용 ─────────────────────
        # 삽입이 끝난 뒤 tag_add → right-gravity 영향 없음 → 정확한 범위 보장
        for idx, (start, end) in item_positions.items():
            self.content_text.tag_add(f"item_{idx}", start, end)

        self.content_text.config(state=tk.DISABLED)
    
    def update_image_preview(self):
        for widget in self.image_frame.winfo_children():
            widget.destroy()
        self.image_thumbnails.clear()

        if not self.image_paths:
            placeholder = tk.Label(self.image_frame, text="로드된 이미지가 없습니다.", bg="white", width=80, height=10)
            placeholder.pack(padx=10, pady=10)
            return

        for idx, path in enumerate(self.image_paths, 1):
            if not os.path.exists(path):
                continue
            try:
                img = Image.open(path)
                img.thumbnail((180, 180))
                photo = ImageTk.PhotoImage(img)
                self.image_thumbnails.append(photo)
                frame = tk.Frame(self.image_frame, bd=1, relief=tk.RIDGE)
                frame.pack(side=tk.LEFT, padx=5, pady=5)
                label = tk.Label(frame, image=photo)
                label.pack()
                caption = tk.Label(frame, text=os.path.basename(path), wraplength=180)
                caption.pack()
            except Exception as e:
                print(f"이미지 미리보기 오류: {e}")
    
    def _renumber_sections(self):
        """텍스트 내 섹션 번호 자동 재정렬 — 삭제·삽입·이동 모두 대응
        N.M / N.M.P / N.M.P.Q 등 임의 깊이 섹션 번호 처리
        예) 4.8.4.2 삭제 → 4.8.4.3 → 4.8.4.2
            3.1 삭제 → 3.2→3.1, 3.3→3.2
            7.0 삭제 → 8.0→7.0, 8.1→7.1 (캐스케이딩)
        """
        import re

        # 숫자.숫자 이상의 섹션 번호 (최소 점 하나 이상)
        # 뒤에 반드시 공백 또는 줄끝 — "3.14mm" 같은 소수 오매칭 방지
        section_re = re.compile(r'^(\s*)(\d+(?:\.\d+)+)(?=\s|$)([\s\S]*)$', re.DOTALL)

        prefix_counter = {}  # {new_parent_str: 현재 카운터}
        old_to_new = {}      # {old_num_str: new_num_str}  캐스케이딩용
        changed = False

        for item in self.content:
            if item.get('type') != 'text':
                continue
            text = item.get('text', '')
            m = section_re.match(text)
            if not m:
                continue

            leading      = m.group(1)
            old_num_str  = m.group(2)   # 예: "4.8.4.3"
            rest         = m.group(3)   # 예: " Acceptance Standards"

            parts = old_num_str.split('.')
            if len(parts) < 2:
                continue

            last_val       = int(parts[-1])
            old_parent_str = '.'.join(parts[:-1])   # 예: "4.8.4"

            # 부모가 이미 재정렬됐으면 새 번호로 매핑
            new_parent_str = old_to_new.get(old_parent_str, old_parent_str)

            if last_val == 0:
                # N.0 / N.M.0 형식 — 상위 섹션 헤더
                # 조부모 기준으로 카운터 증가
                parent_parts = new_parent_str.split('.')
                if len(parent_parts) == 1:
                    gp = ''  # 최상위 (루트)
                else:
                    gp_raw = '.'.join(parent_parts[:-1])
                    gp = old_to_new.get(gp_raw, gp_raw)

                prefix_counter[gp] = prefix_counter.get(gp, 0) + 1
                new_n = prefix_counter[gp]
                new_parent_actual = f"{gp}.{new_n}" if gp else str(new_n)
                new_num_str = f"{new_parent_actual}.0"

                # 부모 번호 자체를 매핑 → 하위 항목 캐스케이딩 지원
                old_to_new[old_parent_str] = new_parent_actual
                old_to_new[old_num_str]    = new_num_str
            else:
                # 일반 하위 항목 (N.M, N.M.P, N.M.P.Q …)
                prefix_counter[new_parent_str] = prefix_counter.get(new_parent_str, 0) + 1
                new_last    = prefix_counter[new_parent_str]
                new_num_str = f"{new_parent_str}.{new_last}"
                old_to_new[old_num_str] = new_num_str

            if new_num_str != old_num_str:
                item['text'] = f'{leading}{new_num_str}{rest}'
                changed = True

        return changed

    def refresh_content(self):
        self.paragraphs = [item for item in self.content if item.get('type') == 'text']
        self.image_paths = [item.get('path', '') for item in self.content if item.get('type') == 'image']
        self.update_tree_view()
        self.update_image_preview()
        self.update_info_text()
    
    def add_standard(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("표준 절차 추가")
        dialog.geometry("820x620")
        dialog.resizable(True, True)

        # ── 상단 정보 ──
        top = tk.Frame(dialog, bg="#e8f4fd", pady=6)
        top.pack(fill=tk.X)

        selected = self.tree.selection()
        if selected:
            selected_idx = int(selected[0])
            selected_values = self.tree.item(selected[0], 'values')
            selected_desc = selected_values[1] if selected_values and len(selected_values) > 1 else '선택 항목'
            sel_txt = f"현재 선택: {selected_idx + 1}번 ({selected_desc[:55]})"
        else:
            sel_txt = "현재 선택: 없음 (맨 끝에 추가됨)"

        tk.Label(top, text=sel_txt, fg="#1565c0", bg="#e8f4fd",
                 font=("Arial", 9)).pack(side=tk.LEFT, padx=12)

        pos_frame = tk.Frame(top, bg="#e8f4fd")
        pos_frame.pack(side=tk.RIGHT, padx=12)
        tk.Label(pos_frame, text="삽입 위치:", bg="#e8f4fd").pack(side=tk.LEFT)
        position_var = tk.StringVar(value='after')
        tk.Radiobutton(pos_frame, text="위", variable=position_var, value='before',
                       bg="#e8f4fd").pack(side=tk.LEFT, padx=3)
        tk.Radiobutton(pos_frame, text="아래", variable=position_var, value='after',
                       bg="#e8f4fd").pack(side=tk.LEFT)

        standards = self.standards

        # ── 탭 카테고리 정의 ──
        TAB_DEFS = [
            ("전체",  None,                                              "#ffffff"),
            ("PAUT",  ["PAUT","Phased Array","13588","11666","19285","Code Case"], "#fffde7"),
            ("RT",    ["RT","Radiograph","17636","10675","디지털 RT"],   "#e3f2fd"),
            ("MT",    ["MT","Magnetic","17638","23278"],                  "#e8f5e9"),
            ("PT",    ["PT","Penetrant","3452","23277"],                  "#fce4ec"),
            ("PMI",   ["PMI","Material Identification","15011",
                       "E1476","E2191","PCC-2"],                          "#fffde7"),
            ("UT",    ["UT","Ultrasonic","TOFD"],                         "#e8eaf6"),
        ]

        def get_keys(keywords):
            if keywords is None:
                return list(standards.keys())
            return [k for k in standards.keys()
                    if any(w.lower() in k.lower() for w in keywords)]

        # ── Notebook ──
        notebook = ttk.Notebook(dialog)
        notebook.pack(fill=tk.BOTH, expand=True, padx=8, pady=(6, 0))

        tab_state = {}   # tab_name → {'var', 'listbox', 'preview', 'keys'}

        def build_tab(tab_name, keywords, bg):
            frame = tk.Frame(notebook, bg=bg)
            notebook.add(frame, text=f"  {tab_name}  ")

            keys = get_keys(keywords)

            # 좌: 리스트박스  /  우: 미리보기
            pane = tk.PanedWindow(frame, orient=tk.HORIZONTAL, sashrelief=tk.RIDGE)
            pane.pack(fill=tk.BOTH, expand=True, padx=6, pady=6)

            # 좌측 패널
            left = tk.Frame(pane, bg=bg)
            pane.add(left, width=280)

            tk.Label(left, text=f"{len(keys)}개 항목", font=("Arial", 8),
                     fg="gray", bg=bg).pack(anchor='w', padx=4)

            sb = tk.Scrollbar(left)
            sb.pack(side=tk.RIGHT, fill=tk.Y)
            lb = tk.Listbox(left, yscrollcommand=sb.set, selectmode=tk.SINGLE,
                            font=("Arial", 9), activestyle='dotbox', relief=tk.FLAT,
                            selectbackground="#1565c0", selectforeground="white")
            lb.pack(fill=tk.BOTH, expand=True, padx=2)
            sb.config(command=lb.yview)
            for k in keys:
                lb.insert(tk.END, k)

            # 우측 미리보기
            right = tk.Frame(pane, bg=bg)
            pane.add(right)

            tk.Label(right, text="미리보기", font=("Arial", 8), fg="gray",
                     bg=bg).pack(anchor='w', padx=4)
            preview_sb = tk.Scrollbar(right)
            preview_sb.pack(side=tk.RIGHT, fill=tk.Y)
            preview = tk.Text(right, wrap=tk.WORD, yscrollcommand=preview_sb.set,
                              font=("Arial", 9), bg="#fafafa", relief=tk.FLAT,
                              state=tk.DISABLED)
            preview.pack(fill=tk.BOTH, expand=True, padx=2)
            preview_sb.config(command=preview.yview)

            def on_select(evt):
                sel = lb.curselection()
                if not sel:
                    return
                key = lb.get(sel[0])
                preview.config(state=tk.NORMAL)
                preview.delete("1.0", tk.END)
                preview.insert(tk.END, standards.get(key, ""))
                preview.config(state=tk.DISABLED)

            lb.bind('<<ListboxSelect>>', on_select)
            if keys:
                lb.selection_set(0)
                lb.event_generate('<<ListboxSelect>>')

            tab_state[tab_name] = {'lb': lb, 'preview': preview, 'keys': keys}

        for name, kws, bg in TAB_DEFS:
            build_tab(name, kws, bg)

        # ── 하단 버튼 ──
        btn_bar = tk.Frame(dialog, bg="#f5f5f5", pady=6)
        btn_bar.pack(fill=tk.X)

        def get_selected_key():
            tab_name = notebook.tab(notebook.select(), "text").strip()
            data = tab_state.get(tab_name)
            if not data:
                return None
            sel = data['lb'].curselection()
            if not sel:
                return None
            return data['lb'].get(sel[0])

        def do_add():
            key = get_selected_key()
            if not key:
                messagebox.showwarning("선택 없음", "추가할 항목을 선택하세요.", parent=dialog)
                return
            self._push_undo()
            insert_index = self.get_insert_index(position_var.get())
            self.content.insert(insert_index, {'type': 'text', 'text': standards[key],
                                               'style': 'Normal', 'area': 'body'})
            self.refresh_content()
            iid = str(insert_index)
            if self.tree.exists(iid):
                self.tree.selection_set(iid)
                self.tree.focus(iid)
                self.tree.see(iid)
            self._highlight_content_text_item(insert_index)
            self.status_label.config(text=f"로드된 문서: 변경됨 | 포함된 사진: {len(self.image_paths)}개")
            dialog.destroy()
            self.edit_text_dialog(insert_index, standards[key])

        def do_delete():
            key = get_selected_key()
            if not key:
                return
            if not messagebox.askyesno("삭제 확인",
                    f"'{key}' 항목을 목록에서 삭제하시겠습니까?", parent=dialog):
                return
            del standards[key]
            # 모든 탭 리스트 갱신
            for name, kws, _ in TAB_DEFS:
                data = tab_state.get(name)
                if not data:
                    continue
                new_keys = get_keys(kws)
                data['keys'] = new_keys
                data['lb'].delete(0, tk.END)
                for k in new_keys:
                    data['lb'].insert(tk.END, k)
                if new_keys:
                    data['lb'].selection_set(0)
                    data['lb'].event_generate('<<ListboxSelect>>')

        tk.Button(btn_bar, text="✔  추가", command=do_add,
                  bg="#4CAF50", fg="white", padx=16, pady=5,
                  font=("Arial", 10, "bold")).pack(side=tk.LEFT, padx=10)
        tk.Button(btn_bar, text="🗑  삭제", command=do_delete,
                  bg="#ffb3b3", fg="red", padx=12, pady=5,
                  font=("Arial", 9)).pack(side=tk.LEFT, padx=4)
        tk.Button(btn_bar, text="📱  HTML 모바일 내보내기",
                  command=lambda: (dialog.destroy(), self.export_standards_html()),
                  bg="#6c5ce7", fg="white", padx=12, pady=5,
                  font=("Arial", 9)).pack(side=tk.LEFT, padx=4)
        tk.Button(btn_bar, text="닫기", command=dialog.destroy,
                  padx=12, pady=5).pack(side=tk.RIGHT, padx=10)

    def export_standards_html(self):
        """앱 표준절차 프레임과 동일한 구조의 모바일 HTML 내보내기"""
        import datetime, json as _json
        out_path = filedialog.asksaveasfilename(
            defaultextension=".html",
            filetypes=[("HTML 파일", "*.html")],
            initialfile="NDT_Standards_Mobile.html",
            title="HTML 내보내기 위치 선택"
        )
        if not out_path:
            return

        # ── 탭 정의 (앱 add_standard와 동일) ──
        TAB_DEFS = [
            ("전체",  None,
             "#1a237e", "🗂"),
            ("PAUT",  ["PAUT","Phased Array","13588","11666","19285","Code Case"],
             "#6c5ce7", "🔷"),
            ("RT",    ["RT","Radiograph","17636","10675","디지털 RT"],
             "#0984e3", "☢️"),
            ("MT",    ["MT","Magnetic","17638","23278"],
             "#00b894", "🔩"),
            ("PT",    ["PT","Penetrant","3452","23277"],
             "#e17055", "🔴"),
            ("PMI",   ["PMI","Material Identification","15011","E1476","E2191","PCC-2"],
             "#f39c12", "🟡"),
            ("UT",    ["UT","Ultrasonic","TOFD"],
             "#3498db", "🔵"),
        ]

        def get_keys(keywords):
            if keywords is None:
                return list(self.standards.keys())
            return [k for k in self.standards.keys()
                    if any(w.lower() in k.lower() for w in keywords)]

        def safe(s):
            return (s.replace("&","&amp;")
                     .replace("<","&lt;")
                     .replace(">","&gt;")
                     .replace('"','&quot;'))

        # ── 데이터 직렬화 (JS에 주입) ──
        data_dict = {k: v for k, v in self.standards.items()}
        tabs_data = []
        for name, kws, color, icon in TAB_DEFS:
            keys = get_keys(kws)
            tabs_data.append({"name": name, "color": color, "icon": icon, "keys": keys})

        js_data   = _json.dumps(data_dict,   ensure_ascii=False)
        js_tabs   = _json.dumps(tabs_data,   ensure_ascii=False)
        now       = datetime.datetime.now().strftime('%Y-%m-%d %H:%M')
        total     = len(self.standards)

        html = f"""<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>NDT 표준 절차 참조집</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
html,body{{height:100%;overflow:hidden}}
body{{font-family:'Noto Sans KR',system-ui,sans-serif;background:#f0f2f5;
      color:#2d3436;display:flex;flex-direction:column}}

/* ── 헤더 ── */
header{{background:linear-gradient(135deg,#1a237e,#283593);color:#fff;
        padding:10px 14px;flex-shrink:0;
        box-shadow:0 2px 6px rgba(0,0,0,.35)}}
header h1{{font-size:1rem;letter-spacing:.5px}}
header p{{font-size:.68rem;opacity:.75;margin-top:2px}}

/* ── 탭바 ── */
.tab-bar{{display:flex;background:#1e3c72;overflow-x:auto;
          flex-shrink:0;scrollbar-width:none}}
.tab-bar::-webkit-scrollbar{{display:none}}
.tab-btn{{flex:none;padding:9px 14px;font-size:.78rem;
          color:rgba(255,255,255,.65);border:none;background:none;
          cursor:pointer;white-space:nowrap;
          border-bottom:3px solid transparent;transition:all .2s}}
.tab-btn.active{{color:#fff;border-bottom-color:#74b9ff;font-weight:700}}

/* ── 메인 레이아웃 ── */
.main{{display:flex;flex:1;overflow:hidden}}

/* ── 좌측 리스트 ── */
.sidebar{{width:300px;flex-shrink:0;display:flex;flex-direction:column;
          background:#fff;border-right:1px solid #dfe6e9}}
.search-wrap{{padding:8px 8px 4px;border-bottom:1px solid #eee}}
.search-wrap input{{width:100%;padding:7px 10px;font-size:.82rem;
                    border:1px solid #dfe6e9;border-radius:6px;outline:none}}
.search-wrap input:focus{{border-color:#6c5ce7}}
.list-count{{font-size:.7rem;color:#aaa;padding:4px 10px 2px}}
.item-list{{flex:1;overflow-y:auto;padding:4px 0}}
.item-btn{{width:100%;text-align:left;padding:9px 14px;font-size:.82rem;
           border:none;background:none;cursor:pointer;
           border-left:4px solid transparent;
           border-bottom:1px solid #f1f2f6;
           color:#2d3436;line-height:1.4;transition:all .15s}}
.item-btn:hover{{background:#f8f9fa}}
.item-btn.active{{background:#eff0ff;border-left-color:var(--tc);
                  color:#1a237e;font-weight:600}}
.item-btn.hidden{{display:none}}

/* ── 우측 미리보기 ── */
.preview-pane{{flex:1;display:flex;flex-direction:column;overflow:hidden}}
.preview-title{{padding:10px 16px;font-size:.9rem;font-weight:700;
                background:#f8f9fa;border-bottom:2px solid var(--tc,#1a237e);
                color:var(--tc,#1a237e);flex-shrink:0}}
.preview-body{{flex:1;overflow-y:auto;padding:14px 18px;
               font-size:.82rem;line-height:1.85;white-space:pre-wrap;
               color:#2d3436}}
.preview-empty{{display:flex;align-items:center;justify-content:center;
                height:100%;color:#b2bec3;font-size:.9rem;flex-direction:column;gap:8px}}

/* ── 모바일 (≤640px): 세로 전환 ── */
@media(max-width:640px){{
  html,body{{overflow:auto}}
  .main{{flex-direction:column;overflow:visible}}
  .sidebar{{width:100%;border-right:none;border-bottom:1px solid #dfe6e9;
            max-height:240px}}
  .preview-pane{{min-height:300px}}
  .preview-body{{padding:10px 12px}}
}}
</style>
</head>
<body>
<header>
  <h1>📋 NDT 표준 절차 참조집</h1>
  <p>생성일: {now} &nbsp;|&nbsp; 총 {total}개 항목</p>
</header>

<div class="tab-bar" id="tabBar"></div>

<div class="main">
  <div class="sidebar">
    <div class="search-wrap">
      <input id="searchBox" type="text" placeholder="🔍 항목 검색..."
             oninput="filterList(this.value)">
    </div>
    <div class="list-count" id="listCount"></div>
    <div class="item-list" id="itemList"></div>
  </div>
  <div class="preview-pane" id="previewPane">
    <div class="preview-empty" id="previewEmpty">
      <span style="font-size:2rem">📄</span>
      <span>좌측 목록에서 항목을 선택하세요</span>
    </div>
    <div class="preview-title" id="previewTitle" style="display:none"></div>
    <div class="preview-body"  id="previewBody"  style="display:none"></div>
  </div>
</div>

<script>
const DATA = {js_data};
const TABS = {js_tabs};

let curTab  = 0;
let curKey  = null;

/* ── 탭 생성 ── */
const tabBar = document.getElementById('tabBar');
TABS.forEach((t, i) => {{
  const btn = document.createElement('button');
  btn.className = 'tab-btn' + (i===0?' active':'');
  btn.textContent = t.icon + ' ' + t.name;
  btn.onclick = () => switchTab(i);
  tabBar.appendChild(btn);
}});

function switchTab(i) {{
  curTab = i;
  curKey = null;
  document.querySelectorAll('.tab-btn').forEach((b,j)=>b.classList.toggle('active',j===i));
  document.getElementById('searchBox').value = '';
  buildList('');
  showEmpty();
}}

/* ── 리스트 렌더링 ── */
function buildList(q) {{
  const keys   = TABS[curTab].keys;
  const color  = TABS[curTab].color;
  const list   = document.getElementById('itemList');
  const count  = document.getElementById('listCount');
  list.innerHTML = '';
  let shown = 0;
  keys.forEach(k => {{
    const match = !q || k.toLowerCase().includes(q) ||
                  (DATA[k]||'').toLowerCase().includes(q);
    const btn = document.createElement('button');
    btn.className = 'item-btn' + (k===curKey?' active':'') + (match?'':' hidden');
    btn.style.setProperty('--tc', color);
    btn.textContent = k;
    btn.onclick = () => selectItem(k, color);
    list.appendChild(btn);
    if(match) shown++;
  }});
  count.textContent = shown + '개 항목';
}}

/* ── 항목 선택 ── */
function selectItem(key, color) {{
  curKey = key;
  document.querySelectorAll('.item-btn').forEach(b => {{
    b.classList.toggle('active', b.textContent === key);
  }});
  const title = document.getElementById('previewTitle');
  const body  = document.getElementById('previewBody');
  const empty = document.getElementById('previewEmpty');
  title.textContent = key;
  title.style.setProperty('--tc', color);
  body.textContent  = DATA[key] || '';
  title.style.display = '';
  body.style.display  = '';
  empty.style.display = 'none';
  body.scrollTop = 0;
}}

function showEmpty() {{
  document.getElementById('previewTitle').style.display = 'none';
  document.getElementById('previewBody').style.display  = 'none';
  document.getElementById('previewEmpty').style.display = '';
}}

function filterList(q) {{
  buildList(q.toLowerCase());
}}

/* ── 초기화 ── */
buildList('');
</script>
</body>
</html>"""

        with open(out_path, 'w', encoding='utf-8') as f:
            f.write(html)

        if messagebox.askyesno("완료", f"HTML 파일 저장 완료:\n{out_path}\n\n지금 브라우저로 열어볼까요?"):
            import webbrowser
            webbrowser.open(out_path)

    def insert_standard(self, text_area, selected):
        if selected and selected in self.standards:
            if selected == "기타":
                custom_text = simpledialog.askstring("입력", "텍스트를 입력하세요:")
                if custom_text:
                    text_area.insert(tk.INSERT, custom_text)
            else:
                text_area.insert(tk.INSERT, self.standards[selected])
    
    def edit_table_dialog(self, index):
        if not (0 <= index < len(self.content)) or self.content[index].get('type') != 'table':
            return

        data = [list(row) for row in self.content[index]['data']]  # 깊은 복사

        dialog = tk.Toplevel(self.root)
        dialog.title(f"표 편집  (항목 #{index + 1})")
        dialog.geometry("860x540")
        dialog.resizable(True, True)

        # 상단 툴바
        toolbar = tk.Frame(dialog, bg="#e8f4fd", relief=tk.RIDGE, bd=1)
        toolbar.pack(fill=tk.X)

        entry_widgets = []

        def get_current_data():
            return [[e.get() for e in row_ents] for row_ents in entry_widgets]

        def save_table():
            self._push_undo()
            self.content[index]['data'] = get_current_data()
            self.refresh_content()
            self.status_label.config(text=f"로드된 문서: 변경됨 | 포함된 사진: {len(self.image_paths)}개")
            dialog.destroy()

        tk.Button(toolbar, text="💾  저장 (Ctrl+Enter)", command=save_table,
                  bg="#2196F3", fg="white", padx=15, pady=4,
                  font=("Arial", 10, "bold")).pack(side=tk.LEFT, padx=10, pady=5)
        tk.Button(toolbar, text="✕  닫기", command=dialog.destroy,
                  bg="#f44336", fg="white", padx=12, pady=4,
                  font=("Arial", 10)).pack(side=tk.LEFT, padx=5, pady=5)
        tk.Button(toolbar, text="＋ 행 추가", command=lambda: add_row(),
                  bg="#4caf50", fg="white", padx=8, pady=4,
                  font=("Arial", 9)).pack(side=tk.LEFT, padx=5, pady=5)
        tk.Button(toolbar, text="－ 행 삭제", command=lambda: del_row(),
                  bg="#ff9800", fg="white", padx=8, pady=4,
                  font=("Arial", 9)).pack(side=tk.LEFT, padx=5, pady=5)
        tk.Button(toolbar, text="＋ 열 추가", command=lambda: add_col(),
                  bg="#9c27b0", fg="white", padx=8, pady=4,
                  font=("Arial", 9)).pack(side=tk.LEFT, padx=5, pady=5)
        tk.Button(toolbar, text="－ 열 삭제", command=lambda: del_col(),
                  bg="#795548", fg="white", padx=8, pady=4,
                  font=("Arial", 9)).pack(side=tk.LEFT, padx=5, pady=5)
        tk.Label(toolbar, text="※ 셀 클릭 후 직접 입력",
                 font=("Arial", 8), bg="#e8f4fd", fg="#666").pack(side=tk.LEFT, padx=10)

        # 스크롤 가능 영역
        canvas_outer = tk.Frame(dialog)
        canvas_outer.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        canvas = tk.Canvas(canvas_outer, bg="white")
        vbar = tk.Scrollbar(canvas_outer, orient=tk.VERTICAL, command=canvas.yview)
        hbar = tk.Scrollbar(canvas_outer, orient=tk.HORIZONTAL, command=canvas.xview)
        vbar.pack(side=tk.RIGHT, fill=tk.Y)
        hbar.pack(side=tk.BOTTOM, fill=tk.X)
        canvas.pack(fill=tk.BOTH, expand=True)
        canvas.configure(yscrollcommand=vbar.set, xscrollcommand=hbar.set)

        grid_frame = tk.Frame(canvas, bg="white")
        canvas.create_window((0, 0), window=grid_frame, anchor='nw')
        grid_frame.bind('<Configure>', lambda e: canvas.configure(scrollregion=canvas.bbox('all')))

        def build_grid():
            for w in grid_frame.winfo_children():
                w.destroy()
            entry_widgets.clear()
            ncols = len(data[0]) if data else 0
            # 열 헤더
            tk.Label(grid_frame, text="", bg="#bbdefb", relief=tk.RIDGE,
                     width=4).grid(row=0, column=0, padx=1, pady=1, sticky='nsew')
            for c in range(ncols):
                tk.Label(grid_frame, text=f"열 {c+1}", font=("Arial", 8, "bold"),
                         bg="#bbdefb", relief=tk.RIDGE, width=18).grid(
                         row=0, column=c+1, padx=1, pady=1, sticky='nsew')
            # 데이터 행
            for r, row in enumerate(data):
                tk.Label(grid_frame, text=str(r+1), font=("Arial", 8),
                         bg="#e3f2fd", relief=tk.RIDGE, width=4).grid(
                         row=r+1, column=0, padx=1, pady=1, sticky='nsew')
                row_entries = []
                for c, cell in enumerate(row):
                    ent = tk.Entry(grid_frame, font=("Malgun Gothic", 9),
                                   width=20, relief=tk.SOLID, bd=1)
                    ent.insert(0, str(cell) if cell is not None else '')
                    ent.grid(row=r+1, column=c+1, padx=1, pady=1, sticky='nsew')
                    row_entries.append(ent)
                entry_widgets.append(row_entries)

        def add_row():
            nonlocal data
            data = get_current_data()
            ncols = len(data[0]) if data else 1
            data.append([''] * ncols)
            build_grid()

        def del_row():
            nonlocal data
            data = get_current_data()
            if len(data) > 1:
                data.pop()
                build_grid()

        def add_col():
            nonlocal data
            data = get_current_data()
            for row in data:
                row.append('')
            build_grid()

        def del_col():
            nonlocal data
            data = get_current_data()
            if data and len(data[0]) > 1:
                for row in data:
                    row.pop()
                build_grid()

        build_grid()
        dialog.bind('<Control-Return>', lambda e: save_table())
        dialog.bind('<Escape>', lambda e: dialog.destroy())
    
    def add_images(self):
        file_paths = filedialog.askopenfilenames(filetypes=[("Image files", "*.jpg *.jpeg *.png *.gif *.bmp")])
        if file_paths:
            width_str = simpledialog.askstring("이미지 삽입 너비",
                f"{len(file_paths)}개 사진의 문서 삽입 너비 (인치, 기본 5.0):",
                initialvalue="5.0", parent=self.root)
            try:
                width = float(width_str) if width_str else 5.0
                width = max(0.5, min(width, 12.0))
            except (TypeError, ValueError):
                width = 5.0
            self._push_undo()
            insert_index = self.get_insert_index()
            for path in file_paths:
                self.content.insert(insert_index, {
                    'type': 'image', 'path': path,
                    'area': 'body', 'width_inches': width
                })
                insert_index += 1
            self.refresh_content()
            first_new_index = insert_index - len(file_paths)
            first_iid = str(first_new_index)
            if self.tree.exists(first_iid):
                self.tree.selection_set(first_iid)
                self.tree.focus(first_iid)
                self.tree.see(first_iid)
            self.status_label.config(text=f"로드된 문서: 변경됨 | 포함된 사진: {len(self.image_paths)}개")
            messagebox.showinfo("완료", f"{len(file_paths)}개의 사진이 {first_new_index + 1}번째 위치부터 추가되었습니다. (너비: {width}인치)")

    def delete_selected_item(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("선택 필요", "삭제할 항목을 먼저 선택하세요.")
            return

        content_index = int(selected[0])
        if not (0 <= content_index < len(self.content)):
            messagebox.showerror("오류", "유효하지 않은 선택입니다.")
            return

        item = self.content[content_index]
        item_type = item.get('type', 'unknown')

        if not messagebox.askyesno("확인", f"선택한 {item_type} 항목을 삭제하시겠습니까?"):
            return

        self._push_undo()
        del self.content[content_index]
        self._renumber_sections()  # 섹션 번호 자동 재정렬 (N.0 삭제 시 N+1.0 → N.0)
        self.refresh_content()
        self.status_label.config(text=f"로드된 문서: 변경됨 | 포함된 사진: {len(self.image_paths)}개")
    
    def on_tree_right_click(self, event):
        """트리 항목 우클릭 메뉴"""
        item = self.tree.identify_row(event.y)
        if not item:
            return
        
        # 선택 항목 설정
        self.tree.selection_set(item)
        self.tree.focus(item)
        
        try:
            content_index = int(item)
        except (ValueError, IndexError):
            return
        
        if not (0 <= content_index < len(self.content)):
            return
        
        # 컨텍스트 메뉴 생성
        context_menu = tk.Menu(self.root, tearoff=False)
        
        item_type = self.content[content_index].get('type', 'unknown')
        
        if item_type == 'text':
            context_menu.add_command(
                label="편집  (F2)",
                command=lambda: self.edit_selected_text(item)
            )
        elif item_type == 'image':
            context_menu.add_command(
                label="이미지 교체",
                command=lambda: self.replace_image_dialog(content_index)
            )
            context_menu.add_command(
                label="📐 삽입 너비 변경",
                command=lambda: self._change_image_width(content_index)
            )
        elif item_type == 'table':
            context_menu.add_command(
                label="✏ 표 편집",
                command=lambda: self.edit_table_dialog(content_index)
            )
        context_menu.add_command(
            label="📋 복사 (아래에 추가)",
            command=self.duplicate_selected_item
        )
        context_menu.add_separator()
        context_menu.add_command(
            label="↑ 위로 이동  (Alt+↑)",
            command=self.move_item_up
        )
        context_menu.add_command(
            label="↓ 아래로 이동  (Alt+↓)",
            command=self.move_item_down
        )
        context_menu.add_separator()
        context_menu.add_command(
            label="삭제  (Delete)",
            command=self.delete_selected_item,
            foreground="red"
        )
        
        # 메뉴 표시
        try:
            context_menu.tk_popup(event.x_root, event.y_root)
        finally:
            context_menu.grab_release()
    
    def add_text_item(self):
        """새 빈 텍스트 항목을 선택 위치 다음에 추가하고 편집 창 열기"""
        self._push_undo()
        insert_index = self.get_insert_index()
        self.content.insert(insert_index, {'type': 'text', 'text': '', 'style': 'Normal', 'area': 'body'})
        self.refresh_content()
        iid = str(insert_index)
        if self.tree.exists(iid):
            self.tree.selection_set(iid)
            self.tree.focus(iid)
            self.tree.see(iid)
        self.edit_text_dialog(insert_index, '')
        self.status_label.config(text=f"로드된 문서: 변경됨 | 포함된 사진: {len(self.image_paths)}개")

    def move_item_up(self):
        """선택 항목을 한 칸 위로 이동 (단축키: Alt+↑)"""
        selected = self.tree.selection()
        if not selected:
            return
        idx = int(selected[0])
        if idx <= 0:
            return
        self._push_undo()
        self.content[idx], self.content[idx - 1] = self.content[idx - 1], self.content[idx]
        self.refresh_content()
        new_iid = str(idx - 1)
        if self.tree.exists(new_iid):
            self.tree.selection_set(new_iid)
            self.tree.focus(new_iid)
            self.tree.see(new_iid)

    def move_item_down(self):
        """선택 항목을 한 칸 아래로 이동 (단축키: Alt+↓)"""
        selected = self.tree.selection()
        if not selected:
            return
        idx = int(selected[0])
        if idx >= len(self.content) - 1:
            return
        self._push_undo()
        self.content[idx], self.content[idx + 1] = self.content[idx + 1], self.content[idx]
        self.refresh_content()
        new_iid = str(idx + 1)
        if self.tree.exists(new_iid):
            self.tree.selection_set(new_iid)
            self.tree.focus(new_iid)
            self.tree.see(new_iid)

    def duplicate_selected_item(self):
        """선택 항목을 바로 아래에 복사"""
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("선택 필요", "복사할 항목을 먼저 선택하세요.")
            return
        idx = int(selected[0])
        if not (0 <= idx < len(self.content)):
            return
        self._push_undo()
        new_item = copy.deepcopy(self.content[idx])
        self.content.insert(idx + 1, new_item)
        self.refresh_content()
        new_iid = str(idx + 1)
        if self.tree.exists(new_iid):
            self.tree.selection_set(new_iid)
            self.tree.focus(new_iid)
            self.tree.see(new_iid)
        self.status_label.config(text=f"로드된 문서: 변경됨 | 포함된 사진: {len(self.image_paths)}개")

    def clear_all(self):
        if messagebox.askyesno("확인", "모든 내용을 초기화하시겠습니까?"):
            self.content = []
            self.source_file = None
            self._undo_stack.clear()
            self._redo_stack.clear()
            self._update_undo_buttons()
            self.refresh_content()
            self.status_label.config(text="로드된 문서: 없음 | 포함된 사진: 0개")
    
    # ── Undo / Redo ──────────────────────────────────────────
    def _push_undo(self):
        """현재 content 상태를 undo 스택에 저장 (최대 30단계)"""
        self._undo_stack.append((copy.deepcopy(self.content), self.source_file))
        if len(self._undo_stack) > 30:
            self._undo_stack.pop(0)
        self._redo_stack.clear()
        self._update_undo_buttons()

    def _update_undo_buttons(self):
        """Undo/Redo 버튼 활성·비활성 상태 동기화"""
        try:
            self._undo_btn.config(state=tk.NORMAL if self._undo_stack else tk.DISABLED)
            self._redo_btn.config(state=tk.NORMAL if self._redo_stack else tk.DISABLED)
        except Exception:
            pass

    def undo(self):
        """이전 상태로 되돌리기 (Ctrl+Z)"""
        if not self._undo_stack:
            return
        self._redo_stack.append((copy.deepcopy(self.content), self.source_file))
        self.content, self.source_file = self._undo_stack.pop()
        self._update_undo_buttons()
        self.refresh_content()
        self.status_label.config(text=f"↩ 실행 취소  | 항목 수: {len(self.content)}개")

    def redo(self):
        """되돌리기 취소 (Ctrl+Y)"""
        if not self._redo_stack:
            return
        self._undo_stack.append((copy.deepcopy(self.content), self.source_file))
        self.content, self.source_file = self._redo_stack.pop()
        self._update_undo_buttons()
        self.refresh_content()
        self.status_label.config(text=f"↪ 다시 실행  | 항목 수: {len(self.content)}개")

    # ── 초안 저장 / 불러오기 ─────────────────────────────────
    def _git_version_save(self, file_path: str) -> str:
        """
        file_path가 속한 폴더가 git 저장소이면
        해당 파일을 stage하고 자동 커밋한다.
        커밋 메시지에는 파일명 + 저장 시각이 포함된다.
        반환값: 상태 메시지 문자열 (UI 표시용)
        """
        try:
            file_path = os.path.abspath(file_path)
            folder = os.path.dirname(file_path)
            fname = os.path.basename(file_path)

            # git 실행 가능 여부 확인
            try:
                subprocess.run(
                    ['git', '--version'],
                    capture_output=True, check=True
                )
            except (FileNotFoundError, subprocess.CalledProcessError):
                return ''  # git 없으면 조용히 종료

            # 해당 경로가 git 저장소 안에 있는지 확인
            result = subprocess.run(
                ['git', 'rev-parse', '--show-toplevel'],
                cwd=folder,
                capture_output=True, text=True
            )
            if result.returncode != 0:
                return ''  # git 저장소 아님

            repo_root = result.stdout.strip()

            # git add (파일의 git 저장소 내 상대 경로)
            rel_path = os.path.relpath(file_path, repo_root)
            subprocess.run(
                ['git', 'add', rel_path],
                cwd=repo_root,
                capture_output=True, check=True
            )

            # 변경사항이 있는지 확인 (stage된 변경)
            status = subprocess.run(
                ['git', 'diff', '--cached', '--name-only'],
                cwd=repo_root,
                capture_output=True, text=True
            )
            if not status.stdout.strip():
                return '\n[Git] 변경 없음 (커밋 생략)'

            # 커밋 메시지: 파일명 + 저장 시각
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            commit_msg = f'Auto-save: {fname} [{timestamp}]'
            commit_result = subprocess.run(
                ['git', 'commit', '-m', commit_msg],
                cwd=repo_root,
                capture_output=True, text=True
            )
            if commit_result.returncode == 0:
                # 커밋 해시 앞 7자리
                hash_result = subprocess.run(
                    ['git', 'rev-parse', '--short', 'HEAD'],
                    cwd=repo_root,
                    capture_output=True, text=True
                )
                short_hash = hash_result.stdout.strip()

                # 원격 저장소(origin) 존재 시 자동 push
                remote_check = subprocess.run(
                    ['git', 'remote'],
                    cwd=repo_root,
                    capture_output=True, text=True
                )
                push_msg = ''
                if 'origin' in remote_check.stdout:
                    push_result = subprocess.run(
                        ['git', 'push', 'origin', 'HEAD'],
                        cwd=repo_root,
                        capture_output=True, text=True
                    )
                    if push_result.returncode == 0:
                        push_msg = ' → GitHub push 완료'
                    else:
                        push_msg = f' → push 실패: {push_result.stderr.strip()[:60]}'

                return f'\n[Git] 커밋 완료 ({short_hash}) — {timestamp}{push_msg}'
            else:
                return f'\n[Git] 커밋 실패: {commit_result.stderr.strip()[:80]}'

        except Exception as e:
            return f'\n[Git] 오류: {str(e)[:80]}'

    def save_draft(self):
        """작업 중인 content를 JSON 초안으로 저장 (이미지는 경로만 저장)"""
        path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("NDT 초안 파일", "*.json"), ("모든 파일", "*.*")],
            initialfile="ndt_draft.json",
            title="초안 저장"
        )
        if not path:
            return
        serializable = []
        for item in self.content:
            entry = {k: v for k, v in item.items() if k != 'element'}
            serializable.append(entry)
        draft = {
            'source_file': self.source_file or '',
            'title': self.new_title_entry.get(),
            'content': serializable
        }
        try:
            with open(path, 'w', encoding='utf-8') as f:
                json.dump(draft, f, ensure_ascii=False, indent=2)
            git_msg = self._git_version_save(path)
            self.status_label.config(text=f"초안 저장 완료: {os.path.basename(path)}{git_msg}")
        except Exception as e:
            messagebox.showerror("오류", f"초안 저장 실패:\n{e}")

    def load_draft(self):
        """저장된 JSON 초안 불러오기"""
        path = filedialog.askopenfilename(
            filetypes=[("NDT 초안 파일", "*.json"), ("모든 파일", "*.*")],
            title="초안 불러오기"
        )
        if not path:
            return
        try:
            with open(path, 'r', encoding='utf-8') as f:
                draft = json.load(f)
            self._push_undo()
            self.content = draft.get('content', [])
            sf = draft.get('source_file', '')
            self.source_file = sf if sf and os.path.exists(sf) else None
            title = draft.get('title', '')
            if title:
                self.new_title_entry.delete(0, tk.END)
                self.new_title_entry.insert(0, title)
            self.refresh_content()
            self.status_label.config(text=f"초안 불러오기 완료: {os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("오류", f"초안 불러오기 실패:\n{e}")

    # ── 트리뷰 검색 필터 ─────────────────────────────────────
    def _apply_tree_filter(self):
        """검색어에 맞지 않는 트리 항목을 detach로 숨김"""
        try:
            query = self._tree_filter_var.get().lower().strip()
        except Exception:
            return

        # detach된 항목을 별도 집합으로 추적해서 확실하게 reattach
        if not hasattr(self, '_detached_iids'):
            self._detached_iids = set()

        for iid in list(self._detached_iids):
            try:
                self.tree.reattach(iid, '', 'end')
            except Exception:
                pass
        self._detached_iids = set()

        # reattach 후 번호 순서로 정렬 (detach→reattach 시 순서 뒤섞임 방지)
        all_iids = self.tree.get_children('')
        for idx, iid in enumerate(sorted(all_iids, key=lambda x: int(x))):
            self.tree.move(iid, '', idx)

        if not query:
            return

        for iid in list(self.tree.get_children('')):
            try:
                values = self.tree.item(iid, 'values')
                text = ' '.join(str(v) for v in values).lower()
                if query not in text:
                    self.tree.detach(iid)
                    self._detached_iids.add(iid)
            except Exception:
                pass

    # ── 이미지 너비 변경 ─────────────────────────────────────
    def _change_image_width(self, index):
        """선택 이미지 항목의 문서 삽입 너비 변경"""
        if not (0 <= index < len(self.content)):
            return
        current_w = self.content[index].get('width_inches', 5.0)
        new_w_str = simpledialog.askstring(
            "이미지 너비 변경",
            f"새 너비 (인치, 현재: {current_w}인치):",
            initialvalue=str(current_w),
            parent=self.root
        )
        try:
            new_w = float(new_w_str) if new_w_str else current_w
            new_w = max(0.5, min(new_w, 12.0))
        except (TypeError, ValueError):
            return
        self._push_undo()
        self.content[index]['width_inches'] = new_w
        self.refresh_content()
        iid = str(index)
        if self.tree.exists(iid):
            self.tree.selection_set(iid)
            self.tree.focus(iid)
            self.tree.see(iid)

    def generate_document(self):
        if not self.content:
            messagebox.showerror("오류", "로드된 내용이 없습니다.")
            return
        title = self.new_title_entry.get().strip() or "비파괴 검사 절차서"

        # ── 원본 파일을 템플릿으로 사용 (헤더/바닥글 이미지 완전 보존) ──
        if self.source_file and os.path.exists(self.source_file):
            doc = Document(self.source_file)
            # 본문 내용만 삭제 (sectPr·헤더·바닥글 참조는 유지)
            body = doc.element.body
            sectPr = body.find(qn('w:sectPr'))
            for el in list(body):
                if el is not sectPr:
                    body.remove(el)
            # sectPr가 없으면 다시 추가
            if sectPr is not None and body.find(qn('w:sectPr')) is None:
                body.append(sectPr)
        else:
            doc = Document()
            # 헤더 구성 (원본 없을 때)
            try:
                header = doc.sections[0].header
                clear_story_container(header)
                for item in self.content:
                    if item.get('area') != 'header':
                        continue
                    if item['type'] == 'text':
                        header.add_paragraph(item.get('text', ''))
                    elif item['type'] == 'table':
                        add_bordered_table(header, item.get('data', []))
            except:
                pass
            # 바닥글 구성 (원본 없을 때 - 이미지 포함)
            try:
                footer = doc.sections[0].footer
                clear_story_container(footer)
                for item in self.content:
                    if item.get('area') != 'footer':
                        continue
                    if item['type'] == 'text':
                        footer.add_paragraph(item.get('text', ''))
                    elif item['type'] == 'image':
                        path = item.get('path', '')
                        if os.path.exists(path):
                            try:
                                p = footer.add_paragraph()
                                p.add_run().add_picture(path, width=Inches(5))
                            except:
                                pass
                    elif item['type'] == 'table':
                        add_bordered_table(footer, item.get('data', []))
            except:
                pass

        # ── 본문 내용 작성 ──
        inserted_elements = set()  # 동일 element 중복 삽입 방지

        for item in self.content:
            if item.get('area') in ('header', 'footer'):
                continue
            elif item['type'] == 'text':
                text = item.get('text', '').strip()
                if not text:
                    continue
                style = item.get('style', 'Normal')
                if 'Heading' in style:
                    try:
                        level = int(style.split()[-1])
                    except:
                        level = 1
                    doc.add_heading(text, level=level)
                else:
                    doc.add_paragraph(text)
            elif item['type'] == 'image':
                img_elem = item.get('element')
                if self.source_file and img_elem is not None:
                    elem_id = id(img_elem)
                    if elem_id not in inserted_elements:
                        inserted_elements.add(elem_id)
                        # 원본 단락 element deepcopy → 이미지 관계(rId) 완전 보존
                        para_copy = copy.deepcopy(img_elem)
                        _body = doc.element.body
                        _sect = _body.find(qn('w:sectPr'))
                        if _sect is not None:
                            _body.insert(list(_body).index(_sect), para_copy)
                        else:
                            _body.append(para_copy)
                else:
                    path = item.get('path', '')
                    if os.path.exists(path):
                        try:
                            w = item.get('width_inches', 5.0)
                            doc.add_picture(path, width=Inches(w))
                            doc.add_paragraph()
                        except:
                            pass
            elif item['type'] == 'table':
                data = item.get('data', [])
                tbl_elem = item.get('element')
                if tbl_elem is not None:
                    # 원본 XML 깊은 복사 → 병합 셀, 서식, 테두리 완전 보존
                    tbl_copy = copy.deepcopy(tbl_elem)
                    body = doc.element.body
                    sect_pr = body.find(qn('w:sectPr'))
                    if sect_pr is not None:
                        body.insert(list(body).index(sect_pr), tbl_copy)
                    else:
                        body.append(tbl_copy)
                    doc.add_paragraph()
                elif data:
                    add_bordered_table(doc, data)
                    doc.add_paragraph()

        output_file = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if output_file:
            doc.save(output_file)
            git_msg = self._git_version_save(output_file)
            messagebox.showinfo("완료", f"문서가 저장되었습니다:\n{output_file}{git_msg}")
    
    def load_window_geometry(self):
        """저장된 창 크기 로드"""
        try:
            if os.path.exists(self.CONFIG_FILE):
                with open(self.CONFIG_FILE, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    geometry = config.get('window_geometry', '1400x850')
                    return geometry
        except:
            pass
        return '1400x850'

    def _load_sash_position(self):
        """저장된 PanedWindow sash 위치 로드 (없으면 None)"""
        try:
            if os.path.exists(self.CONFIG_FILE):
                with open(self.CONFIG_FILE, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    return config.get('sash_position', None)
        except:
            pass
        return None

    def _save_sash_position(self):
        """현재 PanedWindow sash 위치 저장"""
        try:
            pos = self._paned.sash_coord(0)[0]
            config = {}
            if os.path.exists(self.CONFIG_FILE):
                with open(self.CONFIG_FILE, 'r', encoding='utf-8') as f:
                    config = json.load(f)
            config['sash_position'] = pos
            with open(self.CONFIG_FILE, 'w', encoding='utf-8') as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
        except:
            pass
    
    def save_window_geometry(self):
        """현재 창 크기 저장"""
        try:
            geometry = self.root.geometry()
            config = {}
            if os.path.exists(self.CONFIG_FILE):
                with open(self.CONFIG_FILE, 'r', encoding='utf-8') as f:
                    config = json.load(f)
            config['window_geometry'] = geometry
            with open(self.CONFIG_FILE, 'w', encoding='utf-8') as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
        except:
            pass
    
    def on_window_configure(self, event):
        """창 크기 변경 시 저장"""
        if event.widget == self.root:
            self.save_window_geometry()

if __name__ == "__main__":
    try:
        root = tk.Tk()
        app = NDTProcedureApp(root)
        root.mainloop()
    except KeyboardInterrupt:
        pass
