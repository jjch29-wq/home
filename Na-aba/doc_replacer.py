"""
doc_replacer.py - Word 절차서 Find & Replace 도구
플레이스홀더({{태그}}) 또는 일반 텍스트를 찾아 일괄 교체합니다.
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import json
import re
from docx import Document
from docx.oxml.ns import qn
import datetime

PRESETS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "replacer_presets.json")


# ─────────────────────────────────────────────────────────────
# 핵심 교체 로직
# ─────────────────────────────────────────────────────────────

def _replace_in_paragraph(paragraph, find_text, replace_text):
    """
    run이 분할되어 있어도 전체 텍스트에서 찾아 교체.
    서식(볼드/이탤릭/크기)은 첫 번째 run 기준으로 유지.
    """
    full_text = "".join(r.text for r in paragraph.runs)
    if find_text not in full_text:
        return False

    new_text = full_text.replace(find_text, replace_text)

    # 첫 번째 run에 전체 텍스트 넣고 나머지 비움
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    return True


def _replace_in_table(table, find_text, replace_text):
    count = 0
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if _replace_in_paragraph(para, find_text, replace_text):
                    count += 1
    return count


def replace_in_doc(doc, replacements):
    """
    replacements: [(find, replace), ...]
    반환: {find: 교체 횟수}
    """
    results = {f: 0 for f, _ in replacements}

    for find_text, replace_text in replacements:
        if not find_text:
            continue

        # 본문 단락
        for para in doc.paragraphs:
            if _replace_in_paragraph(para, find_text, replace_text):
                results[find_text] += 1

        # 본문 표
        for table in doc.tables:
            results[find_text] += _replace_in_table(table, find_text, replace_text)

        # 머릿글 / 바닥글
        for section in doc.sections:
            for container in [section.header, section.footer,
                               section.even_page_header, section.even_page_footer,
                               section.first_page_header, section.first_page_footer]:
                try:
                    for para in container.paragraphs:
                        if _replace_in_paragraph(para, find_text, replace_text):
                            results[find_text] += 1
                    for table in container.tables:
                        results[find_text] += _replace_in_table(table, find_text, replace_text)
                except Exception:
                    pass

    return results


# ─────────────────────────────────────────────────────────────
# GUI
# ─────────────────────────────────────────────────────────────

class DocReplacerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("📄 절차서 Find & Replace")
        self.root.geometry("820x680")
        self.root.configure(bg="#f3f4f6")

        self.file_paths = []          # 로드된 파일 목록
        self.rows = []                # (find_var, replace_var, frame) 리스트
        self.presets = {}             # 저장된 프리셋

        self._load_presets()
        self._build_ui()

    # ── UI 빌드 ──────────────────────────────────────────────
    def _build_ui(self):
        # 상단: 파일 선택
        file_frame = ttk.LabelFrame(self.root, text=" 📂 대상 파일 (여러 개 가능) ", padding=10)
        file_frame.pack(fill="x", padx=15, pady=(12, 5))

        btn_row = tk.Frame(file_frame, bg="#f3f4f6")
        btn_row.pack(fill="x")
        ttk.Button(btn_row, text="파일 추가", command=self._add_files).pack(side="left", padx=3)
        ttk.Button(btn_row, text="폴더 추가", command=self._add_folder).pack(side="left", padx=3)
        ttk.Button(btn_row, text="목록 비우기", command=self._clear_files).pack(side="right", padx=3)

        list_frame = tk.Frame(file_frame, bg="white")
        list_frame.pack(fill="x", pady=(6, 0))
        self.file_listbox = tk.Listbox(list_frame, height=4, font=("Consolas", 9),
                                       selectmode="extended", borderwidth=0)
        self.file_listbox.pack(side="left", fill="both", expand=True)
        sb = ttk.Scrollbar(list_frame, orient="vertical", command=self.file_listbox.yview)
        sb.pack(side="right", fill="y")
        self.file_listbox.config(yscrollcommand=sb.set)

        # 중단: 교체 규칙
        rule_outer = ttk.LabelFrame(self.root, text=" 🔄 교체 규칙 ", padding=10)
        rule_outer.pack(fill="both", expand=True, padx=15, pady=5)

        # 헤더 라벨
        hdr = tk.Frame(rule_outer, bg="#f3f4f6")
        hdr.pack(fill="x", pady=(0, 4))
        tk.Label(hdr, text="찾을 텍스트", font=("Malgun Gothic", 9, "bold"),
                 bg="#f3f4f6", width=30, anchor="w").pack(side="left", padx=(30, 0))
        tk.Label(hdr, text="바꿀 텍스트", font=("Malgun Gothic", 9, "bold"),
                 bg="#f3f4f6", anchor="w").pack(side="left", padx=(10, 0))

        # 스크롤 가능한 규칙 목록
        canvas = tk.Canvas(rule_outer, bg="#f3f4f6", highlightthickness=0)
        vscroll = ttk.Scrollbar(rule_outer, orient="vertical", command=canvas.yview)
        self.rule_frame = tk.Frame(canvas, bg="#f3f4f6")
        self.rule_frame.bind("<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas_win = canvas.create_window((0, 0), window=self.rule_frame, anchor="nw")
        canvas.bind("<Configure>", lambda e: canvas.itemconfig(canvas_win, width=e.width))
        canvas.configure(yscrollcommand=vscroll.set)
        canvas.pack(side="left", fill="both", expand=True)
        vscroll.pack(side="right", fill="y")

        # 마우스 휠
        def _wheel(e): canvas.yview_scroll(int(-1*(e.delta/120)), "units")
        canvas.bind("<Enter>", lambda e: canvas.bind_all("<MouseWheel>", _wheel))
        canvas.bind("<Leave>", lambda e: canvas.unbind_all("<MouseWheel>"))

        # 행 추가/삭제 버튼
        rule_btn_row = tk.Frame(rule_outer, bg="#f3f4f6")
        rule_btn_row.pack(fill="x", pady=(6, 0))
        ttk.Button(rule_btn_row, text="＋ 규칙 추가", command=self._add_row).pack(side="left", padx=3)
        ttk.Button(rule_btn_row, text="전체 비우기", command=self._clear_rows).pack(side="left", padx=3)

        # 프리셋
        preset_frame = tk.Frame(rule_btn_row, bg="#f3f4f6")
        preset_frame.pack(side="right")
        tk.Label(preset_frame, text="프리셋:", bg="#f3f4f6",
                 font=("Malgun Gothic", 9)).pack(side="left")
        self.preset_combo = ttk.Combobox(preset_frame, width=18, state="readonly")
        self.preset_combo.pack(side="left", padx=3)
        ttk.Button(preset_frame, text="불러오기", command=self._load_preset).pack(side="left", padx=2)
        ttk.Button(preset_frame, text="저장", command=self._save_preset).pack(side="left", padx=2)
        ttk.Button(preset_frame, text="삭제", command=self._delete_preset).pack(side="left", padx=2)
        self._refresh_preset_combo()

        # 기본 행 3개
        for _ in range(3):
            self._add_row()

        # 옵션
        opt_frame = tk.Frame(self.root, bg="#f3f4f6")
        opt_frame.pack(fill="x", padx=15, pady=3)
        self.backup_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(opt_frame, text="원본 백업 저장 (.bak)", variable=self.backup_var).pack(side="left")
        self.suffix_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(opt_frame, text="새 파일로 저장 (_수정본)", variable=self.suffix_var).pack(side="left", padx=15)

        # 실행 버튼
        run_frame = tk.Frame(self.root, bg="#f3f4f6")
        run_frame.pack(fill="x", padx=15, pady=(5, 5))
        ttk.Button(run_frame, text="▶  교체 실행", command=self._run,
                   style="Accent.TButton").pack(side="right", padx=3, ipadx=10, ipady=4)
        ttk.Button(run_frame, text="미리보기 (첫 번째 파일)",
                   command=self._preview).pack(side="right", padx=3)

        # 로그
        log_frame = ttk.LabelFrame(self.root, text=" 📋 로그 ", padding=8)
        log_frame.pack(fill="x", padx=15, pady=(0, 10))
        self.log_text = tk.Text(log_frame, height=5, font=("Consolas", 9),
                                state="disabled", bg="white")
        self.log_text.pack(side="left", fill="both", expand=True)
        lsb = ttk.Scrollbar(log_frame, orient="vertical", command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=lsb.set)
        lsb.pack(side="right", fill="y")

    # ── 규칙 행 ──────────────────────────────────────────────
    def _add_row(self, find_val="", replace_val=""):
        row_frame = tk.Frame(self.rule_frame, bg="#f3f4f6")
        row_frame.pack(fill="x", pady=2)

        idx_label = tk.Label(row_frame, text=f"{len(self.rows)+1:02d}",
                             font=("Consolas", 9), bg="#f3f4f6", fg="#888", width=3)
        idx_label.pack(side="left")

        find_var = tk.StringVar(value=find_val)
        replace_var = tk.StringVar(value=replace_val)

        find_entry = ttk.Entry(row_frame, textvariable=find_var, width=28,
                               font=("Malgun Gothic", 10))
        find_entry.pack(side="left", padx=(2, 0))

        tk.Label(row_frame, text="→", bg="#f3f4f6", fg="#555",
                 font=("Arial", 12)).pack(side="left", padx=4)

        replace_entry = ttk.Entry(row_frame, textvariable=replace_var, width=28,
                                  font=("Malgun Gothic", 10))
        replace_entry.pack(side="left")

        def _remove():
            row_frame.destroy()
            self.rows = [(f, r, fr) for f, r, fr in self.rows if fr != row_frame]
            self._reindex()

        ttk.Button(row_frame, text="✕", width=3, command=_remove).pack(side="left", padx=4)
        self.rows.append((find_var, replace_var, row_frame))

    def _reindex(self):
        for i, (_, _, frame) in enumerate(self.rows):
            for w in frame.winfo_children():
                if isinstance(w, tk.Label) and len(w.cget("text")) <= 2:
                    try:
                        int(w.cget("text"))
                        w.config(text=f"{i+1:02d}")
                        break
                    except ValueError:
                        pass

    def _clear_rows(self):
        for _, _, frame in self.rows:
            frame.destroy()
        self.rows.clear()
        self._add_row()

    # ── 파일 관리 ─────────────────────────────────────────────
    def _add_files(self):
        paths = filedialog.askopenfilenames(
            filetypes=[("Word 문서", "*.docx"), ("모든 파일", "*.*")])
        for p in paths:
            if p not in self.file_paths:
                self.file_paths.append(p)
                self.file_listbox.insert(tk.END, p)

    def _add_folder(self):
        folder = filedialog.askdirectory()
        if not folder:
            return
        for fn in os.listdir(folder):
            if fn.endswith(".docx") and not fn.startswith("~$"):
                p = os.path.join(folder, fn)
                if p not in self.file_paths:
                    self.file_paths.append(p)
                    self.file_listbox.insert(tk.END, p)

    def _clear_files(self):
        self.file_paths.clear()
        self.file_listbox.delete(0, tk.END)

    # ── 교체 실행 ─────────────────────────────────────────────
    def _get_replacements(self):
        return [(f.get().strip(), r.get()) for f, r, _ in self.rows if f.get().strip()]

    def _run(self):
        if not self.file_paths:
            messagebox.showwarning("경고", "파일을 먼저 추가해주세요.")
            return
        replacements = self._get_replacements()
        if not replacements:
            messagebox.showwarning("경고", "교체 규칙을 입력해주세요.")
            return

        total_files = 0
        total_changes = 0
        errors = []

        for path in self.file_paths:
            try:
                # 백업
                if self.backup_var.get():
                    bak_path = path + ".bak"
                    import shutil
                    shutil.copy2(path, bak_path)

                doc = Document(path)
                results = replace_in_doc(doc, replacements)
                file_changes = sum(results.values())

                # 저장 경로
                if self.suffix_var.get():
                    base, ext = os.path.splitext(path)
                    save_path = base + "_수정본" + ext
                else:
                    save_path = path

                doc.save(save_path)
                total_files += 1
                total_changes += file_changes

                detail = ", ".join(
                    f'"{k}": {v}회' for k, v in results.items() if v > 0
                ) or "변경 없음"
                self._log(f"✅ {os.path.basename(path)} → {file_changes}건 교체 ({detail})")

            except Exception as e:
                errors.append((path, str(e)))
                self._log(f"❌ {os.path.basename(path)} 오류: {e}")

        summary = f"\n완료: {total_files}개 파일, 총 {total_changes}건 교체"
        if errors:
            summary += f" / 오류 {len(errors)}개"
        self._log(summary)
        messagebox.showinfo("완료", summary.strip())

    # ── 미리보기 ──────────────────────────────────────────────
    def _preview(self):
        if not self.file_paths:
            messagebox.showwarning("경고", "파일을 먼저 추가해주세요.")
            return
        replacements = self._get_replacements()
        if not replacements:
            messagebox.showwarning("경고", "교체 규칙을 입력해주세요.")
            return

        path = self.file_paths[0]
        try:
            doc = Document(path)
            results = replace_in_doc(doc, replacements)
        except Exception as e:
            messagebox.showerror("오류", str(e))
            return

        # 미리보기 창
        win = tk.Toplevel(self.root)
        win.title(f"미리보기 - {os.path.basename(path)}")
        win.geometry("700x500")

        text = tk.Text(win, font=("Malgun Gothic", 10), wrap="word")
        sb = ttk.Scrollbar(win, orient="vertical", command=text.yview)
        text.configure(yscrollcommand=sb.set)
        sb.pack(side="right", fill="y")
        text.pack(fill="both", expand=True, padx=10, pady=10)

        # 교체 결과 요약
        text.insert(tk.END, "=== 교체 결과 요약 ===\n")
        for find_text, count in results.items():
            status = f"  '{find_text}' → {count}건 교체\n"
            text.insert(tk.END, status)
        text.insert(tk.END, "\n=== 본문 내용 (교체 후) ===\n\n")

        for para in doc.paragraphs:
            if para.text.strip():
                text.insert(tk.END, para.text + "\n")

        text.config(state="disabled")

        # 저장 버튼
        def _save_preview():
            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word 문서", "*.docx")],
                initialfile=os.path.splitext(os.path.basename(path))[0] + "_수정본.docx"
            )
            if save_path:
                doc.save(save_path)
                messagebox.showinfo("저장 완료", f"저장됨:\n{save_path}")

        ttk.Button(win, text="💾  이 파일 저장", command=_save_preview).pack(pady=8)

    # ── 프리셋 ────────────────────────────────────────────────
    def _load_presets(self):
        if os.path.exists(PRESETS_FILE):
            try:
                with open(PRESETS_FILE, "r", encoding="utf-8") as f:
                    self.presets = json.load(f)
            except Exception:
                self.presets = {}

    def _save_presets_file(self):
        with open(PRESETS_FILE, "w", encoding="utf-8") as f:
            json.dump(self.presets, f, ensure_ascii=False, indent=2)

    def _refresh_preset_combo(self):
        self.preset_combo["values"] = list(self.presets.keys())

    def _save_preset(self):
        name = self._ask_string("프리셋 이름을 입력하세요:", default="프리셋1")
        if not name:
            return
        rules = [(f.get().strip(), r.get()) for f, r, _ in self.rows if f.get().strip()]
        self.presets[name] = rules
        self._save_presets_file()
        self._refresh_preset_combo()
        self.preset_combo.set(name)
        self._log(f"프리셋 '{name}' 저장 완료 ({len(rules)}개 규칙)")

    def _load_preset(self):
        name = self.preset_combo.get()
        if not name or name not in self.presets:
            return
        self._clear_rows()
        for find_val, replace_val in self.presets[name]:
            self._add_row(find_val, replace_val)
        self._log(f"프리셋 '{name}' 불러오기 완료")

    def _delete_preset(self):
        name = self.preset_combo.get()
        if name and name in self.presets:
            del self.presets[name]
            self._save_presets_file()
            self._refresh_preset_combo()
            self.preset_combo.set("")
            self._log(f"프리셋 '{name}' 삭제")

    # ── 유틸 ─────────────────────────────────────────────────
    def _log(self, msg):
        ts = datetime.datetime.now().strftime("%H:%M:%S")
        self.log_text.config(state="normal")
        self.log_text.insert(tk.END, f"[{ts}] {msg}\n")
        self.log_text.see(tk.END)
        self.log_text.config(state="disabled")

    def _ask_string(self, prompt, default=""):
        win = tk.Toplevel(self.root)
        win.title("입력")
        win.geometry("320x110")
        win.grab_set()
        tk.Label(win, text=prompt, font=("Malgun Gothic", 10)).pack(pady=(15, 5))
        var = tk.StringVar(value=default)
        entry = ttk.Entry(win, textvariable=var, width=30)
        entry.pack()
        entry.select_range(0, tk.END)
        entry.focus()
        result = [None]

        def _ok(e=None):
            result[0] = var.get().strip()
            win.destroy()

        entry.bind("<Return>", _ok)
        ttk.Button(win, text="확인", command=_ok).pack(pady=8)
        self.root.wait_window(win)
        return result[0]


# ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    root = tk.Tk()
    app = DocReplacerApp(root)
    root.mainloop()
