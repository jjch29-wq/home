from docx import Document
from docx.shared import Inches
import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk
import os

def create_ndt_procedure_doc(title, purpose, scope, references, equipment, procedure, records, approval, image_paths=None):
    doc = Document()

    # 제목
    doc.add_heading(title, 0)

    # 목적
    doc.add_heading('1. 목적', level=1)
    doc.add_paragraph(purpose)

    # 범위
    doc.add_heading('2. 범위', level=1)
    doc.add_paragraph(scope)

    # 참조 문서
    doc.add_heading('3. 참조 문서', level=1)
    doc.add_paragraph(references)

    # 용어 정의 (기본)
    doc.add_heading('4. 용어 정의', level=1)
    doc.add_paragraph("NDT: Non-Destructive Testing (비파괴 검사)\nPMI: Positive Material Identification (양성 물질 식별)")

    # 장비
    doc.add_heading('5. 장비', level=1)
    doc.add_paragraph(equipment)

    # 절차
    doc.add_heading('6. 절차', level=1)
    doc.add_paragraph(procedure)

    # 사진 추가 (절차와 기록 사이)
    if image_paths:
        doc.add_heading('7. 첨부 사진', level=1)
        for img_path in image_paths:
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(5))
                    doc.add_paragraph()
                except:
                    doc.add_paragraph(f"[사진 삽입 실패: {os.path.basename(img_path)}]")

    # 기록
    doc.add_heading('8. 기록', level=1)
    doc.add_paragraph(records)

    # 승인
    doc.add_heading('9. 승인', level=1)
    doc.add_paragraph(approval)

    return doc

def load_existing_doc(file_path):
    doc = Document(file_path)
    content = {
        'title': '',
        'purpose': '',
        'scope': '',
        'references': '',
        'equipment': '',
        'procedure': '',
        'records': '',
        'approval': ''
    }
    
    current_section = None
    section_markers = {
        '1.': 'purpose',
        '2.': 'scope',
        '3.': 'references',
        '5.': 'equipment',
        '6.': 'procedure',
        '8.': 'records',
        '9.': 'approval'
    }
    
    # 모든 paragraph 읽기
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if not text:
            continue
        
        # 제목 (첫 번째 heading 또는 비어있는 경우 첫 번째 문장)
        if not content['title']:
            if para.style.name == 'Heading 1' or para.style.name == 'Title':
                content['title'] = text
                continue
            elif i == 0:
                content['title'] = text
                continue
        
        # Heading 1 또는 번호가 있는 제목에서 섹션 찾기
        if para.style.name in ['Heading 1', 'Heading 2']:
            for marker, section in section_markers.items():
                if marker in text or section.replace('_', ' ') in text.lower():
                    current_section = section
                    break
            continue
        
        # 현재 섹션에 내용 추가
        if current_section:
            if content[current_section]:
                content[current_section] += '\n' + text
            else:
                content[current_section] = text
    
    # 공백 정리
    for key in content:
        content[key] = content[key].strip()
    
    return content

def extract_images_from_doc(file_path):
    """Word 문서에서 이미지 추출"""
    doc = Document(file_path)
    image_paths = []
    temp_dir = os.path.join(os.path.dirname(file_path), ".temp_images")
    
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)
    
    # 문서의 모든 관계에서 이미지 찾기
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_part = rel.target_part
                image_data = image_part.blob
                # 파일 확장자 결정
                content_type = image_part.content_type
                ext = '.jpg' if 'jpeg' in content_type else '.png' if 'png' in content_type else '.gif'
                # 이미지 저장
                image_file = os.path.join(temp_dir, f"image_{len(image_paths)}{ext}")
                with open(image_file, 'wb') as f:
                    f.write(image_data)
                image_paths.append(image_file)
    except Exception as e:
        print(f"이미지 추출 오류: {e}")
    
    return image_paths

class NDTProcedureApp:
    def __init__(self, root):
        self.root = root
        self.root.title("비파괴 검사 절차서 생성기")
        self.root.geometry("1200x1000")  # 창 크기 더 늘림
        self.root.resizable(True, True)  # 크기 조정 가능

        # 로드 버튼
        self.load_button = tk.Button(root, text="기존 문서 로드", command=self.load_doc)
        self.load_button.pack(pady=10)

        # 사진 추가 버튼
        self.image_button = tk.Button(root, text="사진 추가", command=self.add_images)
        self.image_button.pack(pady=5)

        # 사진 미리보기 프레임
        self.image_preview_frame = tk.Frame(root, bg="lightgray", height=100)
        self.image_preview_frame.pack(fill=tk.X, padx=10, pady=5)
        self.image_preview_label = tk.Label(self.image_preview_frame, text="사진 추가 예정 없음", bg="lightgray")
        self.image_preview_label.pack()

        self.image_paths = []

        # 입력 필드들
        tk.Label(root, text="제목:").pack()
        self.title_entry = tk.Entry(root, width=70)  # 너비 늘림
        self.title_entry.pack()
        self.title_entry.insert(0, "비파괴 검사 절차서")

        tk.Label(root, text="목적:").pack()
        self.purpose_frame = tk.Frame(root)
        self.purpose_frame.pack()
        self.purpose_scrollbar = tk.Scrollbar(self.purpose_frame)
        self.purpose_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.purpose_text = tk.Text(self.purpose_frame, height=4, width=70, yscrollcommand=self.purpose_scrollbar.set)
        self.purpose_text.pack(side=tk.LEFT)
        self.purpose_scrollbar.config(command=self.purpose_text.yview)
        self.purpose_text.insert(tk.END, "본 절차서는 비파괴 검사를 수행하기 위한 방법을 정의한다.")

        tk.Label(root, text="범위:").pack()
        self.scope_frame = tk.Frame(root)
        self.scope_frame.pack()
        self.scope_scrollbar = tk.Scrollbar(self.scope_frame)
        self.scope_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.scope_text = tk.Text(self.scope_frame, height=4, width=70, yscrollcommand=self.scope_scrollbar.set)
        self.scope_text.pack(side=tk.LEFT)
        self.scope_scrollbar.config(command=self.scope_text.yview)
        self.scope_text.insert(tk.END, "본 절차서는 [프로젝트명]에 적용된다.")

        tk.Label(root, text="참조 문서:").pack()
        self.references_frame = tk.Frame(root)
        self.references_frame.pack()
        self.references_scrollbar = tk.Scrollbar(self.references_frame)
        self.references_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.references_text = tk.Text(self.references_frame, height=4, width=70, yscrollcommand=self.references_scrollbar.set)
        self.references_text.pack(side=tk.LEFT)
        self.references_scrollbar.config(command=self.references_text.yview)
        self.references_text.insert(tk.END, "ASTM E1444, ASME V 등")

        tk.Label(root, text="장비:").pack()
        self.equipment_frame = tk.Frame(root)
        self.equipment_frame.pack()
        self.equipment_scrollbar = tk.Scrollbar(self.equipment_frame)
        self.equipment_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.equipment_text = tk.Text(self.equipment_frame, height=4, width=70, yscrollcommand=self.equipment_scrollbar.set)
        self.equipment_text.pack(side=tk.LEFT)
        self.equipment_scrollbar.config(command=self.equipment_text.yview)
        self.equipment_text.insert(tk.END, "XRF 분석기, 초음파 장비 등")

        tk.Label(root, text="절차:").pack()
        self.procedure_frame = tk.Frame(root)
        self.procedure_frame.pack()
        self.procedure_scrollbar = tk.Scrollbar(self.procedure_frame)
        self.procedure_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.procedure_text = tk.Text(self.procedure_frame, height=6, width=70, yscrollcommand=self.procedure_scrollbar.set)
        self.procedure_text.pack(side=tk.LEFT)
        self.procedure_scrollbar.config(command=self.procedure_text.yview)
        self.procedure_text.insert(tk.END, "1. 준비\n2. 검사\n3. 기록")

        tk.Label(root, text="기록:").pack()
        self.records_frame = tk.Frame(root)
        self.records_frame.pack()
        self.records_scrollbar = tk.Scrollbar(self.records_frame)
        self.records_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.records_text = tk.Text(self.records_frame, height=4, width=70, yscrollcommand=self.records_scrollbar.set)
        self.records_text.pack(side=tk.LEFT)
        self.records_scrollbar.config(command=self.records_text.yview)
        self.records_text.insert(tk.END, "검사 결과 기록 양식")

        tk.Label(root, text="승인:").pack()
        self.approval_frame = tk.Frame(root)
        self.approval_frame.pack()
        self.approval_scrollbar = tk.Scrollbar(self.approval_frame)
        self.approval_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.approval_text = tk.Text(self.approval_frame, height=4, width=70, yscrollcommand=self.approval_scrollbar.set)
        self.approval_text.pack(side=tk.LEFT)
        self.approval_scrollbar.config(command=self.approval_text.yview)
        self.approval_text.insert(tk.END, "담당자 서명")

        # 생성 버튼
        self.generate_button = tk.Button(root, text="Word 문서 생성", command=self.generate_doc)
        self.generate_button.pack(pady=20)

    def add_images(self):
        file_paths = filedialog.askopenfilenames(filetypes=[("Image files", "*.jpg *.jpeg *.png *.gif *.bmp")])
        if file_paths:
            self.image_paths.extend(file_paths)
            self.update_image_preview()
            messagebox.showinfo("완료", f"{len(file_paths)}개의 사진이 추가되었습니다.")

    def update_image_preview(self):
        if self.image_paths:
            preview_text = f"추가된 사진: {len(self.image_paths)}개\n"
            for i, path in enumerate(self.image_paths[:3], 1):  # 처음 3개만 표시
                preview_text += f"{i}. {os.path.basename(path)}\n"
            if len(self.image_paths) > 3:
                preview_text += f"... 외 {len(self.image_paths) - 3}개"
            self.image_preview_label.config(text=preview_text)
        else:
            self.image_preview_label.config(text="사진 추가 예정 없음")

    def load_doc(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if file_path:
            try:
                content = load_existing_doc(file_path)
                self.title_entry.delete(0, tk.END)
                self.title_entry.insert(0, content.get('title', ''))
                self.purpose_text.delete("1.0", tk.END)
                self.purpose_text.insert(tk.END, content.get('purpose', ''))
                self.scope_text.delete("1.0", tk.END)
                self.scope_text.insert(tk.END, content.get('scope', ''))
                self.references_text.delete("1.0", tk.END)
                self.references_text.insert(tk.END, content.get('references', ''))
                self.equipment_text.delete("1.0", tk.END)
                self.equipment_text.insert(tk.END, content.get('equipment', ''))
                self.procedure_text.delete("1.0", tk.END)
                self.procedure_text.insert(tk.END, content.get('procedure', ''))
                self.records_text.delete("1.0", tk.END)
                self.records_text.insert(tk.END, content.get('records', ''))
                self.approval_text.delete("1.0", tk.END)
                self.approval_text.insert(tk.END, content.get('approval', ''))
                
                # 문서에서 사진 추출
                extracted_images = extract_images_from_doc(file_path)
                if extracted_images:
                    self.image_paths.extend(extracted_images)
                    self.update_image_preview()
                    messagebox.showinfo("로드 완료", f"기존 문서 내용과 {len(extracted_images)}개의 사진을 로드했습니다.")
                else:
                    messagebox.showinfo("로드 완료", "기존 문서 내용을 로드했습니다.")
            except Exception as e:
                messagebox.showerror("오류", f"문서 로드 실패: {str(e)}")

    def generate_doc(self):
        title = self.title_entry.get()
        purpose = self.purpose_text.get("1.0", tk.END).strip()
        scope = self.scope_text.get("1.0", tk.END).strip()
        references = self.references_text.get("1.0", tk.END).strip()
        equipment = self.equipment_text.get("1.0", tk.END).strip()
        procedure = self.procedure_text.get("1.0", tk.END).strip()
        records = self.records_text.get("1.0", tk.END).strip()
        approval = self.approval_text.get("1.0", tk.END).strip()

        doc = create_ndt_procedure_doc(title, purpose, scope, references, equipment, procedure, records, approval, self.image_paths)

        output_file = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if output_file:
            doc.save(output_file)
            messagebox.showinfo("완료", f"절차서가 {output_file}에 저장되었습니다.")

if __name__ == "__main__":
    root = tk.Tk()
    app = NDTProcedureApp(root)
    root.mainloop()