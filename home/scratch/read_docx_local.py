import docx
try:
    doc = docx.Document('scratch/KS_B_0845_temp.docx')
    texts = []
    for p in doc.paragraphs:
        if p.text.strip(): texts.append(p.text.strip())
    for t in doc.tables:
        for r in t.rows:
            texts.append('|'.join([c.text.strip().replace('
', ' ') for c in r.cells]))
    with open('scratch/doc_dump.txt', 'w', encoding='utf-8') as f:
        f.write('
'.join(texts))
    print('Extract OK')
except Exception as e:
    print('Error:', e)