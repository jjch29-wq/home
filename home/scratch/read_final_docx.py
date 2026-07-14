import docx
try:
    doc = docx.Document(r'C:\Users\-\PMI\home\ks b 0845ks.docx')
    texts = []
    for p in doc.paragraphs:
        if p.text.strip(): texts.append(p.text.strip())
    for t_idx, t in enumerate(doc.tables):
        texts.append(f'--- TABLE {t_idx} ---')
        for r in t.rows:
            texts.append('|'.join([c.text.strip().replace('
', ' ') for c in r.cells]))
    with open('scratch/docx_dump.txt', 'w', encoding='utf-8') as f:
        f.write('
'.join(texts))
    print('Extract OK')
except Exception as e:
    print('Error:', e)