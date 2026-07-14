import docx
path = r'C:\Users\-\OneDrive\바탕 화면\KS B 0845(2026ed)Rev.0.dotx'
try:
    doc = docx.Document(path)
    print('--- PARAGRAPHS ---')
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            print(f'[{i}] {text}')
    print('\n--- TABLES ---')
    for t_idx, table in enumerate(doc.tables):
        print(f'\nTable {t_idx}:')
        for r_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip().replace('
', ' ') for cell in row.cells]
            print(f'  Row {r_idx}: {row_data}')
except Exception as e:
    print('Error:', e)
