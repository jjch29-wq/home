import os
try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
except ImportError:
    print("python-docx 라이브러리가 필요합니다. 'pip install python-docx'를 실행해주세요.")
    exit()

def create_photo_log(output_filename="휴게시설_사진대장.docx", image_dir="images"):
    print("보고서 생성을 시작합니다...")
    doc = Document()
    
    # 문서 제목 추가
    title = doc.add_heading('휴게시설 점검 사진대장', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 기본 정보 추가
    doc.add_paragraph('점검일자: 2026.08.14')
    doc.add_paragraph('점검위치: 현장 휴게시설')
    doc.add_paragraph('') # 빈 줄
    
    # 6가지 점검 항목 정의
    items = [
        "1. 휴게시설 크기는 적정한가?\n- 최소 바닥면적 6㎡ 이상, 천장 높이 2.1m 이상",
        "2. 휴게시설 위치는 편리한가?\n- 왕복 이동시간이 휴게시간의 20% 미만",
        "3. 적정한 온도(18~28℃) 유지\n- 냉난방 기능 작동 상태",
        "4. 적정한 습도(50~55%) 유지\n- 습도 조절 기능 작동 상태",
        "5. 적절한 밝기(100~200Lux) 유지\n- 조명 조절 기능 상태",
        "6. 환기 및 비품, 식수 비치\n- 창문(환기), 의자, 정수기 등 비치 여부"
    ]
    
    # 표 생성 (6행 2열)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid' # 표 테두리 스타일 지정
    
    for i, item_text in enumerate(items):
        row = table.rows[i]
        
        # 첫 번째 열: 사진 칸
        cell_photo = row.cells[0]
        cell_photo.width = Inches(4.0)
        cell_photo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        p = cell_photo.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        img_path = os.path.join(image_dir, f"photo_{i+1}.jpg")
        
        if os.path.exists(img_path):
            # 사진이 있으면 문서에 삽입
            run = p.add_run()
            run.add_picture(img_path, width=Inches(3.5))
        else:
            # 사진이 없으면 안내 문구 표시
            run = p.add_run(f"\n[사진 첨부 공간]\n\n{image_dir} 폴더에 'photo_{i+1}.jpg' 파일을 넣어주세요.\n")
            run.font.color.rgb = docx.shared.RGBColor(255, 0, 0) # 빨간색
            
        # 두 번째 열: 설명 칸
        cell_desc = row.cells[1]
        cell_desc.width = Inches(2.5)
        cell_desc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        p_desc = cell_desc.paragraphs[0]
        run_desc = p_desc.add_run(item_text)
        run_desc.font.size = Pt(11)
        run_desc.font.bold = True
        
    doc.save(output_filename)
    print(f"완료! '{output_filename}' 파일이 성공적으로 생성되었습니다.")

if __name__ == "__main__":
    # 사진을 넣을 폴더 자동 생성
    img_folder = "images"
    os.makedirs(img_folder, exist_ok=True)
    
    print("="*50)
    print(f"1. 먼저 '{img_folder}' 폴더에 6장의 사진을 준비해주세요.")
    print("   (파일이름: photo_1.jpg, photo_2.jpg ... photo_6.jpg)")
    print("2. 스크립트를 실행하면 사진과 설명이 포함된 워드 문서가 만들어집니다.")
    print("="*50)
    
    create_photo_log()
