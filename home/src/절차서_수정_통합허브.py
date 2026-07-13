import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import json

class ProcedureHubApp:
    def __init__(self, root):
        self.root = root
        self.root.title("절차서 수정 통합 허브 (비교 + 다중 파일 일괄 변환 + 프리셋)")
        self.root.geometry("800x850")
        
        style = ttk.Style()
        style.theme_use('clam')
        
        # 탭 생성
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill="both", expand=True, padx=5, pady=5)
        
        # 탭 1 프레임 (수정 및 변환)
        self.tab1 = ttk.Frame(self.notebook)
        self.notebook.add(self.tab1, text="✍️ 절차서 편집 및 다중 일괄 변환")
        
        # 탭 2 프레임 (문서 비교)
        self.tab2 = ttk.Frame(self.notebook)
        self.notebook.add(self.tab2, text="🔍 개정 내용 비교 (구버전 vs 신버전)")
        
        self.create_widgets_tab1()
        self.create_widgets_tab2()
        
    def create_widgets_tab1(self):
        # 1. 파일 다중 선택 프레임
        file_frame = ttk.LabelFrame(self.tab1, text="1. 원본 절차서 파일 선택 (여러 파일 동시 선택 가능)")
        file_frame.pack(fill="x", padx=15, pady=10)
        
        self.file_listbox = tk.Listbox(file_frame, height=4, selectmode=tk.EXTENDED)
        self.file_listbox.pack(side="left", padx=10, pady=10, expand=True, fill="both")
        
        # 리스트박스 스크롤바
        scrollbar = ttk.Scrollbar(self.file_listbox, orient="vertical")
        scrollbar.config(command=self.file_listbox.yview)
        scrollbar.pack(side="right", fill="y")
        self.file_listbox.config(yscrollcommand=scrollbar.set)
        
        btn_frame = ttk.Frame(file_frame)
        btn_frame.pack(side="right", padx=10, pady=10)
        ttk.Button(btn_frame, text="파일 추가", command=self.add_files).pack(fill="x", pady=2)
        ttk.Button(btn_frame, text="선택 삭제", command=self.remove_file).pack(fill="x", pady=2)
        ttk.Button(btn_frame, text="미리보기(Text)", command=self.preview_file).pack(fill="x", pady=2)
        
        # 2. 직접 열어서 편집하기 프레임
        direct_edit_frame = ttk.LabelFrame(self.tab1, text="2. 단일 문서 직접 편집하기 (위 목록에서 선택한 첫 번째 파일이 열립니다)")
        direct_edit_frame.pack(fill="x", padx=15, pady=5)
        
        ttk.Label(direct_edit_frame, text="표, 이미지, 서식 등을 마음껏 자유롭게 수정할 수 있습니다.").pack(pady=(5, 0))
        btn_open = ttk.Button(direct_edit_frame, text="🛠️ 선택한 문서 열어서 직접 수정하기", command=self.open_file_natively)
        btn_open.pack(pady=10, ipady=5)
        
        # 구분선
        ttk.Separator(self.tab1, orient='horizontal').pack(fill='x', padx=15, pady=10)
        
        # 3. 단어 일괄 변환 프레임 (프리셋 기능 포함)
        list_frame = ttk.LabelFrame(self.tab1, text="3. 단어/코드 일괄 자동 변환 (다중 파일 동시 적용 가능)")
        list_frame.pack(fill="both", expand=True, padx=15, pady=5)
        
        preset_frame = ttk.Frame(list_frame)
        preset_frame.pack(fill="x", padx=10, pady=5)
        ttk.Label(preset_frame, text="자주 쓰는 단어 세트를 저장하고 불러올 수 있습니다.").pack(side="left")
        ttk.Button(preset_frame, text="💾 현재 목록 저장", command=self.save_preset).pack(side="right", padx=2)
        ttk.Button(preset_frame, text="📂 목록 불러오기", command=self.load_preset).pack(side="right", padx=2)
        
        columns = ("find", "replace")
        self.tree = ttk.Treeview(list_frame, columns=columns, show="headings", height=6)
        self.tree.heading("find", text="찾을 내용 (기존 코드/문구)")
        self.tree.heading("replace", text="바꿀 내용 (새로운 코드/문구)")
        self.tree.column("find", width=300)
        self.tree.column("replace", width=300)
        self.tree.pack(fill="both", expand=True, padx=10, pady=5)
        
        input_frame = ttk.Frame(list_frame)
        input_frame.pack(fill="x", padx=10, pady=5)
        
        ttk.Label(input_frame, text="찾을 내용:").grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.entry_find = ttk.Entry(input_frame, width=20)
        self.entry_find.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(input_frame, text="바꿀 내용:").grid(row=0, column=2, padx=5, pady=5, sticky="e")
        self.entry_replace = ttk.Entry(input_frame, width=20)
        self.entry_replace.grid(row=0, column=3, padx=5, pady=5)
        
        self.entry_replace.bind("<Return>", lambda e: self.add_item())
        self.entry_find.bind("<Return>", lambda e: self.entry_replace.focus())
        
        ttk.Button(input_frame, text="추가", command=self.add_item, width=8).grid(row=0, column=4, padx=5)
        ttk.Button(input_frame, text="수정", command=self.update_item, width=8).grid(row=0, column=5, padx=5)
        ttk.Button(input_frame, text="삭제", command=self.delete_item, width=8).grid(row=0, column=6, padx=5)
        
        self.tree.bind("<<TreeviewSelect>>", self.on_tree_select)
        
        run_frame = ttk.Frame(list_frame)
        run_frame.pack(fill="x", padx=10, pady=10)
        
        ttk.Button(run_frame, text="위 목록대로 1번에 등록된 모든 파일을 일괄 변환하여 폴더에 자동 저장", command=self.process_files).pack(fill="x", ipady=10)

    def create_widgets_tab2(self):
        desc = ("[문서 개정안 자동 비교기]\n\n"
                "과거 절차서(구버전)와 개정된 절차서(신버전)를 각각 선택하면,\n"
                "워드(Word)의 강력한 자체 비교 기능을 사용하여\n"
                "정확히 어느 문구, 어느 표, 어느 값이 바뀌었는지\n"
                "빨간색 줄과 함께 상세한 비교 화면을 띄워줍니다.\n\n"
                "(⚠️ 현재 이 기능은 워드 파일(.docx, .doc)에 대해서만 작동합니다.)")
        ttk.Label(self.tab2, text=desc, justify="center", font=("맑은 고딕", 11)).pack(pady=30)
        
        comp_frame = ttk.Frame(self.tab2)
        comp_frame.pack(fill="x", padx=30, pady=10)
        
        self.old_file_var = tk.StringVar()
        self.new_file_var = tk.StringVar()
        
        ttk.Button(comp_frame, text="📁 구버전(Old) 절차서 선택", command=lambda: self.browse_comp_file(self.old_file_var), width=25).grid(row=0, column=0, padx=5, pady=10)
        ttk.Entry(comp_frame, textvariable=self.old_file_var, width=50).grid(row=0, column=1, padx=5, pady=10, sticky="ew")
        
        ttk.Button(comp_frame, text="📁 신버전(New) 절차서 선택", command=lambda: self.browse_comp_file(self.new_file_var), width=25).grid(row=1, column=0, padx=5, pady=10)
        ttk.Entry(comp_frame, textvariable=self.new_file_var, width=50).grid(row=1, column=1, padx=5, pady=10, sticky="ew")
        
        ttk.Button(self.tab2, text="✨ 바뀐 내용 전격 비교하기 (워드 화면 열기)", command=self.compare_docs, width=40).pack(pady=40, ipady=15)

    def add_files(self):
        filepaths = filedialog.askopenfilenames(
            title="절차서 파일 선택 (여러 개 선택 가능)",
            filetypes=[
                ("모든 지원 파일", "*.docx *.xlsx *.hwp *.hwpx *.txt"),
                ("Word 문서", "*.docx"),
                ("Excel 문서", "*.xlsx"),
                ("한글 문서", "*.hwp *.hwpx"),
                ("텍스트 문서", "*.txt"),
                ("모든 파일", "*.*")
            ]
        )
        for path in filepaths:
            if path not in self.file_listbox.get(0, tk.END):
                self.file_listbox.insert(tk.END, path)

    def remove_file(self):
        selected = self.file_listbox.curselection()
        for index in reversed(selected):
            self.file_listbox.delete(index)

    def browse_comp_file(self, var):
        filepath = filedialog.askopenfilename(
            title="비교할 Word 문서 선택",
            filetypes=[("Word 문서", "*.docx *.doc")]
        )
        if filepath:
            var.set(filepath)

    def open_file_natively(self):
        files = self.file_listbox.get(0, tk.END)
        if not files:
            messagebox.showerror("오류", "먼저 1번 목록에 원본 파일을 추가해주세요.")
            return
        
        # 선택된 항목이 있으면 그걸 열고, 없으면 첫번째 항목 염
        selected = self.file_listbox.curselection()
        filepath = self.file_listbox.get(selected[0]) if selected else files[0]
            
        try:
            os.startfile(filepath)
            messagebox.showinfo("안내", "문서를 열었습니다.\n표, 이미지, 서식 등을 수정한 후 해당 프로그램에서 직접 [저장]을 눌러주세요.")
        except Exception as e:
            messagebox.showerror("실행 오류", f"문서를 여는 데 실패했습니다:\n{e}")

    def add_item(self):
        f_text = self.entry_find.get().strip()
        r_text = self.entry_replace.get().strip()
        if f_text:
            self.tree.insert("", "end", values=(f_text, r_text))
            self.entry_find.delete(0, tk.END)
            self.entry_replace.delete(0, tk.END)
            self.entry_find.focus()
        else:
            messagebox.showwarning("입력 오류", "찾을 내용을 입력해주세요.")
            
    def update_item(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("선택 오류", "수정할 항목을 위 목록에서 선택해주세요.")
            return
        f_text = self.entry_find.get().strip()
        r_text = self.entry_replace.get().strip()
        if f_text:
            self.tree.item(selected[0], values=(f_text, r_text))
        else:
            messagebox.showwarning("입력 오류", "찾을 내용을 입력해주세요.")
            
    def on_tree_select(self, event):
        selected = self.tree.selection()
        if selected:
            item = self.tree.item(selected[0])
            val = item['values']
            self.entry_find.delete(0, tk.END)
            self.entry_find.insert(0, val[0])
            self.entry_replace.delete(0, tk.END)
            self.entry_replace.insert(0, val[1])

    def delete_item(self):
        selected = self.tree.selection()
        if selected:
            for item in selected:
                self.tree.delete(item)
                
    def preview_file(self):
        selected = self.file_listbox.curselection()
        if not selected:
            messagebox.showwarning("선택 오류", "미리보기할 파일을 위 목록에서 먼저 선택해주세요.")
            return
            
        filepath = self.file_listbox.get(selected[0])
        ext = os.path.splitext(filepath)[1].lower()
        
        content = ""
        try:
            if ext == '.docx':
                import docx
                doc = docx.Document(filepath)
                # 문단 텍스트만 추출
                content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            elif ext == '.txt':
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        content = f.read()
                except:
                    with open(filepath, 'r', encoding='euc-kr') as f:
                        content = f.read()
            else:
                messagebox.showinfo("안내", f"{ext} 파일은 텍스트 미리보기를 지원하지 않습니다.\n[선택한 문서 열어서 직접 수정하기] 버튼을 이용해주세요.")
                return
                
            # 텍스트가 비어있는 경우
            if not content.strip():
                content = "(추출된 텍스트가 없습니다. 문서가 비어있거나 스캔 이미지 형태일 수 있습니다.)"
                
            # 팝업 띄우기
            top = tk.Toplevel(self.root)
            top.title(f"텍스트 미리보기 - {os.path.basename(filepath)}")
            top.geometry("600x600")
            
            top_frame = ttk.Frame(top)
            top_frame.pack(fill="both", expand=True, padx=10, pady=10)
            
            txt_widget = tk.Text(top_frame, wrap="word", font=("맑은 고딕", 10))
            
            scrollbar = ttk.Scrollbar(top_frame, orient="vertical", command=txt_widget.yview)
            scrollbar.pack(side="right", fill="y")
            txt_widget.pack(side="left", fill="both", expand=True)
            txt_widget.config(yscrollcommand=scrollbar.set)
            
            txt_widget.insert("1.0", content)
            txt_widget.config(state="disabled") # 읽기 전용
            
        except Exception as e:
            messagebox.showerror("오류", f"미리보기를 불러오는 중 오류가 발생했습니다:\n{e}")

    # 프리셋 저장/불러오기 기능
    def save_preset(self):
        items = self.tree.get_children()
        if not items:
            messagebox.showwarning("경고", "저장할 단어 목록이 없습니다.")
            return
            
        preset_data = []
        for item in items:
            val = self.tree.item(item, 'values')
            preset_data.append({"find": val[0], "replace": val[1]})
            
        filepath = filedialog.asksaveasfilename(
            title="단어 목록 저장",
            defaultextension=".json",
            filetypes=[("JSON 파일", "*.json")],
            initialfile="자주쓰는단어.json"
        )
        if filepath:
            try:
                with open(filepath, 'w', encoding='utf-8') as f:
                    json.dump(preset_data, f, ensure_ascii=False, indent=4)
                messagebox.showinfo("저장 완료", "단어 목록이 성공적으로 저장되었습니다.")
            except Exception as e:
                messagebox.showerror("오류", f"저장 중 오류 발생:\n{e}")

    def load_preset(self):
        filepath = filedialog.askopenfilename(
            title="단어 목록 불러오기",
            filetypes=[("JSON 파일", "*.json")]
        )
        if filepath:
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    preset_data = json.load(f)
                    
                # 기존 목록 초기화 후 불러오기
                for item in self.tree.get_children():
                    self.tree.delete(item)
                    
                for entry in preset_data:
                    self.tree.insert("", "end", values=(entry.get("find", ""), entry.get("replace", "")))
                    
                messagebox.showinfo("불러오기 완료", "단어 목록을 성공적으로 불러왔습니다.")
            except Exception as e:
                messagebox.showerror("오류", f"불러오기 중 오류 발생:\n{e}")

    # ==================== 단어 일괄 변환 처리부 ====================
    def replace_text_in_paragraph(self, paragraph, find_text, replace_text):
        if find_text in paragraph.text:
            for run in paragraph.runs:
                if find_text in run.text:
                    run.text = run.text.replace(find_text, replace_text)
                    return
            inline = paragraph.runs
            if not inline: return
            text = paragraph.text.replace(find_text, replace_text)
            for i in range(len(inline)): inline[i].text = ''
            inline[0].text = text

    def process_docx(self, input_file, output_file, replacements):
        import docx
        doc = docx.Document(input_file)
        for paragraph in doc.paragraphs:
            for f_text, r_text in replacements:
                self.replace_text_in_paragraph(paragraph, f_text, r_text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for f_text, r_text in replacements:
                            self.replace_text_in_paragraph(paragraph, f_text, r_text)
        doc.save(output_file)

    def process_xlsx(self, input_file, output_file, replacements):
        import openpyxl
        wb = openpyxl.load_workbook(input_file)
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        new_val = cell.value
                        for f_text, r_text in replacements:
                            new_val = new_val.replace(f_text, r_text)
                        if new_val != cell.value:
                            cell.value = new_val
        wb.save(output_file)

    def process_hwp(self, input_file, output_file, replacements):
        import win32com.client as win32
        hwp = win32.gencache.EnsureDispatch("HWPFrame.HwpObject")
        hwp.RegisterModule("FilePathCheckDLL", "SecurityModule")
        hwp.XHwpWindows.Item(0).Visible = False
        
        if not hwp.Open(input_file):
            hwp.Quit()
            raise Exception("HWP 파일을 여는데 실패했습니다.")
            
        hwp.HAction.GetDefault("AllReplace", hwp.HParameterSet.HFindReplace.HSet)
        hwp.HParameterSet.HFindReplace.IgnoreMessage = 1
        
        for f_text, r_text in replacements:
            hwp.HParameterSet.HFindReplace.FindString = f_text
            hwp.HParameterSet.HFindReplace.ReplaceString = r_text
            hwp.HParameterSet.HFindReplace.Direction = hwp.FindDir("AllDoc")
            hwp.HAction.Execute("AllReplace", hwp.HParameterSet.HFindReplace.HSet)
            
        hwp.SaveAs(output_file)
        hwp.Quit()

    def process_txt(self, input_file, output_file, replacements):
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                content = f.read()
            enc = 'utf-8'
        except:
            with open(input_file, 'r', encoding='euc-kr') as f:
                content = f.read()
            enc = 'euc-kr'
            
        for f_text, r_text in replacements:
            content = content.replace(f_text, r_text)
            
        with open(output_file, 'w', encoding=enc) as f:
            f.write(content)

    def process_files(self):
        files = self.file_listbox.get(0, tk.END)
        if not files:
            messagebox.showerror("오류", "변환할 파일을 먼저 추가해주세요.")
            return
            
        items = self.tree.get_children()
        if not items:
            messagebox.showwarning("경고", "변경할 단어 목록이 비어있습니다.")
            return
            
        replacements = []
        for item in items:
            val = self.tree.item(item, 'values')
            replacements.append((val[0], val[1]))
            
        # 저장할 폴더 선택
        output_dir = filedialog.askdirectory(title="변환된 새 파일들을 저장할 폴더 선택")
        if not output_dir:
            return
            
        try:
            self.root.config(cursor="wait")
            self.root.update()
            
            success_count = 0
            for input_file in files:
                if not os.path.exists(input_file): continue
                
                filename = os.path.basename(input_file)
                ext = os.path.splitext(filename)[1].lower()
                output_file = os.path.join(output_dir, f"일괄변환_{filename}")
                
                if ext == '.docx':
                    self.process_docx(input_file, output_file, replacements)
                elif ext == '.xlsx':
                    self.process_xlsx(input_file, output_file, replacements)
                elif ext in ['.hwp', '.hwpx']:
                    self.process_hwp(input_file, output_file, replacements)
                else:
                    self.process_txt(input_file, output_file, replacements)
                    
                success_count += 1
                
            self.root.config(cursor="")
            messagebox.showinfo("완료", f"총 {success_count}개의 파일이 성공적으로 일괄 변환 및 저장되었습니다!\n\n저장 폴더: {output_dir}")
            os.startfile(output_dir)
            
        except Exception as e:
            self.root.config(cursor="")
            messagebox.showerror("실행 오류", f"파일을 자동 변환하는 중 오류가 발생했습니다:\n{e}")

    def compare_docs(self):
        old_file = self.old_file_var.get()
        new_file = self.new_file_var.get()
        if not old_file or not new_file or not os.path.exists(old_file) or not os.path.exists(new_file):
            messagebox.showerror("오류", "구버전과 신버전 파일을 모두 정상적으로 선택해주세요.")
            return
            
        try:
            self.root.config(cursor="wait")
            self.root.update()
            
            import win32com.client as win32
            word = win32.gencache.EnsureDispatch('Word.Application')
            word.Visible = True
            
            doc_old = word.Documents.Open(os.path.abspath(old_file), ReadOnly=True, Visible=False)
            doc_new = word.Documents.Open(os.path.abspath(new_file), ReadOnly=True, Visible=False)
            
            word.CompareDocuments(doc_old, doc_new)
            
            doc_old.Close(False)
            doc_new.Close(False)
            
            self.root.config(cursor="")
            messagebox.showinfo("비교 완료", "변경점이 표시된 비교 문서가 워드로 열렸습니다!\n워드 창을 확인해 주세요.")
            
        except Exception as e:
            self.root.config(cursor="")
            messagebox.showerror("실행 오류", f"문서를 비교하는 중 오류가 발생했습니다:\n{e}\n\n(Word가 정상 설치되어 있는지 확인하세요.)")

if __name__ == "__main__":
    root = tk.Tk()
    app = ProcedureHubApp(root)
    root.mainloop()
