import os
import tkinter as tk
from tkinter import filedialog, messagebox

try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
except ImportError:
    messagebox.showerror("오류", "python-docx 라이브러리가 필요합니다.\n명령프롬프트에서 'pip install python-docx'를 실행해주세요.")
    exit()

def create_photo_log(image_paths, output_filename="휴게시설_사진대장.docx"):
    doc = Document()
    
    # 여백을 최소화하여 한 장에 무조건 들어가도록 설정
    from docx.enum.section import WD_ORIENT
    for section in doc.sections:
        # 가로 방향(Landscape) 설정
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
    title = doc.add_heading('휴게시설 점검 사진대장', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 공간을 절약하기 위해 한 줄로 합침
    info_p = doc.add_paragraph('점검업체: 서울검사(주)    |    점검일자: 2026.08.14    |    점검위치: 현장 휴게시설')
    info_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    items = [
        "1. 크기 적정성\n(면적 6㎡, 높이 2.1m 이상)",
        "2. 위치 편의성\n(이동시간 20% 미만)",
        "3. 적정 온도\n(18~28℃ 냉난방 기능)",
        "4. 적정 습도\n(50~55% 습도 조절)",
        "5. 적절한 밝기\n(100~200Lux 조명)",
        "6. 환기 및 비품\n(창문, 의자, 정수기)"
    ]
    
    # 4행 3열 표 생성 (한 줄에 사진 3장씩)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    for i in range(6):
        col_idx = i % 3
        # 0~2번 항목은 0행(사진), 1행(설명) / 3~5번 항목은 2행(사진), 3행(설명)
        if i < 3:
            row_photo = table.rows[0]
            row_desc = table.rows[1]
        else:
            row_photo = table.rows[2]
            row_desc = table.rows[3]
            
        # 사진 칸 설정
        cell_photo = row_photo.cells[col_idx]
        cell_photo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_photo = cell_photo.paragraphs[0]
        p_photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        img_path = image_paths[i]
        
        if img_path and os.path.exists(img_path):
            run = p_photo.add_run()
            # 세로로 긴 사진이 들어와도 다음 장으로 넘어가지 않도록 '높이'를 고정 (2.4인치)
            run.add_picture(img_path, height=Inches(2.4))
        else:
            run = p_photo.add_run(f"\n[사진 {i+1}]\n선택 안됨\n")
            run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
            
        # 설명 칸 설정
        cell_desc = row_desc.cells[col_idx]
        cell_desc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_desc = cell_desc.paragraphs[0]
        p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run_desc = p_desc.add_run(items[i])
        run_desc.font.size = Pt(10)
        run_desc.font.bold = True
        
    doc.save(output_filename)

class SafetyFormApp:
    def __init__(self, root):
        self.root = root
        self.root.title("휴게시설 점검 사진대장 생성 프로그램")
        self.root.geometry("600x450")
        
        self.image_paths = [""] * 6
        self.labels = []
        
        items_short = [
            "1. 크기 적정성",
            "2. 위치 편의성",
            "3. 적정 온도",
            "4. 적정 습도",
            "5. 적절한 밝기",
            "6. 환기 및 비품"
        ]
        
        tk.Label(root, text="각 항목의 [사진 찾기] 버튼을 눌러 점검 사진을 등록하세요.", pady=10, font=("맑은 고딕", 11, "bold")).grid(row=0, column=0, columnspan=3)
        
        for i in range(6):
            tk.Label(root, text=items_short[i], width=15, anchor="w", font=("맑은 고딕", 10)).grid(row=i+1, column=0, padx=10, pady=10)
            
            lbl_path = tk.Label(root, text="선택된 파일 없음", width=40, anchor="w", bg="white", relief="sunken")
            lbl_path.grid(row=i+1, column=1, padx=5)
            self.labels.append(lbl_path)
            
            btn_browse = tk.Button(root, text="사진 찾기", command=lambda idx=i: self.browse_file(idx))
            btn_browse.grid(row=i+1, column=2, padx=10)
            
        btn_generate = tk.Button(root, text="워드 문서 자동 생성하기", font=("맑은 고딕", 12, "bold"), bg="#4CAF50", fg="white", command=self.generate)
        btn_generate.grid(row=7, column=0, columnspan=3, pady=20, ipadx=20, ipady=10)
        
    def browse_file(self, idx):
        filename = filedialog.askopenfilename(
            title=f"사진 선택 - 항목 {idx+1}",
            filetypes=[("Image files", "*.jpg *.jpeg *.png *.bmp")]
        )
        if filename:
            self.image_paths[idx] = filename
            self.labels[idx].config(text=os.path.basename(filename))
            
    def generate(self):
        output_file = "휴게시설_사진대장_최종.docx"
        try:
            create_photo_log(self.image_paths, output_file)
            messagebox.showinfo("성공", f"보고서가 성공적으로 생성되었습니다!\n\n저장위치: {os.path.abspath(output_file)}")
            os.startfile(output_file)
        except Exception as e:
            messagebox.showerror("오류 발생", f"문서 생성 중 오류가 발생했습니다.\n{str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = SafetyFormApp(root)
    root.mainloop()
