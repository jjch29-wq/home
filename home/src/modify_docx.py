import docx

file_path = r'C:\Users\-\OneDrive\바탕 화면\PAUT_SIS-H-264 (2024ed)KI rev.0(최종)수정.docx'
doc = docx.Document(file_path)

para = doc.paragraphs[294]

eng_text = '\nIn addition, an encoder calibration check (Cal-In and Cal-Out) shall be performed at the beginning and end of each examination or shift to verify its accuracy as per RLNG-000-MT-SP-6313.'
kor_text = '\n(또한, RLNG-000-MT-SP-6313 규격에 따라 장비의 정확성을 검증하기 위해 매 검사 혹은 교대 근무의 시작(Cal-In)과 종료(Cal-Out) 시점에 엔코더 교정 확인이 수행되어야 한다.)'

run1 = para.add_run(eng_text)
run2 = para.add_run(kor_text)

output_path = r'C:\Users\-\OneDrive\바탕 화면\PAUT_SIS-H-264 (2024ed)KI rev.0(최종)_코멘트반영.docx'
doc.save(output_path)
print(f"Successfully saved to:\n{output_path}")
