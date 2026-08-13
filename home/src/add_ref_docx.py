import docx

file_path = r'C:\Users\-\OneDrive\바탕 화면\PAUT_SIS-H-264 (2024ed)KI rev.0(최종)_코멘트반영.docx'
doc = docx.Document(file_path)

for i, p in enumerate(doc.paragraphs):
    if 'RLNG-000-MT-SP-6302' in p.text:
        # We found (8) RLNG-000-MT-SP-6302
        new_p = p.insert_paragraph_before('  (9) RLNG-000-MT-SP-6313 : Specification for Phased Array Ultrasonic Testing (PAUT)')
        
        # move new_p after p by inserting before the next paragraph
        next_p = doc.paragraphs[i+1]
        new_p2 = next_p.insert_paragraph_before('  (9) RLNG-000-MT-SP-6313 : Specification for Phased Array Ultrasonic Testing (PAUT)')
        
        new_p2.style = p.style
        
        # Delete the first incorrectly placed paragraph
        p._element.getparent().remove(new_p._element)
        break

doc.save(file_path)
print("Added Reference 9 successfully.")
