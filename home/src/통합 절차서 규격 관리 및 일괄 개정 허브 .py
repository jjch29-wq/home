import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import json
import os
import re

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from PIL import Image, ImageTk
except ImportError:
    Image = None
    ImageTk = None

try:
    from tkinterweb import HtmlFrame
except ImportError:
    HtmlFrame = None

try:
    import mammoth
except ImportError:
    mammoth = None

DB_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "codebook_db.json")

# =========================================================
# 1. DatabaseManager: DB 로딩, 저장, 내보내기, 규격 추가 등
# =========================================================
class DatabaseManager:
    def __init__(self, db_path):
        self.db_path = db_path
        self.data = []
        self.load_data()

    def load_data(self):
        if os.path.exists(self.db_path):
            try:
                with open(self.db_path, 'r', encoding='utf-8') as f:
                    self.data = json.load(f)
            except Exception as e:
                messagebox.showerror("로딩 오류", f"DB 파일을 불러오지 못했습니다:\n{e}")
                self.data = []
        else:
            self.data = [
                {"category": "Article 2 (RT - 방사선투과)", "find": "ASME Sec.V Art.2", "replace": "ISO 17636-1", "details": "RT 절차서 관련 규격 예시"},
                {"category": "Article 4 (UT/PAUT - 초음파탐상)", "find": "ASME Sec.V Art.4", "replace": "ISO 11666", "details": "UT 절차서 관련 규격 예시"},
                {"category": "Article 7 (MT - 자분탐상)", "find": "ASME Sec.V Art.7", "replace": "ISO 17638", "details": "MT 절차서 관련 규격 예시"},
                {"category": "Article 6 (PT - 침투탐상)", "find": "ASME Sec.V Art.6", "replace": "ISO 3452-1", "details": "PT 절차서 관련 규격 예시"},
                {"category": "PMI (재질분석 - API/ASTM 등)", "find": "API RP 578", "replace": "ASTM E1476", "details": "PMI 절차서 관련 규격 예시"},
                {"category": "공통 (프로젝트/용어/기타)", "find": "기존프로젝트명", "replace": "가산~가평 천연가스 공급시설", "details": "문서 내 일괄 수정될 프로젝트명"}
            ]

    def save_data(self):
        try:
            with open(self.db_path, 'w', encoding='utf-8') as f:
                json.dump(self.data, f, ensure_ascii=False, indent=4)
        except Exception as e:
            messagebox.showerror("저장 오류", f"DB 파일을 저장하지 못했습니다:\n{e}")

# =========================================================
# 2. DocumentProcessor: 문서 파일 변환 엔진
# =========================================================
class DocumentProcessor:
    @staticmethod
    def replace_text_in_paragraph(paragraph, find_text, replace_text):
        if find_text in paragraph.text:
            for run in paragraph.runs:
                if find_text in run.text:
                    run.text = run.text.replace(find_text, replace_text)
                    return
            inline = paragraph.runs
            if not inline: return
            text = paragraph.text.replace(find_text, replace_text)
            for i in range(len(inline)): inline[i].text = ''
            inline[0].text = text

    @staticmethod
    def process_docx(input_file, output_file, replacements):
        try:
            import docx
            doc = docx.Document(input_file)
            
            def process_paragraphs(paragraphs):
                for paragraph in paragraphs:
                    for f_text, r_text in replacements:
                        DocumentProcessor.replace_text_in_paragraph(paragraph, f_text, r_text)
                        
            def process_tables(tables):
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            process_paragraphs(cell.paragraphs)
            
            # 본문 처리
            process_paragraphs(doc.paragraphs)
            process_tables(doc.tables)
            
            # 머릿글 및 바닥글 처리 (첫 페이지, 짝수 페이지 등 모든 옵션 포함)
            for section in doc.sections:
                for header in [section.header, section.first_page_header, section.even_page_header]:
                    if header:
                        process_paragraphs(header.paragraphs)
                        process_tables(header.tables)
                for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                    if footer:
                        process_paragraphs(footer.paragraphs)
                        process_tables(footer.tables)
                
            doc.save(output_file)
        except Exception as e:
            raise Exception(f"docx 처리 실패: {e}")

    @staticmethod
    def process_xlsx(input_file, output_file, replacements):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(input_file)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            new_val = cell.value
                            for f_text, r_text in replacements:
                                new_val = new_val.replace(f_text, r_text)
                            if new_val != cell.value:
                                cell.value = new_val
            wb.save(output_file)
        except Exception as e:
            raise Exception(f"xlsx 처리 실패: {e}")

    @staticmethod
    def process_hwp(input_file, output_file, replacements):
        try:
            import win32com.client as win32
            hwp = win32.gencache.EnsureDispatch("HWPFrame.HwpObject")
            hwp.RegisterModule("FilePathCheckDLL", "SecurityModule")
            hwp.XHwpWindows.Item(0).Visible = False
            
            if not hwp.Open(input_file):
                hwp.Quit()
                raise Exception("HWP 파일을 여는데 실패했습니다.")
                
            hwp.HAction.GetDefault("AllReplace", hwp.HParameterSet.HFindReplace.HSet)
            hwp.HParameterSet.HFindReplace.IgnoreMessage = 1
            
            for f_text, r_text in replacements:
                hwp.HParameterSet.HFindReplace.FindString = f_text
                hwp.HParameterSet.HFindReplace.ReplaceString = r_text
                hwp.HParameterSet.HFindReplace.Direction = hwp.FindDir("AllDoc")
                hwp.HAction.Execute("AllReplace", hwp.HParameterSet.HFindReplace.HSet)
                
            hwp.SaveAs(output_file)
            hwp.Quit()
        except Exception as e:
            raise Exception(f"hwp 처리 실패: {e}")

    @staticmethod
    def process_txt(input_file, output_file, replacements):
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                content = f.read()
            enc = 'utf-8'
        except:
            with open(input_file, 'r', encoding='euc-kr') as f:
                content = f.read()
            enc = 'euc-kr'
            
        for f_text, r_text in replacements:
            content = content.replace(f_text, r_text)
            
        with open(output_file, 'w', encoding=enc) as f:
            f.write(content)

    @staticmethod
    def process_doc(input_file, output_file, replacements):
        import os
        try:
            import win32com.client as win32
            word = win32.gencache.EnsureDispatch("Word.Application")
            word.Visible = False
            
            abs_input = os.path.abspath(input_file)
            abs_output = os.path.abspath(output_file)
            
            doc = word.Documents.Open(abs_input)
            
            # 모든 영역(본문, 머릿글, 바닥글 등) 순회하며 찾기 및 바꾸기 (Replace=2 는 모두 바꾸기)
            for f_text, r_text in replacements:
                for story in doc.StoryRanges:
                    story.Find.Execute(FindText=f_text, ReplaceWith=r_text, Replace=2)
                    while story.NextStoryRange:
                        story = story.NextStoryRange
                        story.Find.Execute(FindText=f_text, ReplaceWith=r_text, Replace=2)
                
            doc.SaveAs(abs_output)
            doc.Close()
            word.Quit()
        except Exception as e:
            try:
                word.Quit()
            except:
                pass
            raise Exception(f"doc 처리 오류: {e}")

    @staticmethod
    def process_xls(input_file, output_file, replacements):
        import os
        try:
            import win32com.client as win32
            excel = win32.gencache.EnsureDispatch("Excel.Application")
            excel.Visible = False
            excel.DisplayAlerts = False
            
            abs_input = os.path.abspath(input_file)
            abs_output = os.path.abspath(output_file)
            
            wb = excel.Workbooks.Open(abs_input)
            
            for sheet in wb.Worksheets:
                for f_text, r_text in replacements:
                    # LookAt=2 (xlPart: 부분 일치)
                    sheet.Cells.Replace(What=f_text, Replacement=r_text, LookAt=2, SearchOrder=1, MatchCase=False)
                    
            wb.SaveAs(abs_output)
            wb.Close(SaveChanges=False)
            excel.Quit()
        except Exception as e:
            try:
                excel.Quit()
            except:
                pass
            raise Exception(f"xls 처리 오류: {e}")

    @staticmethod
    def process_single_document(input_file, output_file, rules):
        ext_lower = os.path.splitext(input_file)[1].lower()
        if ext_lower == '.docx':
            DocumentProcessor.process_docx(input_file, output_file, rules)
        elif ext_lower == '.doc':
            DocumentProcessor.process_doc(input_file, output_file, rules)
        elif ext_lower == '.xlsx':
            DocumentProcessor.process_xlsx(input_file, output_file, rules)
        elif ext_lower == '.xls':
            DocumentProcessor.process_xls(input_file, output_file, rules)
        elif ext_lower in ['.hwp', '.hwpx']:
            DocumentProcessor.process_hwp(input_file, output_file, rules)
        elif ext_lower == '.txt':
            DocumentProcessor.process_txt(input_file, output_file, rules)
        else:
            raise ValueError(f"지원하지 않는 파일 형식입니다: {ext_lower}")


# =========================================================
# 3. MacroController: 운영체제 및 외부 프로그램 제어
# =========================================================
class MacroController:
    @staticmethod
    def sync_search_to_external_app(root, filepath, search_query=""):
        import os, time, threading
        import win32gui, win32con
        
        base_name = os.path.basename(filepath)
        name_without_ext = os.path.splitext(base_name)[0]
        
        def _macro_thread():
            for _ in range(10):
                time.sleep(0.5)
                found_hwnd = [0]
                def callback(hwnd, _):
                    if win32gui.IsWindowVisible(hwnd):
                        title = win32gui.GetWindowText(hwnd)
                        # 엄격한 매칭: PAUT절차서가 PAUT절차서_rev를 잡지 않도록 함
                        if title.startswith(base_name) or title.startswith(name_without_ext + " -") or title.startswith(name_without_ext + " [") or title == name_without_ext:
                            found_hwnd[0] = hwnd
                win32gui.EnumWindows(callback, None)
                
                if found_hwnd[0] != 0:
                    hwnd = found_hwnd[0]
                    screen_width = root.winfo_screenwidth()
                    screen_height = root.winfo_screenheight()
                    win32gui.ShowWindow(hwnd, win32con.SW_RESTORE)
                    win32gui.MoveWindow(hwnd, screen_width//2, 0, screen_width//2, screen_height-40, True)
                    
                    if search_query:
                        try: win32gui.SetForegroundWindow(hwnd)
                        except: pass
                        time.sleep(0.3)
                        try:
                            import win32com.client
                            shell = win32com.client.Dispatch("WScript.Shell")
                            shell.SendKeys("^f")
                            time.sleep(0.3)
                            shell.SendKeys("^v")
                            time.sleep(0.2)
                            shell.SendKeys("{ENTER}")
                        except Exception as e:
                            print("매크로 전송 오류:", e)
                    break
                    
        if search_query:
            root.clipboard_clear()
            root.clipboard_append(search_query)
            root.update()
        threading.Thread(target=_macro_thread, daemon=True).start()


# =========================================================
# 4. CodebookApp: 메인 UI 컨트롤러 (View 담당)
# =========================================================
class CodebookApp:
    def __init__(self, root):
        self.root = root
        self.root.title("🔥 통합 절차서 규격 관리 및 일괄 개정 허브 🔥")
        self.root.geometry("850x650")
        
        style = ttk.Style()
        style.theme_use('clam')
        
        # 1. 의존성 생성 (Managers & Controllers)
        self.db_manager = DatabaseManager(DB_FILE)
        
        # 2. 탭(Notebook) 구성
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill="both", expand=True)
        
        self.tab_viewer = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_viewer, text="📖 절차서 뷰어 및 추출")
        
        self.tab_code = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_code, text="📚 규격(코드) 관리 DB")
        
        self.tab_batch = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_batch, text="✍️ 다중 일괄 변환 (프리셋)")
        
        # 3. 상태 관리 변수
        self.current_doc_text = ""
        self.current_filepath = ""
        self.current_search_index = 0
        self.last_search_query = ""
        
        # 4. UI 생성 및 초기화
        self.create_viewer_widgets()
        self.create_widgets()
        self.create_batch_widgets()

        self.refresh_list()
        
    def create_widgets(self):
        # 상단 검색 및 필터 프레임
        search_frame = ttk.LabelFrame(self.tab_code, text="🔍 검색 및 필터")
        
        ttk.Label(search_frame, text="카테고리:").pack(side="left", padx=5, pady=5)
        self.combo_filter_cat = ttk.Combobox(search_frame, state="readonly", width=15)
        self.combo_filter_cat.pack(side="left", padx=5, pady=5)
        self.combo_filter_cat.bind("<<ComboboxSelected>>", lambda e: self.refresh_list())
        
        ttk.Label(search_frame, text="검색어:").pack(side="left", padx=(15, 5), pady=5)
        self.entry_search = ttk.Entry(search_frame, width=30)
        self.entry_search.pack(side="left", padx=5, pady=5)
        self.entry_search.bind("<KeyRelease>", lambda e: self.refresh_list())
        
        ttk.Button(search_frame, text="초기화", command=self.reset_filter).pack(side="left", padx=10, pady=5)

        # 메인 리스트 및 상세 보기 팬(PanedWindow)
        paned = ttk.PanedWindow(self.tab_code, orient=tk.HORIZONTAL)
        
        list_frame = ttk.Frame(paned)
        paned.add(list_frame, weight=5)
        
        detail_frame = ttk.LabelFrame(paned, text="📖 상세 내용 보기")
        paned.add(detail_frame, weight=3)
        
        self.detail_text = tk.Text(detail_frame, wrap="word", font=("맑은 고딕", 11))
        
        detail_scrollbar = ttk.Scrollbar(detail_frame, orient="vertical", command=self.detail_text.yview)
        detail_scrollbar.pack(side="right", fill="y")
        self.detail_text.pack(side="left", fill="both", expand=True, padx=5, pady=5)
        self.detail_text.config(yscrollcommand=detail_scrollbar.set)
        self.detail_text.config(state="disabled")
        
        columns = ("category", "find", "replace")
        self.tree = ttk.Treeview(list_frame, columns=columns, show="headings", height=15)
        self.tree.heading("category", text="분류 (카테고리)")
        self.tree.heading("find", text="찾을 내용 (기존 코드/문구)")
        self.tree.heading("replace", text="바꿀 내용 (새로운 코드/문구)")
        
        self.tree.column("category", width=150, anchor="center")
        self.tree.column("find", width=300)
        self.tree.column("replace", width=300)
        
        scrollbar = ttk.Scrollbar(list_frame, orient="vertical", command=self.tree.yview)
        scrollbar.pack(side="right", fill="y")
        self.tree.pack(side="left", fill="both", expand=True)
        self.tree.config(yscrollcommand=scrollbar.set)
        
        self.tree.bind("<<TreeviewSelect>>", self.on_tree_select)
        
        # 버튼 프레임
        btn_frame = ttk.Frame(self.tab_code)
        ttk.Button(btn_frame, text="📋 찾을 내용 복사", command=lambda: self.copy_to_clipboard("find")).pack(side="left", padx=5)
        ttk.Button(btn_frame, text="📋 바꿀 내용 복사", command=lambda: self.copy_to_clipboard("replace")).pack(side="left", padx=5)
        ttk.Button(btn_frame, text="📤 현재 목록을 통합 허브용(JSON)으로 내보내기", command=self.export_preset).pack(side="right", padx=5)

        # 하단 입력 프레임
        input_frame = ttk.LabelFrame(self.tab_code, text="✍️ 코드 추가 및 수정")
        
        ttk.Label(input_frame, text="카테고리:").grid(row=0, column=0, padx=5, pady=10, sticky="e")
        self.combo_cat = ttk.Combobox(input_frame, width=15)
        self.combo_cat.grid(row=0, column=1, padx=5, pady=10, sticky="w")
        
        ttk.Label(input_frame, text="찾을 내용:").grid(row=0, column=2, padx=5, pady=10, sticky="e")
        self.entry_find = ttk.Entry(input_frame, width=25)
        self.entry_find.grid(row=0, column=3, padx=5, pady=10, sticky="w")
        
        ttk.Label(input_frame, text="바꿀 내용:").grid(row=0, column=4, padx=5, pady=10, sticky="e")
        self.entry_replace = ttk.Entry(input_frame, width=25)
        self.entry_replace.grid(row=0, column=5, padx=5, pady=10, sticky="w")
        
        ttk.Label(input_frame, text="코드 내용\n(상세 설명):").grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.text_details_input = tk.Text(input_frame, width=32, height=4, font=("맑은 고딕", 10))
        self.text_details_input.grid(row=1, column=1, columnspan=2, padx=5, pady=5, sticky="w")
        
        ttk.Label(input_frame, text="개정 내역\n(변경 사유):").grid(row=1, column=3, padx=5, pady=5, sticky="e")
        self.text_revision_input = tk.Text(input_frame, width=32, height=4, font=("맑은 고딕", 10))
        self.text_revision_input.grid(row=1, column=4, columnspan=2, padx=5, pady=5, sticky="w")
        
        action_frame = ttk.Frame(input_frame)
        action_frame.grid(row=2, column=0, columnspan=6, pady=10)
        
        ttk.Button(action_frame, text="✨ 추가", command=self.add_code, width=15).pack(side="left", padx=10)
        ttk.Button(action_frame, text="💾 수정", command=self.update_code, width=15).pack(side="left", padx=10)
        ttk.Button(action_frame, text="🗑️ 삭제", command=self.delete_code, width=15).pack(side="left", padx=10)

        # 화면 배치
        search_frame.pack(side="top", fill="x", padx=10, pady=10)
        input_frame.pack(side="bottom", fill="x", padx=10, pady=10)
        btn_frame.pack(side="bottom", fill="x", padx=10, pady=5)
        paned.pack(side="top", fill="both", expand=True, padx=10, pady=5)

    def update_categories(self):
        cats = sorted(list(set([d.get("category", "") for d in self.db_manager.data if d.get("category", "")])))
        self.combo_cat['values'] = cats
        self.combo_filter_cat['values'] = ["전체"] + cats
        if self.combo_filter_cat.get() not in ["전체"] + cats:
            self.combo_filter_cat.set("전체")

    def refresh_list(self):
        self.update_categories()
        for item in self.tree.get_children():
            self.tree.delete(item)
            
        filter_cat = self.combo_filter_cat.get()
        search_kw = self.entry_search.get().lower()
        
        for idx, row in enumerate(self.db_manager.data):
            cat = row.get("category", "")
            f_txt = row.get("find", "")
            r_txt = row.get("replace", "")
            
            if filter_cat and filter_cat != "전체" and cat != filter_cat:
                continue
            if search_kw:
                if search_kw not in cat.lower() and search_kw not in f_txt.lower() and search_kw not in r_txt.lower():
                    continue
            self.tree.insert("", "end", iid=str(idx), values=(cat, f_txt, r_txt))

    def reset_filter(self):
        self.combo_filter_cat.set("전체")
        self.entry_search.delete(0, tk.END)
        self.refresh_list()

    def on_tree_select(self, event):
        selected = self.tree.selection()
        if selected:
            idx = int(selected[0])
            row = self.db_manager.data[idx]
            
            self.combo_cat.set(row.get("category", ""))
            self.entry_find.delete(0, tk.END)
            self.entry_find.insert(0, row.get("find", ""))
            self.entry_replace.delete(0, tk.END)
            self.entry_replace.insert(0, row.get("replace", ""))
            
            self.text_details_input.delete("1.0", tk.END)
            self.text_details_input.insert("1.0", row.get("details", ""))
            
            self.text_revision_input.delete("1.0", tk.END)
            self.text_revision_input.insert("1.0", row.get("revision_note", ""))
            
            cat_text = row.get("category", "")
            f_text = row.get("find", "")
            r_text = row.get("replace", "")
            d_text = row.get("details", "")
            rev_text = row.get("revision_note", "")
            
            detail_content = f"■ 분류 (카테고리)\n{cat_text}\n\n"
            detail_content += f"■ 찾을 내용 (설명/기존문구)\n{f_text}\n\n"
            detail_content += f"■ 바꿀 내용 (적용할 규격/코드)\n{r_text}\n\n"
            detail_content += f"■ 개정 내역 및 변경 사유\n{rev_text if rev_text else '(등록된 개정 내역이 없습니다.)'}\n\n"
            detail_content += f"■ 실제 코드 내용 (상세 설명)\n{d_text if d_text else '(등록된 상세 내용이 없습니다.)'}"
            
            self.detail_text.config(state="normal")
            self.detail_text.delete("1.0", tk.END)
            self.detail_text.insert("1.0", detail_content)
            self.detail_text.config(state="disabled")

    def add_code(self):
        cat = self.combo_cat.get().strip()
        f_txt = self.entry_find.get().strip()
        r_txt = self.entry_replace.get().strip()
        d_txt = self.text_details_input.get("1.0", tk.END).strip()
        rev_txt = self.text_revision_input.get("1.0", tk.END).strip()
        
        if not f_txt:
            messagebox.showwarning("입력 오류", "찾을 내용을 입력해주세요.")
            return
            
        # Check for duplicates to merge
        merged = False
        for item in self.db_manager.data:
            if item.get("category") == cat and item.get("find") == f_txt:
                # Merge details if different and not empty
                if d_txt and d_txt not in item.get("details", ""):
                    if item.get("details"):
                        item["details"] += f"\n\n[추가 내용]\n{d_txt}"
                    else:
                        item["details"] = d_txt
                
                # Merge revision_note if different and not empty
                if rev_txt and rev_txt not in item.get("revision_note", ""):
                    if item.get("revision_note"):
                        item["revision_note"] += f"\n\n[추가 개정내역]\n{rev_txt}"
                    else:
                        item["revision_note"] = rev_txt
                
                if r_txt:
                    item["replace"] = r_txt
                
                merged = True
                break

        if not merged:
            self.db_manager.data.append({
                "category": cat,
                "find": f_txt,
                "replace": r_txt,
                "details": d_txt,
                "revision_note": rev_txt
            })
        self.db_manager.save_data()
        self.refresh_list()
        
        self.combo_cat.set("")
        self.entry_find.delete(0, tk.END)
        self.entry_replace.delete(0, tk.END)
        self.text_details_input.delete("1.0", tk.END)
        self.text_revision_input.delete("1.0", tk.END)
        
        self.detail_text.config(state="normal")
        self.detail_text.delete("1.0", tk.END)
        self.detail_text.config(state="disabled")
        
        self.entry_find.focus()

    def update_code(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("선택 오류", "수정할 코드를 위 목록에서 선택해주세요.")
            return
            
        idx = int(selected[0])
        cat = self.combo_cat.get().strip()
        f_txt = self.entry_find.get().strip()
        r_txt = self.entry_replace.get().strip()
        d_txt = self.text_details_input.get("1.0", tk.END).strip()
        rev_txt = self.text_revision_input.get("1.0", tk.END).strip()
        
        if not f_txt:
            messagebox.showwarning("입력 오류", "찾을 내용을 입력해주세요.")
            return
            
        self.db_manager.data[idx] = {
            "category": cat,
            "find": f_txt,
            "replace": r_txt,
            "details": d_txt,
            "revision_note": rev_txt
        }
        self.db_manager.save_data()
        self.refresh_list()
        
    def delete_code(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("선택 오류", "삭제할 코드를 위 목록에서 선택해주세요.")
            return
            
        if messagebox.askyesno("삭제 확인", "선택한 코드를 정말 삭제하시겠습니까?"):
            idxs = sorted([int(s) for s in selected], reverse=True)
            for idx in idxs:
                del self.db_manager.data[idx]
                
            self.db_manager.save_data()
            self.refresh_list()
            
            self.combo_cat.set("")
            self.entry_find.delete(0, tk.END)
            self.entry_replace.delete(0, tk.END)
            self.text_details_input.delete("1.0", tk.END)
            self.text_revision_input.delete("1.0", tk.END)
            
            self.detail_text.config(state="normal")
            self.detail_text.delete("1.0", tk.END)
            self.detail_text.config(state="disabled")

    def copy_to_clipboard(self, field="find"):
        selected = self.tree.selection()
        if not selected:
            messagebox.showinfo("안내", "복사할 항목을 먼저 선택해주세요.")
            return
        idx = int(selected[0])
        text = self.db_manager.data[idx].get(field, "")
        self.root.clipboard_clear()
        self.root.clipboard_append(text)
        self.root.update()
        messagebox.showinfo("복사 완료", f"클립보드에 복사되었습니다:\n\n{text}")

    def export_preset(self):
        items = self.tree.get_children()
        if not items:
            messagebox.showwarning("경고", "내보낼 항목이 없습니다.")
            return
            
        preset_data = []
        for item in items:
            val = self.tree.item(item, 'values')
            preset_data.append({"find": val[1], "replace": val[2]})
            
        filepath = filedialog.asksaveasfilename(
            title="절차서 수정 헬퍼용 단어 목록으로 내보내기",
            defaultextension=".json",
            filetypes=[("JSON 파일", "*.json")],
            initialfile="내보낸_코드세트.json"
        )
        
        if filepath:
            try:
                with open(filepath, 'w', encoding='utf-8') as f:
                    json.dump(preset_data, f, ensure_ascii=False, indent=4)
                messagebox.showinfo("내보내기 완료", f"성공적으로 내보냈습니다!\n\n저장 경로: {filepath}")
            except Exception as e:
                messagebox.showerror("오류", f"저장 중 오류 발생:\n{e}")

    def create_viewer_widgets(self):
        ctrl_frame = ttk.Frame(self.tab_viewer)
        ctrl_frame.pack(side="top", fill="x", padx=10, pady=5)
        
        ttk.Button(ctrl_frame, text="📂 문서 열기 (Word, Excel, HWP, TXT)", command=self.load_document).pack(side="left", padx=5)
        ttk.Button(ctrl_frame, text="✨ 2.0 Reference 추출", command=self.extract_references).pack(side="left", padx=10)
        
        self.btn_edit_doc = ttk.Button(ctrl_frame, text="📝 원본 프로그램으로 열어서 직접 수정하기", command=self.open_current_document, state="disabled")
        self.btn_edit_doc.pack(side="left", padx=10)
        
        self.btn_apply_db = ttk.Button(ctrl_frame, text="✨ 전체 규격 코드 DB 일괄 적용", command=self.apply_db_to_current, state="disabled")
        self.btn_apply_db.pack(side="left", padx=10)
        
        self.btn_popup_db = ttk.Button(ctrl_frame, text="🔍 DB 검색/수정 (팝업)", command=self.open_db_popup)
        self.btn_popup_db.pack(side="right", padx=10)
        
        if HtmlFrame is None or mammoth is None:
            ttk.Label(ctrl_frame, text="⚠️ 필수 모듈(tkinterweb, mammoth)이 부족합니다.", foreground="red").pack(side="right", padx=10)
            
        search_frame = ttk.Frame(self.tab_viewer)
        search_frame.pack(side="top", fill="x", padx=10, pady=0)
        
        ttk.Label(search_frame, text="🔍 찾을 단어:").pack(side="left", padx=5)
        self.entry_viewer_search = ttk.Entry(search_frame, width=20)
        self.entry_viewer_search.pack(side="left", padx=5)
        self.entry_viewer_search.bind("<Return>", lambda e: self.search_in_viewer(direction=1))
        
        ttk.Button(search_frame, text="◀ 이전", command=lambda: self.search_in_viewer(direction=-1)).pack(side="left", padx=2)
        ttk.Button(search_frame, text="다음 ▶", command=lambda: self.search_in_viewer(direction=1)).pack(side="left", padx=2)
        ttk.Button(search_frame, text="초기화", command=self.reset_viewer_search).pack(side="left", padx=2)
        
        self.lbl_viewer_search_result = ttk.Label(search_frame, text="")
        self.lbl_viewer_search_result.pack(side="left", padx=5)
        
        ttk.Label(search_frame, text=" | ").pack(side="left", padx=2)
        
        ttk.Label(search_frame, text="➡ 바꿀 단어:").pack(side="left", padx=5)
        self.entry_viewer_replace = ttk.Entry(search_frame, width=20)
        self.entry_viewer_replace.pack(side="left", padx=5)
        
        self.btn_quick_replace = ttk.Button(search_frame, text="✨ 즉시 고치기", command=self.quick_replace_viewer_text, state="disabled")
        self.btn_quick_replace.pack(side="left", padx=5)
        
        viewer_frame = ttk.Frame(self.tab_viewer)
        viewer_frame.pack(side="top", fill="both", expand=True, padx=10, pady=5)
        
        if HtmlFrame:
            self.html_viewer = HtmlFrame(viewer_frame, messages_enabled=False)
            self.html_viewer.pack(fill="both", expand=True)
        else:
            self.html_viewer = None

        # Status label
        self.lbl_viewer_status = ttk.Label(self.tab_viewer, text="준비됨", foreground="gray")
        self.lbl_viewer_status.pack(side="bottom", fill="x", padx=10, pady=5)

    def open_db_popup(self):
        """뷰어에서 즉시 DB를 검색/수정할 수 있는 팝업창 열기"""
        popup = tk.Toplevel(self.root)
        popup.title("규격 DB 빠른 검색 및 수정")
        popup.geometry("900x600")
        popup.attributes('-topmost', True)
        
        # 1. Search Frame
        search_f = ttk.Frame(popup, padding=10)
        search_f.pack(fill='x')
        ttk.Label(search_f, text="검색 (카테고리, 찾을값 등):").pack(side='left', padx=5)
        ent_search = ttk.Entry(search_f, width=40)
        ent_search.pack(side='left', padx=5)
        
        # 2. Treeview
        tree_f = ttk.Frame(popup, padding=10)
        tree_f.pack(fill='both', expand=True)
        columns = ("ID", "Category", "Find", "Replace", "Details")
        tree = ttk.Treeview(tree_f, columns=columns, show="headings", height=12)
        for col in columns: tree.heading(col, text=col)
        tree.column("ID", width=40, anchor="center")
        tree.column("Category", width=120)
        tree.column("Find", width=250)
        tree.column("Replace", width=250)
        tree.column("Details", width=180)
        scroll = ttk.Scrollbar(tree_f, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=scroll.set)
        tree.pack(side='left', fill='both', expand=True)
        scroll.pack(side='right', fill='y')
        
        # 3. Edit Frame
        edit_f = ttk.LabelFrame(popup, text="항목 편집", padding=10)
        edit_f.pack(fill='x', padx=10, pady=10)
        
        ttk.Label(edit_f, text="카테고리:").grid(row=0, column=0, padx=2, pady=5, sticky='e')
        c_cat = ttk.Combobox(edit_f, width=15)
        c_cat.grid(row=0, column=1, padx=2, pady=5, sticky='w')
        
        ttk.Label(edit_f, text="찾을 값:").grid(row=0, column=2, padx=2, pady=5, sticky='e')
        e_find = ttk.Entry(edit_f, width=25)
        e_find.grid(row=0, column=3, padx=2, pady=5, sticky='w')
        
        ttk.Label(edit_f, text="바꿀 값:").grid(row=0, column=4, padx=2, pady=5, sticky='e')
        e_rep = ttk.Entry(edit_f, width=25)
        e_rep.grid(row=0, column=5, padx=2, pady=5, sticky='w')
        
        ttk.Label(edit_f, text="비고:").grid(row=1, column=0, padx=2, pady=5, sticky='e')
        e_note = ttk.Entry(edit_f, width=70)
        e_note.grid(row=1, column=1, columnspan=5, sticky='w', padx=2, pady=5)
        
        # Functions
        def refresh_popup(e=None):
            for i in tree.get_children(): tree.delete(i)
            q = ent_search.get().lower()
            cats = set()
            for idx, item in enumerate(self.db_manager.data):
                cat = item.get('category', '')
                f = item.get('find', '')
                r = item.get('replace', '')
                d = item.get('details', '')
                cats.add(cat)
                if q in cat.lower() or q in f.lower() or q in r.lower() or q in d.lower():
                    tree.insert("", "end", iid=str(idx), values=(idx, cat, f, r, d))
            c_cat['values'] = sorted(list(cats))
            
        def on_select(e):
            sel = tree.selection()
            if not sel: return
            idx = int(sel[0])
            item = self.db_manager.data[idx]
            c_cat.set(item.get('category', ''))
            e_find.delete(0, tk.END); e_find.insert(0, item.get('find', ''))
            e_rep.delete(0, tk.END); e_rep.insert(0, item.get('replace', ''))
            e_note.delete(0, tk.END); e_note.insert(0, item.get('details', ''))
            
        def do_add():
            if not e_find.get(): return messagebox.showwarning("경고", "찾을 값을 입력하세요.", parent=popup)
            self.db_manager.data.append({"category": c_cat.get(), "find": e_find.get(), "replace": e_rep.get(), "details": e_note.get()})
            self.db_manager.save_data()
            refresh_popup(); self.refresh_list(); self.update_categories()
            messagebox.showinfo("추가", "추가되었습니다.", parent=popup)
            
        def do_edit():
            sel = tree.selection()
            if not sel: return messagebox.showwarning("경고", "수정할 항목을 선택하세요.", parent=popup)
            self.db_manager.data[int(sel[0])] = {"category": c_cat.get(), "find": e_find.get(), "replace": e_rep.get(), "details": e_note.get()}
            self.db_manager.save_data()
            refresh_popup(); self.refresh_list(); self.update_categories()
            messagebox.showinfo("수정", "수정되었습니다.", parent=popup)
            
        def do_del():
            sel = tree.selection()
            if not sel: return messagebox.showwarning("경고", "삭제할 항목을 선택하세요.", parent=popup)
            if messagebox.askyesno("삭제", "정말 삭제하시겠습니까?", parent=popup):
                del self.db_manager.data[int(sel[0])]
                self.db_manager.save_data()
                refresh_popup(); self.refresh_list(); self.update_categories()
                
        # Bindings & Buttons
        ent_search.bind("<KeyRelease>", refresh_popup)
        tree.bind("<<TreeviewSelect>>", on_select)
        
        btn_f = ttk.Frame(popup, padding=10)
        btn_f.pack(fill='x')
        ttk.Button(btn_f, text="➕ 새 항목 추가", command=do_add).pack(side='left', padx=5)
        ttk.Button(btn_f, text="✏️ 선택 항목 수정", command=do_edit).pack(side='left', padx=5)
        ttk.Button(btn_f, text="🗑️ 선택 항목 삭제", command=do_del).pack(side='left', padx=5)
        ttk.Button(btn_f, text="닫기", command=popup.destroy).pack(side='right', padx=5)
        
        refresh_popup()

    def load_document(self):
        if self.html_viewer is None or mammoth is None:
            messagebox.showerror("오류", "문서를 렌더링하기 위한 필수 라이브러리가 없습니다.")
            return
            
        filepath = filedialog.askopenfilename(
            title="절차서 문서 열기",
            filetypes=[
                ("모든 지원 파일", "*.docx *.xlsx *.xls *.hwp *.hwpx *.txt *.doc *.pdf"),
                ("워드 파일", "*.docx *.doc"),
                ("엑셀 파일", "*.xlsx *.xls"),
                ("한글 파일", "*.hwp *.hwpx"),
                ("텍스트 파일", "*.txt"),
                ("PDF 파일", "*.pdf"),
                ("모든 파일", "*.*")
            ]
        )
        if filepath:
            self._load_document_by_path(filepath)

    def _load_document_by_path(self, filepath):
        ext = filepath.lower().split('.')[-1]
        try:
            if ext == 'docx':
                with open(filepath, "rb") as docx_file:
                    result = mammoth.convert_to_html(docx_file)
                    html = result.value
                    docx_file.seek(0)
                    text_result = mammoth.extract_raw_text(docx_file)
                    self.current_doc_text = text_result.value
                    
                styled_html = f"""
                <html>
                <head>
                <style>
                    body {{ font-family: 'Malgun Gothic', sans-serif; line-height: 1.6; padding: 20px; }}
                    table {{ border-collapse: collapse; width: 100%; margin-bottom: 20px; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background-color: #f2f2f2; }}
                    img {{ max-width: 100%; height: auto; }}
                </style>
                </head>
                <body>{html}</body>
                </html>
                """
                self.html_viewer.load_html(styled_html)
                
            elif ext == 'txt':
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        text_content = f.read()
                except:
                    with open(filepath, 'r', encoding='euc-kr') as f:
                        text_content = f.read()
                self.current_doc_text = text_content
                html_content = text_content.replace('\n', '<br>')
                styled_html = f"<html><body style=\"font-family: 'Malgun Gothic', sans-serif; padding: 20px; line-height: 1.6;\">{html_content}</body></html>"
                self.html_viewer.load_html(styled_html)
                
            elif ext in ['xlsx', 'xls', 'hwp', 'hwpx', 'doc', 'pdf']:
                self.current_doc_text = ""
                styled_html = f"""
                <html>
                <body style="font-family: 'Malgun Gothic', sans-serif; padding: 40px; text-align: center; color: #555; line-height: 1.8;">
                    <h2>{ext.upper()} 파일이 로드되었습니다.</h2>
                    <p>현재 뷰어 화면의 미리보기는 최신 워드(.docx) 및 텍스트(.txt) 문서만 지원합니다.</p>
                    <p>하지만 상단의 <b>[📝 원본 프로그램으로 열어서 직접 수정하기]</b> 버튼을 누르시면<br>
                    정상적으로 해당 문서를 여실 수 있습니다.</p>
                </body>
                </html>
                """
                self.html_viewer.load_html(styled_html)
            else:
                messagebox.showinfo("안내", "지원하지 않는 파일 형식입니다.")
                return

            self.notebook.update()
            self.current_filepath = filepath
            self.btn_edit_doc.config(state="normal")
            self.btn_apply_db.config(state="normal")
            if hasattr(self, 'btn_quick_replace'):
                self.btn_quick_replace.config(state="normal")
            
        except Exception as e:
            messagebox.showerror("오류", f"문서를 여는 중 오류가 발생했습니다:\\n{e}")

    def apply_db_to_current(self):
        if not self.current_filepath:
            messagebox.showwarning("경고", "먼저 문서를 열어주세요.")
            return
            
        rules = [(d["find"], d["replace"]) for d in self.db_manager.data if d.get("find") and d.get("replace")]
        if not rules:
            messagebox.showwarning("경고", "코드 DB에 바꿀 내용(Replace)이 설정된 규격이 하나도 없습니다.")
            return
            
        if not messagebox.askyesno("일괄 적용 확인", f"현재 문서에 DB의 변환 규칙 {len(rules)}개를 모두 적용하시겠습니까?"):
            return
            
        try:
            self.root.config(cursor="wait")
            self.root.update()
            
            dir_name = os.path.dirname(self.current_filepath)
            base_name = os.path.basename(self.current_filepath)
            name, ext = os.path.splitext(base_name)
            output_filepath = os.path.join(dir_name, f"{name}_수정본{ext}")
            
            try:
                DocumentProcessor.process_single_document(self.current_filepath, output_filepath, rules)
            except Exception as fe:
                self.root.config(cursor="")
                messagebox.showinfo("안내", str(fe))
                return
            
            self._load_document_by_path(output_filepath)
            self.root.config(cursor="")
            messagebox.showinfo("적용 완료", f"적용 완료! 저장 경로: {output_filepath}")
        except Exception as e:
            self.root.config(cursor="")
            messagebox.showerror("오류", f"문서 변환 중 오류가 발생했습니다:\\n{e}")

    def quick_replace_viewer_text(self):
        if not self.current_filepath:
            return
            
        find_txt = self.entry_viewer_search.get().strip()
        replace_txt = self.entry_viewer_replace.get().strip()
        if not find_txt:
            messagebox.showwarning("입력 오류", "찾을 단어를 입력해주세요.")
            return
            
        if not messagebox.askyesno("빠른 치환 확인", f"'{find_txt}' 단어를 '{replace_txt}'(으)로 즉시 변경하시겠습니까?"):
            return
            
        try:
            with open(self.current_filepath, 'a'):
                pass
        except PermissionError:
            messagebox.showwarning("문서 사용 중", "현재 이 문서가 Word, Excel, HWP 등 다른 프로그램에서 열려 있습니다!\n\n수정된 내용을 덮어쓰려면 열려 있는 원본 프로그램을 완전히 닫은 후 다시 [즉시 고치기]를 눌러주세요.\n(또는 열려있는 창에서 직접 찾기/바꾸기 단축키를 이용하세요.)")
            return
            
        try:
            self.root.config(cursor="wait")
            self.root.update()
            
            rules = [(find_txt, replace_txt)]
            import os, shutil
            dir_name = os.path.dirname(self.current_filepath)
            base_name = os.path.basename(self.current_filepath)
            name, ext = os.path.splitext(base_name)
            temp_filepath = os.path.join(dir_name, f"{name}_temp_replace{ext}")
            
            try:
                DocumentProcessor.process_single_document(self.current_filepath, temp_filepath, rules)
            except Exception as fe:
                self.root.config(cursor="")
                messagebox.showinfo("안내", str(fe))
                return
                
            try:
                os.remove(self.current_filepath)
                shutil.move(temp_filepath, self.current_filepath)
                final_path = self.current_filepath
            except Exception:
                final_path = os.path.join(dir_name, f"{name}_즉시수정본{ext}")
                shutil.move(temp_filepath, final_path)
                messagebox.showwarning("저장 안내", "원본 파일이 열려 있어 덮어쓰지 못했습니다.\\n대신 '_즉시수정본'으로 저장되었습니다.")
                
            self._load_document_by_path(final_path)
            self.root.config(cursor="")
            messagebox.showinfo("수정 완료", "성공적으로 수정되었습니다!")
        except Exception as e:
            self.root.config(cursor="")
            messagebox.showerror("오류", f"문서 수정 중 오류:\\n{e}")

    def open_current_document(self):
        if self.current_filepath:
            try:
                if self.root.state() == 'zoomed':
                    self.root.state('normal')
                self.root.update_idletasks()
                screen_width = self.root.winfo_screenwidth()
                screen_height = self.root.winfo_screenheight()
                self.root.geometry(f"{screen_width//2}x{screen_height-80}+0+0")
                
                os.startfile(self.current_filepath)
                search_term = self.entry_viewer_search.get().strip()
                MacroController.sync_search_to_external_app(self.root, self.current_filepath, search_term)
            except Exception as e:
                messagebox.showerror("오류", f"실행 오류:\\n{e}")

    def search_in_viewer(self, direction=1):
        if not self.html_viewer: return
        query = self.entry_viewer_search.get().strip()
        if not query:
            self.reset_viewer_search()
            return
            
        safe_query = re.escape(query)
        
        # Check total matches first without selecting
        try:
            matches_count = self.html_viewer.find_text(safe_query, select=0, ignore_case=True, highlight_all=True)
        except Exception:
            self.lbl_viewer_search_result.config(text="검색 오류", foreground="red")
            return
            
        if matches_count == 0:
            self.lbl_viewer_search_result.config(text="결과 없음", foreground="red")
            self.last_search_query = query
            return
            
        if query != self.last_search_query:
            self.last_search_query = query
            self.current_search_index = 1 if direction == 1 else matches_count
        else:
            self.current_search_index += direction
            if self.current_search_index > matches_count:
                self.current_search_index = 1
            elif self.current_search_index < 1:
                self.current_search_index = matches_count
                
        try:
            self.html_viewer.find_text(safe_query, select=self.current_search_index, ignore_case=True, highlight_all=True)
            self.lbl_viewer_search_result.config(text=f"{self.current_search_index} / {matches_count} 개", foreground="blue")
        except Exception:
            self.lbl_viewer_search_result.config(text="검색 오류", foreground="red")
            
        if self.current_filepath:
            MacroController.sync_search_to_external_app(self.root, self.current_filepath, query)

    def reset_viewer_search(self):
        self.entry_viewer_search.delete(0, tk.END)
        self.last_search_query = ""
        self.current_search_index = 0
        self.lbl_viewer_search_result.config(text="")
        if self.html_viewer:
            self.html_viewer.find_text("")

    def extract_references(self):
        if not self.current_doc_text:
            messagebox.showwarning("경고", "먼저 문서를 열어주세요.")
            return
            
        full_text = self.current_doc_text
        lines = full_text.split('\n')
        extracted_count = 0
        found_codes = set()
        
        # 문서 제목 및 앞부분 텍스트로 절차서 종류(RT, UT 등) 자동 판별
        head_text = ""
        if self.current_filepath:
            head_text += os.path.basename(self.current_filepath).lower() + " "
        head_text += "\n".join(lines[:50]).lower()
        
        doc_type_category = None
        if "paut" in head_text or "phased array" in head_text or "위상배열" in head_text or "ut" in head_text.split() or "초음파" in head_text or "ultrasonic" in head_text:
            doc_type_category = "Article 4 (UT/PAUT - 용접부 초음파)"
        elif "rt" in head_text.split() or "방사선" in head_text or "radiographic" in head_text:
            doc_type_category = "Article 2 (RT - 방사선투과)"
        elif "mt" in head_text.split() or "자분" in head_text or "magnetic" in head_text:
            doc_type_category = "Article 7 (MT - 자분탐상)"
        elif "pt" in head_text.split() or "침투" in head_text or "penetrant" in head_text:
            doc_type_category = "Article 6 (PT - 침투탐상)"
        elif "pmi" in head_text.split() or "재질분석" in head_text or "positive material" in head_text:
            doc_type_category = "PMI (재질분석 - API/ASTM 등)"
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line or len(line) < 4: continue
            
            pattern = r'\b(ASME\s+(?:Sec(?:tion|\.)?\s*[IVX]+(?:\s*Art(?:icle|\.)?\s*\d+(?:\s*(?:SE|SA|SB)-[\w\d\-]+)?)?(?:\s*Mandatory\s*Appendix\s*[IVX]+)?|B\s*\d+(?:\.\d+)?)|(?:ASME\s*)?BPV[C]?\s*Code\s*\d{2,4}(?:\s*Ed\.?)?|API\s*(?:Std|Spec)?\s*\d+[A-Z]?|AWS\s*[A-Z]\d+(?:\.\d+)?|ISO\s*\d+(?:-\d+)?|KS\s*[A-Z]\s*\d+(?:-\d+)?|ASTM\s*[A-Z]\s*\d+|EN\s*\d+(?:-\d+)?|ASNT\s*SNT-TC-1A(?:(?:\s*\(\d{4}\s*Ed\.?\))?)|KEPIC\s*[A-Z\d]+)'
            match = re.search(pattern, line, re.IGNORECASE)
            
            if match:
                code_name = match.group(1).strip()
                code_name_norm = " ".join(code_name.split())
                
                if code_name_norm.lower() in found_codes:
                    continue
                found_codes.add(code_name_norm.lower())
                
                details = f"[규격 명칭]\n{line}"
                if i + 1 < len(lines) and lines[i+1].strip() and not re.search(r'^\d+\.', lines[i+1]):
                    details += " " + lines[i+1].strip()
                    
                usages = []
                for usage_match in re.finditer(r'.{0,40}' + re.escape(code_name) + r'.{0,40}', full_text, re.IGNORECASE):
                    usage_text = usage_match.group(0).replace('\n', ' ').strip()
                    if usage_text not in usages and "Reference" not in usage_text:
                        usages.append(usage_text)
                
                if usages:
                    details += "\n\n[문서 내 검색된 사용 예시]\n"
                    for u in usages[:5]:
                        details += f"- ...{u}...\n"
                        
                existing_item = None
                for d in self.db_manager.data:
                    if code_name_norm.lower() in d.get("find", "").lower():
                        existing_item = d
                        break
                        
                if not existing_item:
                    prefix = code_name_norm.split()[0].upper()
                    if doc_type_category:
                        category_name = doc_type_category
                    elif prefix in ["ASME", "API", "AWS", "ISO", "KS", "ASTM", "EN", "KEPIC", "ASNT", "BPV", "BPVC"]:
                        category_name = f"{prefix} 규격"
                    else:
                        category_name = "규격 (자동추출)"
                        
                    self.db_manager.data.append({
                        "category": category_name,
                        "find": code_name_norm,
                        "replace": "",
                        "details": details
                    })
                    extracted_count += 1
                else:
                    old_details = existing_item.get("details", "")
                    new_usages = [u for u in usages[:5] if u not in old_details]
                    if new_usages:
                        if "[문서 내 검색된 사용 예시]" not in old_details:
                            old_details += "\n\n[문서 내 검색된 추가 사용 예시]\n"
                        for u in new_usages:
                            old_details += f"- ...{u}...\n"
                        existing_item["details"] = old_details
                        extracted_count += 1
                    
        if extracted_count > 0:
            self.db_manager.save_data()
            self.refresh_list()
            messagebox.showinfo("추출 완료", f"성공적으로 {extracted_count}개의 규격을 신규 등록/업데이트했습니다!")
            self.notebook.select(self.tab_code)
        else:
            if found_codes:
                messagebox.showinfo("추출 완료", f"문서에서 {len(found_codes)}개의 규격을 찾았으나, 모두 이미 등록되어 있습니다.")
            else:
                messagebox.showinfo("추출 실패", "문서에서 규격 코드를 찾지 못했습니다.")

    def create_batch_widgets(self):
        file_frame = ttk.LabelFrame(self.tab_batch, text="1. 원본 절차서 파일 선택 (여러 파일 동시 선택 가능)")
        file_frame.pack(fill="x", padx=15, pady=10)
        
        self.file_listbox = tk.Listbox(file_frame, height=4, selectmode=tk.EXTENDED)
        self.file_listbox.pack(side="left", padx=10, pady=10, expand=True, fill="both")
        
        scrollbar = ttk.Scrollbar(self.file_listbox, orient="vertical")
        scrollbar.config(command=self.file_listbox.yview)
        scrollbar.pack(side="right", fill="y")
        self.file_listbox.config(yscrollcommand=scrollbar.set)
        
        btn_frame = ttk.Frame(file_frame)
        btn_frame.pack(side="right", padx=10, pady=10)
        ttk.Button(btn_frame, text="파일 추가", command=self.add_files).pack(fill="x", pady=2)
        ttk.Button(btn_frame, text="선택 삭제", command=self.remove_file).pack(fill="x", pady=2)
        ttk.Button(btn_frame, text="미리보기(Text)", command=self.preview_file).pack(fill="x", pady=2)
        
        list_frame = ttk.LabelFrame(self.tab_batch, text="2. 단어/코드 일괄 자동 변환 (다중 파일 동시 적용 가능)")
        list_frame.pack(fill="both", expand=True, padx=15, pady=5)
        
        preset_frame = ttk.Frame(list_frame)
        preset_frame.pack(fill="x", padx=10, pady=5)
        ttk.Label(preset_frame, text="자주 쓰는 단어 세트를 저장하고 불러올 수 있습니다.").pack(side="left")
        ttk.Button(preset_frame, text="💾 현재 목록 저장", command=self.save_preset).pack(side="right", padx=2)
        ttk.Button(preset_frame, text="📂 목록 불러오기", command=self.load_preset).pack(side="right", padx=2)
        ttk.Button(preset_frame, text="📚 코드 DB에서 최신 규격 끌어오기", command=self.load_from_code_db).pack(side="right", padx=10)
        
        columns = ("find", "replace")
        self.batch_tree = ttk.Treeview(list_frame, columns=columns, show="headings", height=6)
        self.batch_tree.heading("find", text="찾을 내용 (기존 코드/문구)")
        self.batch_tree.heading("replace", text="바꿀 내용 (새로운 코드/문구)")
        self.batch_tree.column("find", width=300)
        self.batch_tree.column("replace", width=300)
        self.batch_tree.pack(fill="both", expand=True, padx=10, pady=5)
        
        input_frame = ttk.Frame(list_frame)
        input_frame.pack(fill="x", padx=10, pady=5)
        
        ttk.Label(input_frame, text="찾을 내용:").grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.batch_entry_find = ttk.Entry(input_frame, width=20)
        self.batch_entry_find.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(input_frame, text="바꿀 내용:").grid(row=0, column=2, padx=5, pady=5, sticky="e")
        self.batch_entry_replace = ttk.Entry(input_frame, width=20)
        self.batch_entry_replace.grid(row=0, column=3, padx=5, pady=5)
        
        self.batch_entry_replace.bind("<Return>", lambda e: self.batch_add_item())
        self.batch_entry_find.bind("<Return>", lambda e: self.batch_entry_replace.focus())
        
        ttk.Button(input_frame, text="추가", command=self.batch_add_item, width=8).grid(row=0, column=4, padx=5)
        ttk.Button(input_frame, text="수정", command=self.batch_update_item, width=8).grid(row=0, column=5, padx=5)
        ttk.Button(input_frame, text="삭제", command=self.batch_delete_item, width=8).grid(row=0, column=6, padx=5)
        
        self.batch_tree.bind("<<TreeviewSelect>>", self.batch_on_tree_select)
        
        run_frame = ttk.Frame(list_frame)
        run_frame.pack(fill="x", padx=10, pady=10)
        ttk.Button(run_frame, text="위 목록대로 1번에 등록된 모든 파일을 일괄 변환하여 폴더에 자동 저장", command=self.process_files).pack(fill="x", ipady=10)

    def add_files(self):
        filepaths = filedialog.askopenfilenames(
            title="절차서 파일 선택 (여러 개 선택 가능)",
            filetypes=[
                ("모든 지원 파일", "*.docx *.xlsx *.hwp *.hwpx *.txt"),
                ("모든 파일", "*.*")
            ]
        )
        for path in filepaths:
            if path not in self.file_listbox.get(0, tk.END):
                self.file_listbox.insert(tk.END, path)

    def remove_file(self):
        selected = self.file_listbox.curselection()
        for index in reversed(selected):
            self.file_listbox.delete(index)

    def preview_file(self):
        selected = self.file_listbox.curselection()
        if not selected:
            messagebox.showwarning("선택 오류", "미리보기할 파일을 위 목록에서 먼저 선택해주세요.")
            return
            
        filepath = self.file_listbox.get(selected[0])
        ext = os.path.splitext(filepath)[1].lower()
        content = ""
        try:
            if ext == '.docx':
                try:
                    import docx
                    doc = docx.Document(filepath)
                    content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                except ImportError:
                    content = "(docx 모듈이 설치되어 있지 않습니다.)"
            elif ext == '.txt':
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        content = f.read()
                except:
                    with open(filepath, 'r', encoding='euc-kr') as f:
                        content = f.read()
            else:
                messagebox.showinfo("안내", f"{ext} 파일은 텍스트 미리보기를 지원하지 않습니다.")
                return
                
            if not content.strip():
                content = "(추출된 텍스트가 없습니다. 문서가 비어있거나 스캔 이미지 형태일 수 있습니다.)"
                
            top = tk.Toplevel(self.root)
            top.title(f"텍스트 미리보기 - {os.path.basename(filepath)}")
            top.geometry("600x600")
            
            top_frame = ttk.Frame(top)
            top_frame.pack(fill="both", expand=True, padx=10, pady=10)
            
            txt_widget = tk.Text(top_frame, wrap="word", font=("맑은 고딕", 10))
            scrollbar = ttk.Scrollbar(top_frame, orient="vertical", command=txt_widget.yview)
            scrollbar.pack(side="right", fill="y")
            txt_widget.pack(side="left", fill="both", expand=True)
            txt_widget.config(yscrollcommand=scrollbar.set)
            
            txt_widget.insert("1.0", content)
            txt_widget.config(state="disabled")
            
        except Exception as e:
            messagebox.showerror("오류", f"미리보기를 불러오는 중 오류가 발생했습니다:\\n{e}")

    def save_preset(self):
        items = self.batch_tree.get_children()
        if not items:
            messagebox.showwarning("경고", "저장할 단어 목록이 없습니다.")
            return
        preset_data = []
        for item in items:
            val = self.batch_tree.item(item, 'values')
            preset_data.append({"find": val[0], "replace": val[1]})
            
        filepath = filedialog.asksaveasfilename(
            title="단어 목록 저장",
            defaultextension=".json",
            filetypes=[("JSON 파일", "*.json")],
            initialfile="자주쓰는단어.json"
        )
        if filepath:
            try:
                with open(filepath, 'w', encoding='utf-8') as f:
                    json.dump(preset_data, f, ensure_ascii=False, indent=4)
                messagebox.showinfo("저장 완료", "단어 목록이 성공적으로 저장되었습니다.")
            except Exception as e:
                messagebox.showerror("오류", f"저장 중 오류 발생:\\n{e}")

    def load_preset(self):
        filepath = filedialog.askopenfilename(
            title="단어 목록 불러오기",
            filetypes=[("JSON 파일", "*.json")]
        )
        if filepath:
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    preset_data = json.load(f)
                for item in self.batch_tree.get_children():
                    self.batch_tree.delete(item)
                for entry in preset_data:
                    self.batch_tree.insert("", "end", values=(entry.get("find", ""), entry.get("replace", "")))
                messagebox.showinfo("불러오기 완료", "단어 목록을 성공적으로 불러왔습니다.")
            except Exception as e:
                messagebox.showerror("오류", f"불러오기 중 오류 발생:\\n{e}")

    def load_from_code_db(self):
        rules = [(d["find"], d["replace"]) for d in self.db_manager.data if d.get("find") and d.get("replace")]
        if not rules:
            messagebox.showwarning("경고", "코드 관리 DB에 '바꿀 내용'이 설정된 규격이 없습니다.")
            return
            
        if messagebox.askyesno("불러오기 확인", f"코드 DB에 저장된 {len(rules)}개의 변환 규칙을 목록에 가져오시겠습니까?"):
            for item in self.batch_tree.get_children():
                self.batch_tree.delete(item)
            for f_text, r_text in rules:
                self.batch_tree.insert("", "end", values=(f_text, r_text))
            messagebox.showinfo("불러오기 완료", "성공적으로 코드 DB에서 규칙을 불러왔습니다!")

    def batch_add_item(self):
        f_text = self.batch_entry_find.get().strip()
        r_text = self.batch_entry_replace.get().strip()
        if f_text:
            self.batch_tree.insert("", "end", values=(f_text, r_text))
            self.batch_entry_find.delete(0, tk.END)
            self.batch_entry_replace.delete(0, tk.END)
            self.batch_entry_find.focus()
        else:
            messagebox.showwarning("입력 오류", "찾을 내용을 입력해주세요.")
            
    def batch_update_item(self):
        selected = self.batch_tree.selection()
        if not selected:
            messagebox.showwarning("선택 오류", "수정할 항목을 위 목록에서 선택해주세요.")
            return
        f_text = self.batch_entry_find.get().strip()
        r_text = self.batch_entry_replace.get().strip()
        if f_text:
            self.batch_tree.item(selected[0], values=(f_text, r_text))
        else:
            messagebox.showwarning("입력 오류", "찾을 내용을 입력해주세요.")
            
    def batch_on_tree_select(self, event):
        selected = self.batch_tree.selection()
        if selected:
            item = self.batch_tree.item(selected[0])
            val = item['values']
            self.batch_entry_find.delete(0, tk.END)
            self.batch_entry_find.insert(0, val[0])
            self.batch_entry_replace.delete(0, tk.END)
            self.batch_entry_replace.insert(0, val[1])

    def batch_delete_item(self):
        selected = self.batch_tree.selection()
        if selected:
            for item in selected:
                self.batch_tree.delete(item)

    def process_files(self):
        files = self.file_listbox.get(0, tk.END)
        if not files:
            messagebox.showerror("오류", "변환할 파일을 먼저 추가해주세요.")
            return
            
        items = self.batch_tree.get_children()
        if not items:
            messagebox.showwarning("경고", "변경할 단어 목록이 비어있습니다.")
            return
            
        replacements = []
        for item in items:
            val = self.batch_tree.item(item, 'values')
            replacements.append((val[0], val[1]))
            
        output_dir = filedialog.askdirectory(title="변환된 새 파일들을 저장할 폴더 선택")
        if not output_dir:
            return
            
        try:
            self.root.config(cursor="wait")
            self.root.update()
            
            success_count = 0
            for input_file in files:
                if not os.path.exists(input_file): continue
                filename = os.path.basename(input_file)
                output_file = os.path.join(output_dir, f"일괄변환_{filename}")
                
                try:
                    DocumentProcessor.process_single_document(input_file, output_file, replacements)
                    success_count += 1
                except Exception as e:
                    print(f"파일 변환 실패: {filename}, {e}")
                    continue
                    
            self.root.config(cursor="")
            messagebox.showinfo("완료", f"총 {success_count}개의 파일이 성공적으로 일괄 변환 및 저장되었습니다!\n\n저장 폴더: {output_dir}")
            os.startfile(output_dir)
            
        except Exception as e:
            self.root.config(cursor="")
            messagebox.showerror("실행 오류", f"파일을 자동 변환하는 중 오류가 발생했습니다:\n{e}")

if __name__ == "__main__":
    root = tk.Tk()
    app = CodebookApp(root)
    root.mainloop()
