import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
try:
    import docx
except ImportError:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

class DocxEditorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("PAUT 절차서(.docx) 코드/문구 일괄 수정 프로그램")
        self.root.geometry("650x550")
        
        style = ttk.Style()
        style.theme_use('clam')
        
        # UI 구성
        self.create_widgets()
        
    def create_widgets(self):
        # 1. 파일 선택
        file_frame = ttk.LabelFrame(self.root, text="1. 원본 PAUT 절차서(.docx) 선택")
        file_frame.pack(fill="x", padx=15, pady=10)
        
        self.file_path_var = tk.StringVar()
        ttk.Entry(file_frame, textvariable=self.file_path_var, width=50).pack(side="left", padx=10, pady=10, expand=True, fill="x")
        ttk.Button(file_frame, text="파일 찾기", command=self.browse_file).pack(side="right", padx=10, pady=10)
        
        # 2. 변경할 단어 목록
        list_frame = ttk.LabelFrame(self.root, text="2. 변경할 코드 및 문구 목록 (찾을 내용 -> 바꿀 내용)")
        list_frame.pack(fill="both", expand=True, padx=15, pady=5)
        
        columns = ("find", "replace")
        self.tree = ttk.Treeview(list_frame, columns=columns, show="headings", height=8)
        self.tree.heading("find", text="찾을 내용 (기존 코드/문구)")
        self.tree.heading("replace", text="바꿀 내용 (새로운 코드/문구)")
        self.tree.column("find", width=250)
        self.tree.column("replace", width=250)
        self.tree.pack(fill="both", expand=True, padx=10, pady=10)
        
        # 추가/삭제 입력부
        input_frame = ttk.Frame(list_frame)
        input_frame.pack(fill="x", padx=10, pady=5)
        
        ttk.Label(input_frame, text="찾을 내용:").grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.entry_find = ttk.Entry(input_frame, width=20)
        self.entry_find.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(input_frame, text="바꿀 내용:").grid(row=0, column=2, padx=5, pady=5, sticky="e")
        self.entry_replace = ttk.Entry(input_frame, width=20)
        self.entry_replace.grid(row=0, column=3, padx=5, pady=5)
        
        ttk.Button(input_frame, text="목록에 추가", command=self.add_item).grid(row=0, column=4, padx=10)
        ttk.Button(input_frame, text="선택 삭제", command=self.delete_item).grid(row=0, column=5, padx=5)
        
        # 기본 예시 추가
        self.tree.insert("", "end", values=("ASME Sec.V", "ISO 10863"))
        self.tree.insert("", "end", values=("기존 프로젝트명", "가산~가평 천연가스 공급시설"))
        
        # 3. 저장 및 실행
        run_frame = ttk.Frame(self.root)
        run_frame.pack(fill="x", padx=15, pady=15)
        
        ttk.Button(run_frame, text="수정된 내용으로 새로운 파일 저장", command=self.process_file, width=40).pack(pady=10, ipady=10)

    def browse_file(self):
        filepath = filedialog.askopenfilename(
            title="절차서 파일 선택",
            filetypes=[("Word 문서", "*.docx")]
        )
        if filepath:
            self.file_path_var.set(filepath)
            
    def add_item(self):
        f_text = self.entry_find.get().strip()
        r_text = self.entry_replace.get().strip()
        if f_text:
            self.tree.insert("", "end", values=(f_text, r_text))
            self.entry_find.delete(0, tk.END)
            self.entry_replace.delete(0, tk.END)
        else:
            messagebox.showwarning("입력 오류", "찾을 내용을 입력해주세요.")
            
    def delete_item(self):
        selected = self.tree.selection()
        if selected:
            for item in selected:
                self.tree.delete(item)
                
    def replace_text_in_paragraph(self, paragraph, find_text, replace_text):
        if find_text in paragraph.text:
            # Word 문서는 하나의 문단이 여러개의 'run'으로 쪼개져 있어 단순 replace가 까다롭습니다.
            # 가장 안전한 방법은 문단의 전체 텍스트를 교체하고 기존 서식을 첫번째 run에 몰아주는 것입니다.
            # (복잡한 서식이 섞여있다면 일부 서식이 풀릴 수 있으나 내용 변환에는 가장 확실합니다)
            inline = paragraph.runs
            if not inline:
                return
                
            text = paragraph.text.replace(find_text, replace_text)
            
            # 문단 내 텍스트 모두 지우기
            for i in range(len(inline)):
                inline[i].text = ''
                
            # 첫번째 run에 바뀐 텍스트 넣기
            inline[0].text = text

    def process_file(self):
        input_file = self.file_path_var.get()
        if not input_file or not os.path.exists(input_file):
            messagebox.showerror("오류", "유효한 원본 파일을 선택해주세요.")
            return
            
        items = self.tree.get_children()
        if not items:
            messagebox.showwarning("경고", "변경할 단어 목록이 비어있습니다.")
            return
            
        replacements = []
        for item in items:
            val = self.tree.item(item, 'values')
            replacements.append((val[0], val[1]))
            
        # 저장 경로 묻기
        output_file = filedialog.asksaveasfilename(
            title="저장할 파일 이름 설정",
            defaultextension=".docx",
            filetypes=[("Word 문서", "*.docx")],
            initialfile="수정된_PAUT_절차서.docx"
        )
        if not output_file:
            return
            
        try:
            doc = docx.Document(input_file)
            
            # 문단 내용 바꾸기
            for paragraph in doc.paragraphs:
                for f_text, r_text in replacements:
                    self.replace_text_in_paragraph(paragraph, f_text, r_text)
            
            # 표(Table) 안의 내용 바꾸기
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for f_text, r_text in replacements:
                                self.replace_text_in_paragraph(paragraph, f_text, r_text)
                                
            doc.save(output_file)
            messagebox.showinfo("완료", f"파일이 성공적으로 수정 및 저장되었습니다!\n\n저장 위치: {output_file}")
            
        except Exception as e:
            messagebox.showerror("실행 오류", f"파일을 수정하는 중 오류가 발생했습니다:\n{e}")

if __name__ == "__main__":
    root = tk.Tk()
    app = DocxEditorApp(root)
    root.mainloop()
