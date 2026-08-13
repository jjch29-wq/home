import docx

file_path = r'C:\Users\-\OneDrive\바탕 화면\PAUT_SIS-H-264 (2024ed)KI rev.0(최종)_코멘트반영.docx'
doc = docx.Document(file_path)

for i, para in enumerate(doc.paragraphs):
    if "The demonstration block's weld joint geometry shall be representative of the production joint's" in para.text:
        para.text = "The block shall be representative of the component to be inspected. The weld joint geometry shall be representative of the production joint's details. The weld caps conditions being the same as those encountered for production weld testing."
        
        kr_para = doc.paragraphs[i+1]
        if '검증 시험편의 용접접합부 형상은' in kr_para.text:
            kr_para.text = "(검증 시험편은 검사 대상 부품을 대표해야 한다. 용접 접합부의 기하학적 형상은 생산 접합부의 세부 사항을 대표해야 한다. 용접 덧살(캡) 상태는 생산 용접 검사 시 마주치는 것과 동일해야 한다.)"
            kr_para.style = para.style
            
        break

output_path = r'C:\Users\-\OneDrive\바탕 화면\PAUT_SIS-H-264 (2024ed)KI rev.0(최종)_코멘트2반영.docx'
doc.save(output_path)
print("Saved to:", output_path)
